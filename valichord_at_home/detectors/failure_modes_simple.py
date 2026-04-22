"""
ValiChord Auto-Generate
Simple (pattern-matching) failure mode detectors
Implements failure modes from ValiChord Specification v15
"""
from __future__ import annotations

import re
from pathlib import Path


# ── file classification helpers ──────────────────────────────────────────────

CODE_EXTENSIONS = {
    '.py', '.r', '.rmd', '.qmd', '.jl', '.m', '.sh', '.bash', '.smk', '.nf', '.groovy',
    '.do', '.sas', '.ado', '.c', '.cpp', '.f', '.f90',
    '.sql', '.rs', '.go', '.java', '.js', '.ts',
    '.ipynb',   # Jupyter notebooks — treat as code for has_code checks
}

NOTEBOOK_EXTENSIONS = {'.ipynb', '.mlx', '.rmd', '.qmd'}

DATA_EXTENSIONS = {
    '.csv', '.tsv', '.xlsx', '.xls', '.json', '.jsonl', '.ndjson',
    '.parquet', '.feather', '.arrow',
    '.rds', '.rda', '.rdata', '.dta', '.sav', '.por', '.zsav', '.sas7bdat', '.xpt',
    '.mat', '.pkl', '.npy', '.npz', '.hdf5', '.h5', '.nc',
    '.dif', '.gdt',
    # Shapefile components
    '.shp', '.dbf', '.shx', '.prj', '.cpg', '.sbn', '.sbx',
    # Geospatial
    '.geojson', '.gpkg', '.kml', '.kmz',
    # .gdb is an ESRI directory format — won't match f.is_file() but included for completeness
}

ARCHIVE_EXTENSIONS = {'.zip', '.rar', '.7z', '.tar', '.gz', '.tgz', '.bz2'}

ENCRYPTED_EXTENSIONS = {'.gpg', '.enc', '.secret', '.age', '.asc'}

DEPENDENCY_FILES = {
    'requirements.txt', 'requirements_extra.txt', 'environment.yml', 'environment.yaml',
    'pipfile', 'pipfile.lock', 'poetry.lock', 'setup.py',
    'pyproject.toml', 'setup.cfg', 'conda-lock.yml',
    'description', 'renv.lock', 'packrat.lock',
    'cargo.toml', 'cargo.lock', 'go.mod', 'go.sum',
    'package.json', 'package-lock.json', 'yarn.lock',
    'pom.xml', 'build.gradle',
    'project.toml', 'manifest.toml',
    'manifest-v1.6.toml', 'manifest-v1.7.toml',
    'manifest-v1.8.toml', 'manifest-v1.9.toml',
    'manifest-v1.10.toml', 'manifest-v1.11.toml',

}

README_NAMES = {'readme.md', 'readme.txt', 'readme.rst', 'readme'}

# Exact filenames that unambiguously function as codebooks / data dictionaries.
# Used by detect_E (suppression), detect_BA (exclusion), and INVENTORY (type).
CODEBOOK_FILENAMES = {
    'metadata.csv', 'metadata.xlsx', 'metadata.txt',
    'data_dictionary.csv', 'data_dictionary.xlsx',
    'codebook.csv', 'codebook.xlsx', 'codebook.txt',
    'variables.csv', 'variables.txt',
    'column_descriptions.csv', 'field_descriptions.csv',
}


def _looks_like_codebook(path) -> bool:
    """Return True if a CSV/TSV file is structured as a variable codebook.

    Detects delimiter by counting occurrences in the first line (semicolons
    and tabs before comma), then checks that the second column contains
    description-like text (average length > 12 chars).
    """
    try:
        with path.open(encoding='utf-8', errors='ignore') as fh:
            content = fh.read(8192)  # 8 KB is enough to inspect header rows
        lines = [ln for ln in content.split('\n') if ln.strip()][:10]
        if len(lines) < 3:
            return False
        first = lines[0]
        if first.count(';') > first.count(','):
            delim = ';'
        elif first.count('\t') > first.count(','):
            delim = '\t'
        else:
            delim = ','
        rows = [ln.split(delim) for ln in lines]
        col1 = [r[0].strip() for r in rows[1:] if r[0].strip()]
        col2 = [r[1].strip() for r in rows[1:] if len(r) > 1 and r[1].strip()]
        if not col2 or not col1:
            return False
        avg_col1_len = sum(len(c) for c in col1) / len(col1)
        avg_col2_len = sum(len(c) for c in col2) / len(col2)
        # A real codebook has short variable names in col1 (< 35 chars avg)
        # and long descriptions in col2 (> 20 chars avg).  Data CSVs often
        # have both columns long (e.g. compound names + activity values) or
        # col2 short (e.g. numeric measurements), so this combination is a
        # stronger signal than col2 length alone.
        return avg_col1_len < 35 and avg_col2_len > 20
    except Exception as e:
        import sys
        print(
            f'[ValiChord WARNING] _looks_like_codebook failed on '
            f'{path.name!r}: {type(e).__name__}: {e}',
            file=sys.stderr,
        )
        return False


CODEBOOK_SHEET_NAMES = {
    'variables', 'variable list', 'list of variables', 'data dictionary',
    'codebook', 'code book', 'metadata', 'column descriptions',
    'variable descriptions', 'field descriptions', 'legend',
}


def _xlsx_has_codebook_sheet(path) -> str | None:
    """Return the sheet name if any sheet in an xlsx file looks like a variable codebook, else None."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        for sheet_name in wb.sheetnames:
            if sheet_name.lower().strip() in CODEBOOK_SHEET_NAMES:
                wb.close()
                return sheet_name
            ws = wb[sheet_name]
            headers = [
                str(c).lower().strip() for c in
                next(ws.iter_rows(min_row=1, max_row=1, values_only=True), [])
                if c is not None
            ]
            if 'variable' in headers and any(
                kw in headers for kw in ('description', 'abbr', 'label', 'definition')
            ):
                wb.close()
                return sheet_name
        wb.close()
    except Exception as e:
        import sys
        print(
            f'[ValiChord WARNING] _xlsx_has_codebook_sheet failed on '
            f'{path.name!r}: {type(e).__name__}: {e}',
            file=sys.stderr,
        )
        return None
    return None


# ── centralised content inspection ───────────────────────────────────────────


class FileInspectionResult:
    """Structured result of inspecting the content of an opaque file format.

    Fields
    ------
    extracted_text : str | None
        Full text extracted from the file (docx paragraphs, PDF pages, …).
    variable_labels : dict | None
        Mapping varname → label string (stat files, netCDF, HDF5).
    sheet_names : list | None
        Ordered list of sheet names (xlsx only).
    has_codebook : bool
        True when inspection evidence indicates a variable codebook or
        sufficient variable labels to suppress [E].
    inspection_note : str | None
        Human-readable note for CLEANING_REPORT positive observations,
        or a brief failure message if reading failed.
    """
    __slots__ = ('extracted_text', 'variable_labels', 'sheet_names',
                 'has_codebook', 'inspection_note')

    def __init__(self, *, extracted_text=None, variable_labels=None,
                 sheet_names=None, has_codebook=False, inspection_note=None):
        self.extracted_text = extracted_text
        self.variable_labels = variable_labels
        self.sheet_names = sheet_names
        self.has_codebook = has_codebook
        self.inspection_note = inspection_note


# Cache keyed on resolved absolute path — shared across all detector calls
# within a single run so each file is opened at most once.
_file_inspection_cache: dict = {}


def _inspect_file_content(path) -> FileInspectionResult:
    """Return a FileInspectionResult for the given file, with per-run caching.

    Dispatches to a format-specific inspector based on suffix. Any format
    that cannot be read falls back gracefully — has_codebook=False and
    inspection_note records why. Callers must NOT suppress findings based
    on a failed read.
    """
    key = path.resolve()
    cached = _file_inspection_cache.get(key)
    if cached is not None:
        return cached
    ext = path.suffix.lower()
    if ext == '.docx':
        result = _inspect_docx(path)
    elif ext in {'.xlsx', '.xls'}:
        result = _inspect_xlsx_content(path)
    elif ext in {'.dta', '.sav', '.zsav', '.sas7bdat', '.xpt'}:
        result = _inspect_stat_file(path)
    elif ext == '.pdf':
        result = _inspect_pdf_file(path)
    elif ext == '.nc':
        result = _inspect_netcdf_file(path)
    elif ext in {'.h5', '.hdf5'}:
        result = _inspect_hdf5_file(path)
    elif ext in {'.sqlite', '.db', '.sqlite3'}:
        result = _inspect_sqlite_file(path)
    else:
        result = FileInspectionResult()
    _file_inspection_cache[key] = result
    return result


def _inspect_docx(path) -> FileInspectionResult:
    """Extract text and detect codebook structure from a .docx file."""
    try:
        import docx as _docx
        doc = _docx.Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # Inspect tables for variable+description structure
        has_var_desc_table = False
        for tbl in doc.tables:
            if tbl.rows:
                hdrs = [cell.text.lower().strip() for cell in tbl.rows[0].cells]
                if 'variable' in hdrs and any(
                    kw in hdrs for kw in ('description', 'label', 'definition', 'abbr')
                ):
                    has_var_desc_table = True
                for row in tbl.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)
        full_text = '\n'.join(parts)
        lower = full_text.lower()
        _CB_PHRASES = {'variables', 'variable list', 'codebook', 'data dictionary',
                       'column descriptions', 'variable descriptions', 'field descriptions'}
        has_codebook = has_var_desc_table or any(ph in lower for ph in _CB_PHRASES)
        return FileInspectionResult(
            extracted_text=full_text,
            has_codebook=has_codebook,
            inspection_note=(
                f'Variable codebook content detected in: {path.name}' if has_codebook else None
            ),
        )
    except Exception as e:
        return FileInspectionResult(
            inspection_note=f'docx read failed ({type(e).__name__}: {e})'
        )


def _inspect_xlsx_content(path) -> FileInspectionResult:
    """Inspect xlsx/xls for codebook sheets (delegates to _xlsx_has_codebook_sheet)."""
    sheet = _xlsx_has_codebook_sheet(path)  # str (matched name) or None
    if sheet:
        return FileInspectionResult(
            has_codebook=True,
            inspection_note=f'Variable codebook detected in: {path.name} (sheet: "{sheet}")',
        )
    # Try to return sheet names even when no codebook detected
    try:
        import openpyxl as _opxl
        wb = _opxl.load_workbook(str(path), read_only=True, data_only=True)
        snames = list(wb.sheetnames)
        wb.close()
        return FileInspectionResult(sheet_names=snames)
    except Exception as e:
        import sys
        print(
            f'[ValiChord WARNING] _inspect_xlsx_content (sheet names) failed on '
            f'{path.name!r}: {type(e).__name__}: {e}',
            file=sys.stderr,
        )
        return FileInspectionResult()


def _inspect_stat_file(path) -> FileInspectionResult:
    """Read variable labels from Stata / SPSS / SAS files via pyreadstat."""
    try:
        import pyreadstat as _prs
        _READERS = {
            '.dta':      _prs.read_dta,
            '.sav':      _prs.read_sav,
            '.zsav':     _prs.read_sav,
            '.sas7bdat': _prs.read_sas7bdat,
            '.xpt':      _prs.read_xport,
        }
        reader = _READERS.get(path.suffix.lower())
        if reader is None:
            return FileInspectionResult()
        _, meta = reader(str(path), metadataonly=True)
        col_names  = list(getattr(meta, 'column_names',  []) or [])
        col_labels = list(getattr(meta, 'column_labels', []) or [])
        labels = {
            col: str(lbl).strip()
            for col, lbl in zip(col_names, col_labels)
            if lbl and str(lbl).strip()
        }
        n_vars     = len(col_names)
        n_labelled = len(labels)
        has_codebook = n_vars > 0 and (n_labelled / n_vars) > 0.5
        note = (
            f'Variable labels found in {path.name}: '
            f'{n_labelled} of {n_vars} variables labelled'
        ) if n_vars > 0 else None
        return FileInspectionResult(
            variable_labels=labels or None,
            has_codebook=has_codebook,
            inspection_note=note,
        )
    except Exception as e:
        return FileInspectionResult(
            inspection_note=f'stat file read failed ({type(e).__name__}: {e})'
        )


def _inspect_pdf_file(path) -> FileInspectionResult:
    """Extract text from a PDF via pdfplumber; skip gracefully if image-only."""
    try:
        import pdfplumber as _pp
        parts = []
        with _pp.open(str(path)) as pdf:
            for page in pdf.pages[:20]:   # cap at 20 pages for performance
                t = page.extract_text()
                if t:
                    parts.append(t)
        full_text = '\n'.join(parts)
        if not full_text.strip():
            return FileInspectionResult(
                inspection_note=f'PDF appears image-only (no extractable text): {path.name}'
            )
        lower = full_text.lower()
        _CB_PHRASES = {'variables', 'variable list', 'codebook', 'data dictionary',
                       'column descriptions', 'variable descriptions', 'field descriptions'}
        has_codebook = any(ph in lower for ph in _CB_PHRASES)
        return FileInspectionResult(
            extracted_text=full_text,
            has_codebook=has_codebook,
            inspection_note=(
                f'Variable codebook content detected in: {path.name}' if has_codebook else None
            ),
        )
    except Exception as e:
        return FileInspectionResult(
            inspection_note=f'PDF read failed ({type(e).__name__}: {e})'
        )


def _inspect_netcdf_file(path) -> FileInspectionResult:
    """Read long_name / units variable metadata from a netCDF4 file."""
    try:
        import netCDF4 as _nc4
        labels: dict = {}
        with _nc4.Dataset(str(path), 'r') as ds:
            for vname, var in ds.variables.items():
                ln = getattr(var, 'long_name', None)
                if ln and str(ln).strip():
                    labels[vname] = str(ln).strip()
            n_vars = len(ds.variables)
        n_labelled = len(labels)
        has_codebook = n_vars > 0 and (n_labelled / n_vars) > 0.5
        note = (
            f'Variable metadata found in {path.name}: '
            f'{n_labelled} of {n_vars} variables with long_name'
        ) if n_vars > 0 else None
        return FileInspectionResult(
            variable_labels=labels or None,
            has_codebook=has_codebook,
            inspection_note=note,
        )
    except Exception as e:
        return FileInspectionResult(
            inspection_note=f'netCDF read failed ({type(e).__name__}: {e})'
        )


def _inspect_hdf5_file(path) -> FileInspectionResult:
    """Read long_name attributes from HDF5 datasets via h5py."""
    try:
        import h5py as _h5
        labels: dict = {}

        def _visit(name, obj):
            if isinstance(obj, _h5.Dataset):
                ln = obj.attrs.get('long_name')
                if ln is not None:
                    labels[name] = str(ln)

        with _h5.File(str(path), 'r') as hf:
            hf.visititems(_visit)
            n_top = len(hf)
        n_labelled = len(labels)
        has_codebook = n_labelled > 0 and (n_labelled >= 0.5 * n_top if n_top else False)
        note = (
            f'Variable metadata found in {path.name}: '
            f'{n_labelled} datasets with long_name attribute'
        ) if n_labelled > 0 else None
        return FileInspectionResult(
            variable_labels=labels or None,
            has_codebook=has_codebook,
            inspection_note=note,
        )
    except Exception as e:
        return FileInspectionResult(
            inspection_note=f'HDF5 read failed ({type(e).__name__}: {e})'
        )


def _inspect_sqlite_file(path) -> FileInspectionResult:
    """Check SQLite DBs for codebook tables or meaningful (non-generic) column names."""
    try:
        import sqlite3 as _sql3
        _CB_TABLES = frozenset({
            'metadata', 'codebook', 'variables', 'data_dictionary',
            'column_descriptions', 'field_descriptions', 'variable_descriptions',
        })
        _GENERIC_COL = re.compile(r'^(col|field|column|f)\d+$', re.IGNORECASE)
        with _sql3.connect(str(path)) as conn:
            cur = conn.cursor()
            cur.execute("SELECT name FROM sqlite_master WHERE type='table'")
            tables = [row[0] for row in cur.fetchall()]
        if not tables:
            return FileInspectionResult()
        has_codebook_table = bool({t.lower().replace(' ', '_') for t in tables} & _CB_TABLES)
        all_meaningful = True
        with _sql3.connect(str(path)) as conn:
            cur = conn.cursor()
            for tbl in tables:
                cur.execute(f'PRAGMA table_info("{tbl}")')
                cols = [row[1] for row in cur.fetchall()]
                if not cols or all(_GENERIC_COL.match(c) for c in cols):
                    all_meaningful = False
                    break
        has_codebook = has_codebook_table or (bool(tables) and all_meaningful)
        note = None
        if has_codebook_table:
            note = f'Codebook table detected in {path.name} (tables: {", ".join(tables)})'
        elif has_codebook:
            note = f'Meaningful column names in {path.name} (tables: {", ".join(tables)})'
        return FileInspectionResult(has_codebook=has_codebook, inspection_note=note)
    except Exception as e:
        return FileInspectionResult(
            inspection_note=f'SQLite read failed ({type(e).__name__}: {e})'
        )


LICENCE_NAMES = {
    'licence', 'license', 'licence.md', 'license.md',
    'licence.txt', 'license.txt'
}

# Directories that contain third-party / vendored code — excluded from
# researcher-code detectors (B, U, AC, AI, AJ, AK, AS, AT, AU, AA, BB …)
VENDOR_DIRS = {'weka', 'vendor', 'lib', 'dist', 'node_modules', 'target'}

# Directories whose contents are input assets — stimuli, fonts, icons, etc.
# Files inside these directories are not research data or generated outputs and
# must not trigger [Y], [E], or [AA].  Path check covers any depth:
#   images/Training/001.jpeg  → excluded by 'images' in parts
#   static/fonts/glyph.woff  → excluded by 'fonts' in parts
_ASSET_INPUT_DIRS = frozenset({
    'images', 'img',
    'stimuli', 'stimulus',
    'audio', 'video',
    'fonts', 'font',
    'assets', 'static',
    'media', 'resources',
    'materials', 'icons', 'icon',
    'maps', 'tiles',
})


def _in_asset_dir(f, repo_dir) -> bool:
    """Return True if this file lives inside an asset input directory."""
    try:
        parts = f.relative_to(repo_dir).parts
    except ValueError:
        parts = f.parts
    # Check all parts except the filename itself
    return any(part.lower() in _ASSET_INPUT_DIRS for part in parts[:-1])

# R package library directories — contain installed third-party package code,
# not researcher-authored scripts.  When a deposit commits renv/library/ or
# packrat/lib/, those R files must not be scanned for researcher practices
# (sessionInfo, library() calls, source() chains, etc.) because the third-party
# code triggers the same patterns as researcher scripts, causing false suppression.
# Matches the top-two-parts rule: renv/library/<pkg>/…  packrat/lib/<pkg>/…
_R_PKG_LIB_ROOTS = frozenset({'renv', 'packrat'})
_R_PKG_LIB_SUBDIRS = frozenset({'library', 'lib', 'src'})

# Minified / bundled frontend assets — excluded from all researcher-code detectors.
# .min.js / .min.css are double-suffix patterns; MINIFIED_FILE_STEMS covers
# common webpack chunk/bundle stems.
MINIFIED_FILE_EXTENSIONS = ('.min.js', '.min.css')
MINIFIED_FILE_STEMS = frozenset({'lib.min', 'vendor.min', 'bundle.min'})

# Frontend-directory detection.
# A directory is a "frontend dir" if it contains at least one frontend-marker
# extension AND no analysis-code extension.  We do NOT require all extensions
# to be in an allowlist — the directory may contain .sb3, .wav, .mp3, .zip,
# or other assets we haven't enumerated.
_FRONTEND_MARKERS = frozenset({'.js', '.html', '.css', '.jsx', '.ts', '.tsx', '.vue'})
_ANALYSIS_CODE_EXTS = frozenset({
    '.py', '.r', '.do', '.jl', '.m', '.sas', '.ipynb',
    '.rmd', '.qmd', '.sh', '.bash', '.f90', '.f', '.cpp', '.c',
    '.java', '.nf', '.smk', '.ado', '.rs', '.go', '.scala', '.sql',
})


def _is_minified(f):
    """Return True if a file appears to be a minified or bundled frontend asset."""
    name_lower = f.name.lower()
    stem_lower = f.stem.lower()
    return (
        any(name_lower.endswith(ext) for ext in MINIFIED_FILE_EXTENSIONS)
        or stem_lower in MINIFIED_FILE_STEMS
        or stem_lower.startswith('chunk')
    )


def _researcher_r_files(all_files, repo_dir):
    """Return R/Rmd/Qmd files that are researcher-authored scripts.

    Only excludes committed R package library trees (renv/library/,
    packrat/lib/, packrat/src/) where the top-two directory parts match
    the known root+subdir pattern.  Does NOT apply VENDOR_DIRS — that
    set includes 'lib' which is a common name for a researcher's own
    helper-function folder (e.g. scripts/lib/utils.R), and excluding it
    would silently drop legitimate code files, breaking the "same traversal
    as all_files" contract that [E] and [BA] satisfy.
    """
    result = []
    for f in all_files:
        if f.suffix.lower() not in {'.r', '.rmd', '.qmd'}:
            continue
        try:
            parts = f.relative_to(repo_dir).parts
        except ValueError:
            result.append(f)
            continue
        # Skip renv/library/<pkg>/…  and  packrat/lib/<pkg>/…
        if (len(parts) >= 2
                and parts[0].lower() in _R_PKG_LIB_ROOTS
                and parts[1].lower() in _R_PKG_LIB_SUBDIRS):
            continue
        result.append(f)
    return result


def _is_frontend_dir(directory):
    """Return True if a directory contains only frontend assets and no analysis code.

    Used to exclude bundled framework directories (React, Scratch, etc.) from
    execution-order and config-file detectors.
    """
    _fe_exts = {'.js', '.html', '.css', '.json', '.svg', '.png', '.gif',
                '.woff', '.ttf', '.woff2', '.eot', '.ico', '.webp'}
    _analysis_exts = {'.py', '.r', '.do', '.jl', '.m', '.sas', '.ipynb',
                      '.rmd', '.qmd', '.sh', '.bash'}
    try:
        files = [f for f in directory.rglob('*') if f.is_file()]
    except Exception:
        return False
    if not files:
        return False
    extensions = {f.suffix.lower() for f in files}
    return extensions.issubset(_fe_exts) and not (extensions & _analysis_exts)


def _inspect_rar(abs_path):
    """Return list of file entry names from a RAR archive, or None on total failure.

    Tries rarfile package first; falls back to the unrar binary via subprocess.
    """
    # Method 1: rarfile package (preferred)
    try:
        import rarfile as _rf
        with _rf.RarFile(abs_path) as rf:
            return [n for n in rf.namelist() if not n.endswith('/')]
    except ImportError:
        pass
    except Exception:
        return []  # file exists but is unreadable — stop here

    # Method 2: unrar binary fallback
    try:
        import subprocess as _sp
        result = _sp.run(
            ['unrar', 'l', '-p-', abs_path],
            capture_output=True, text=True, timeout=15
        )
        if result.returncode == 0:
            # Parse "unrar l" output: file lines contain size + date + name columns
            names = []
            in_list = False
            for line in result.stdout.splitlines():
                if line.startswith('----'):
                    in_list = not in_list
                    continue
                if in_list and line.strip():
                    # Last field is filename; skip directory entries (end with /)
                    parts = line.split()
                    if parts and not parts[-1].endswith('/'):
                        names.append(parts[-1])
            return names
    except (FileNotFoundError, Exception):
        pass

    return None  # both methods failed


def _inspect_archive(path):
    """Return a content-summary string for an archive file.

    Returns a string starting with ' — ', e.g.:
      ' — 123 files (csv, xlsx)'
      ' — contents not inspectable (BadRarName: ...)'
    Uses path.resolve() to ensure an absolute path is passed to archive libs.
    Supports .zip/.tar/.gz/.tgz/.bz2 (stdlib) and .rar (rarfile) / .7z (py7zr).
    """
    ext = path.suffix.lower()
    try:
        abs_path = str(path.resolve())
        if ext == '.zip':
            import zipfile as _zf
            with _zf.ZipFile(abs_path) as zf:
                names = [n for n in zf.namelist() if not n.endswith('/')]
        elif ext == '.rar':
            names = _inspect_rar(abs_path)
            if names is None:
                raise ImportError('rarfile not available and unrar binary not found')
        elif ext in ('.tar', '.gz', '.tgz', '.bz2'):
            import tarfile as _tf
            with _tf.open(abs_path) as tf:
                names = [m.name for m in tf.getmembers() if m.isfile()]
        elif ext == '.7z':
            import py7zr as _7z  # pip install py7zr
            with _7z.SevenZipFile(abs_path, mode='r') as zf:
                names = zf.getnames()
        else:
            return ''
        exts = sorted({n.rsplit('.', 1)[-1].lower() for n in names if '.' in n})
        ext_str = f' ({", ".join(exts)})' if exts else ''
        return f' — {len(names)} files{ext_str}'
    except Exception:
        return ' — contents not inspectable (not a valid archive)'


def _archive_contents_note(path):
    """Thin wrapper — returns _inspect_archive string, defaulting to not-inspectable."""
    note = _inspect_archive(path)
    return note if note else ' — contents not inspectable'


def _is_single_file_compressed(path) -> bool:
    """Return True for single-file compressed formats like .csv.gz, .jsonl.gz.

    Excludes .tar.gz, .tar.bz2 (true nested archives).
    Works on any Path-like object without filesystem access.
    """
    from pathlib import Path as _Path
    p = _Path(path) if not hasattr(path, 'suffix') else path
    if p.suffix.lower() not in {'.gz', '.bz2', '.xz', '.zst'}:
        return False
    inner = _Path(p.stem)   # strip outer compression extension
    return inner.suffix.lower() not in {'.tar', ''}


def finding(mode, severity, title, detail, evidence=None):
    """Create a standardised finding dictionary."""
    return {
        'mode': mode,
        'severity': severity,
        'title': title,
        'detail': detail,
        'evidence': evidence or []
    }


_MAX_READ_BYTES = 2 * 1024 * 1024  # 2 MB cap — prevents OOM on large data files


def read_file_safe(path):
    """Read a file, trying utf-8 then latin-1. Return empty string on failure.

    Capped at 2 MB: READMEs and code files are always smaller; large CSV/binary
    data files need only their first bytes for pattern-matching purposes.
    """
    for encoding in ('utf-8', 'latin-1'):
        try:
            with path.open('rb') as fh:
                raw = fh.read(_MAX_READ_BYTES)
            return raw.decode(encoding)
        except Exception:
            pass
    return ''


# Standard-library top-level module names (Python 3.6+).
# Used to decide whether a script has zero external dependencies.
_STDLIB_TOPLEVEL = frozenset({
    'abc', 'ast', 'asyncio', 'base64', 'binascii', 'builtins',
    'calendar', 'cmath', 'codecs', 'collections', 'concurrent',
    'contextlib', 'copy', 'csv', 'ctypes', 'dataclasses', 'datetime',
    'decimal', 'difflib', 'email', 'enum', 'errno', 'faulthandler',
    'filecmp', 'fnmatch', 'fractions', 'functools', 'gc', 'glob',
    'gzip', 'hashlib', 'heapq', 'hmac', 'html', 'http', 'imaplib',
    'importlib', 'inspect', 'io', 'ipaddress', 'itertools', 'json',
    'keyword', 'linecache', 'locale', 'logging', 'lzma', 'math',
    'mimetypes', 'multiprocessing', 'numbers', 'operator', 'os',
    'pathlib', 'pickle', 'platform', 'pprint', 'queue', 'random',
    're', 'shlex', 'shutil', 'signal', 'socket', 'sqlite3',
    'ssl', 'stat', 'statistics', 'string', 'struct', 'subprocess',
    'sys', 'tarfile', 'tempfile', 'textwrap', 'threading', 'time',
    'timeit', 'tkinter', 'traceback', 'typing', 'unicodedata',
    'unittest', 'urllib', 'uuid', 'warnings', 'weakref', 'xml',
    'xmlrpc', 'zipfile', 'zipimport', 'zlib',
})

_IMPORT_PAT = re.compile(
    r'^\s*(?:import\s+(\w+)|from\s+(\w+)\s+import)',
    re.MULTILINE
)


def _has_only_stdlib_imports(f):
    """Return True if a Python file imports nothing outside the stdlib."""
    if f.suffix.lower() != '.py':
        return False
    modules = [
        m.group(1) or m.group(2)
        for m in _IMPORT_PAT.finditer(read_file_safe(f))
    ]
    return all(m in _STDLIB_TOPLEVEL for m in modules)


_TRIVIAL_NAMES = frozenset({'reader', 'loader', 'parser', 'helper'})


def _is_zero_dep_reader(all_files):
    """Return True when the deposit is a zero-dependency reader/loader helper.

    Conditions (both must hold):
    1. Every non-vendored code file is either <1 KB or named *reader/loader/
       parser/helper* — i.e. there are no substantial analysis scripts.
    2. No external (non-stdlib) imports — checked via requirements_DRAFT.txt
       (authoritative, already computed) when present, or by direct import
       scan for pure-Python deposits on first run.
    """
    code_files = [
        f for f in all_files
        if f.suffix.lower() in CODE_EXTENSIONS
        and not any(p.name.lower() in VENDOR_DIRS for p in f.parents)
    ]
    if not code_files:
        return False

    # Condition 1: every code file is small OR trivially named
    if not all(
        f.stat().st_size < 1024
        or any(kw in f.stem.lower() for kw in _TRIVIAL_NAMES)
        for f in code_files
    ):
        return False

    # Condition 2: no external imports
    # Fast path: requirements_DRAFT.txt from a prior run is authoritative
    for f in all_files:
        if f.name.lower() == 'requirements_draft.txt':
            return 'no external imports detected' in read_file_safe(f).lower()

    # Fallback: scan Python files directly (only valid for pure-Python deposits)
    non_py = [f for f in code_files if f.suffix.lower() != '.py']
    if non_py:
        return False  # non-Python code present; can't determine stdlib-only
    py_files = [f for f in code_files if f.suffix.lower() == '.py']
    return bool(py_files) and all(_has_only_stdlib_imports(f) for f in py_files)


_CODE_TXT_STEM_KEYWORDS = frozenset({
    'code', 'script', 'analysis', 'replication', 'pipeline', 'main', 'run'
})
_CODE_TXT_CONTENT_RE = re.compile(
    r'library\s*\(|import\s+\w|^\s*def\s+\w|\bfunction\s*\(|\bcd\s+|\buse\s+',
    re.MULTILINE
)


def _is_code_txt(f):
    """Return True if a .txt file's stem and content suggest it is actually code."""
    if f.suffix.lower() != '.txt':
        return False
    if not any(kw in f.stem.lower() for kw in _CODE_TXT_STEM_KEYWORDS):
        return False
    return bool(_CODE_TXT_CONTENT_RE.search(read_file_safe(f)))


# ── individual detectors ─────────────────────────────────────────────────────

def detect_A_no_readme(repo_dir, all_files):
    """Failure Mode A: No README or inadequate README."""
    findings = []
    names = {f.name.lower() for f in all_files}

    root_readme = [f for f in all_files if f.name.lower() in README_NAMES and len(f.relative_to(repo_dir).parts) <= 4]
    if not root_readme:
        # .docx README: common pattern deserves its own specific message
        _README_STEMS = ('readme', 'read me', 'read_me')
        _docx_readmes = [
            f for f in all_files
            if f.stem.lower().startswith(_README_STEMS)
            and f.suffix.lower() == '.docx'
            and len(f.relative_to(repo_dir).parts) <= 2
        ]
        # Other proprietary/non-machine-readable formats
        _PROP_EXTS = {'.pdf', '.pages', '.doc'}
        _prop_readmes = [
            f for f in all_files
            if f.stem.lower().startswith(_README_STEMS)
            and f.suffix.lower() in _PROP_EXTS
            and len(f.relative_to(repo_dir).parts) <= 4
        ]
        if _docx_readmes:
            _docx = min(_docx_readmes, key=lambda x: len(x.relative_to(repo_dir).parts))
            findings.append(finding(
                'A', 'SIGNIFICANT',
                f'README found as .docx file ({_docx.name})',
                f'README found as .docx file ({_docx.name}) — consider converting to '
                'README.md or README.txt for better accessibility and plain-text '
                'searchability.',
                [f'Found: {_docx.name}',
                 'Fix: export content to README.md using a suitable conversion '
                 'tool or by copy-pasting the text into a Markdown file']
            ))
        elif _prop_readmes:
            # Deduplicate by filename: keep shallowest copy of each distinct name
            _seen_names = {}
            for f in sorted(_prop_readmes, key=lambda x: len(x.relative_to(repo_dir).parts)):
                if f.name.lower() not in _seen_names:
                    _seen_names[f.name.lower()] = f
            _prop_readmes = list(_seen_names.values())
            _names_str = ', '.join(f.name for f in _prop_readmes)
            _fmts_str = ', '.join(sorted({f.suffix.lower() for f in _prop_readmes}))
            findings.append(finding(
                'A', 'SIGNIFICANT',
                f'README found in proprietary format ({_names_str}) — convert to README.md',
                f'A README in proprietary format ({_fmts_str}) was detected. '
                'These formats are not machine-readable and cannot be parsed by '
                'automated validators or indexed by repository platforms. '
                'Convert to README.md (plain text or Markdown) so that all '
                'required sections are accessible. '
                'README_DRAFT.md will be generated as a starting template.',
                [f'Found: {f.name}' for f in _prop_readmes] +
                ['Fix: export content to README.md using a suitable conversion '
                 'tool or by copy-pasting the text into a Markdown file']
            ))
        else:
            # Check for same-named .txt files alongside data files
            # (e.g. GDP_FDI_Dataset.txt next to GDP_FDI_Dataset.csv)
            # These may serve as data descriptions rather than a proper README
            data_names = {f.stem.lower() for f in all_files
                          if f.suffix.lower() in DATA_EXTENSIONS}
            companion_txt = [f for f in all_files
                             if f.suffix.lower() == '.txt'
                             and f.stem.lower() in data_names
                             and len(f.relative_to(repo_dir).parts) <= 3]
            if companion_txt:
                findings.append(finding(
                    'A', 'SIGNIFICANT',
                    'No README file found — a text file may serve as documentation',
                    'No file named README.md, README.txt, or README.rst was found. '
                    'A text file with the same name as your data files was detected '
                    '— if this describes your dataset, rename it to README.md and '
                    'ensure it covers all required sections. README_DRAFT.md will '
                    'be generated as a template.',
                    [f'Text file found: {f.name} — verify whether this is your README'
                     for f in companion_txt]
                ))
            else:
                # Check for partial README: any file starting with 'readme' that
                # isn't a standard README name (e.g. README_variables.md).
                _readme_prefix = [
                    f for f in all_files
                    if f.name.lower().startswith('readme')
                    and f.name.lower() not in README_NAMES
                    and f.suffix.lower() not in {'.docx', '.pdf', '.doc', '.pages'}
                    and len(f.relative_to(repo_dir).parts) <= 2
                ]
                if _readme_prefix:
                    _rp_names = ', '.join(
                        f.name for f in sorted(_readme_prefix, key=lambda x: x.name.lower())
                    )
                    findings.append(finding(
                        'A', 'SIGNIFICANT',
                        f'No standard README found — README-like file detected ({_rp_names})',
                        f'No README.md, README.txt, or README.rst was found, but a '
                        f'README-like file ({_rp_names}) is present. This may cover '
                        f'partial documentation but validators typically look for standard '
                        f'README filenames. Create README.md at the repository root '
                        f'covering a study overview and reproduction instructions (it may '
                        f'reference this file). README_DRAFT.md will be generated.',
                        [f'Found: {_rp_names}',
                         'Fix: create README.md at repository root covering study overview, '
                         'requirements, and reproduction steps',
                         'Validators may not recognise non-standard README filenames']
                    ))
                else:
                    # Check for filenames that CONTAIN 'readme' but don't start with it
                    # (e.g. 61622524_README_RTMS.txt, Study_README_v2.txt)
                    _readme_contains = [
                        f for f in all_files
                        if 'readme' in f.name.lower()
                        and not f.name.lower().startswith('readme')
                        and f.name.lower() not in README_NAMES
                        and f.suffix.lower() not in {'.docx', '.pdf', '.doc', '.pages'}
                        and len(f.relative_to(repo_dir).parts) <= 2
                    ]
                    if _readme_contains:
                        _rc_names = ', '.join(
                            f.name for f in sorted(_readme_contains, key=lambda x: x.name.lower())
                        )
                        findings.append(finding(
                            'A', 'SIGNIFICANT',
                            f'No standard README found — README-like file detected ({_rc_names})',
                            f'No README.md, README.txt, or README.rst was found, but a '
                            f'file whose name contains "README" ({_rc_names}) is present. '
                            f'Validators typically look for standard README filenames. '
                            f'Create README.md at the repository root covering a study overview '
                            f'and reproduction instructions (it may reference this file). '
                            f'README_DRAFT.md will be generated.',
                            [f'Found: {_rc_names}',
                             'Fix: create README.md at repository root covering study overview, '
                             'requirements, and reproduction steps',
                             'Validators may not recognise non-standard README filenames']
                        ))
                    else:
                        _sub_readme = next(
                            (f for f in all_files if f.name.lower() in README_NAMES), None
                        )
                        if _sub_readme:
                            _sub_rel = _sub_readme.relative_to(repo_dir)
                            findings.append(finding(
                                'A', 'CRITICAL',
                                'No README at repository root — README found in subdirectory',
                                f'No README at repository root. A README was found at {_sub_rel} '
                                '— move it to the repository root so validators can find it '
                                'immediately on download. Validators typically look only at the '
                                'root level. README_DRAFT.md will be generated.',
                                [f'README found at: {_sub_rel} — move to repository root',
                                 'No README.md, README.txt, or README.rst at root level']
                            ))
                        else:
                            # Before firing CRITICAL, check for readable PDF/docx
                            # with README-like or methods/protocol names that the
                            # filename-only checks above would have missed.
                            _README_LIKE_A = {
                                'readme', 'read_me', 'methods', 'protocol',
                                'documentation', 'instructions',
                            }
                            _readable_doc_a = None
                            for _f in sorted(
                                all_files,
                                key=lambda x: len(x.relative_to(repo_dir).parts)
                            ):
                                if (_f.suffix.lower() in {'.pdf', '.docx'}
                                        and _f.stem.lower() in _README_LIKE_A
                                        and len(_f.relative_to(repo_dir).parts) <= 3):
                                    _res = _inspect_file_content(_f)
                                    if _res.extracted_text and len(_res.extracted_text.strip()) >= 300:
                                        _readable_doc_a = _f
                                        break
                            if _readable_doc_a:
                                findings.append(finding(
                                    'A', 'SIGNIFICANT',
                                    f'README found in proprietary format ({_readable_doc_a.name}) — convert to README.md',
                                    f'A readable {_readable_doc_a.suffix.lower()} document '
                                    f'({_readable_doc_a.name}) was found and may serve as the '
                                    'repository README. This format is not machine-readable by '
                                    'automated validators. Convert to README.md for best '
                                    'accessibility. README_DRAFT.md will be generated.',
                                    [f'Found: {_readable_doc_a.name}',
                                     'Fix: export content to README.md']
                                ))
                            else:
                                findings.append(finding(
                                    'A', 'CRITICAL',
                                    'No README file found',
                                    'Every research repository requires a README. '
                                    'README_DRAFT.md will be generated.',
                                    ['No README.md, README.txt, or README.rst found at repository root level']
                                ))
    else:
        # check if readme is too short to be useful
        for f in all_files:
            if f.name.lower() in README_NAMES and len(f.relative_to(repo_dir).parts) <= 4:
                content = read_file_safe(f)
                if len(content.strip()) < 300:
                    findings.append(finding(
                        'A', 'SIGNIFICANT',
                        'README is present but appears inadequate',
                        f'README is only {len(content.strip())} characters. '
                        'A useful README requires study identification, '
                        'system requirements, installation instructions, '
                        'and execution instructions.',
                        [f'Evidence: {f.name} ({len(content.strip())} chars)']
                    ))
    return findings


def detect_B_no_dependencies(repo_dir, all_files):
    """Failure Mode B: Unpinned or missing dependencies."""
    findings = []
    names_lower = {f.name.lower() for f in all_files}

    # Build frontend-dir exclusion set from all_files — immune to zip wrappers
    # and __MACOSX artefacts because all_files is already filtered.
    # A dir is frontend if it has JS/HTML/CSS marker AND no analysis-code ext.
    # We do NOT require all extensions to be in an allowlist so that dirs
    # containing .sb3, .wav, .mp3, etc. are still correctly identified.
    _dir_exts_b: dict = {}
    for _f in all_files:
        for _anc in _f.parents:
            if _anc == repo_dir:
                break
            _dir_exts_b.setdefault(_anc, set()).add(_f.suffix.lower())
    _frontend_dirs_b = {
        d for d, exts in _dir_exts_b.items()
        if bool(exts & _FRONTEND_MARKERS) and not bool(exts & _ANALYSIS_CODE_EXTS)
    }

    code_files = [f for f in all_files
                  if (f.suffix.lower() in CODE_EXTENSIONS or _is_code_txt(f))
                  and not _is_minified(f)
                  and not any(part.lower() in VENDOR_DIRS
                              for part in f.relative_to(repo_dir).parts)
                  and not any(fd in f.parents for fd in _frontend_dirs_b)]

    has_dep_file = bool(names_lower.intersection(DEPENDENCY_FILES))
    # R install scripts are valid dependency specifications
    if not has_dep_file:
        has_dep_file = any(
            bool(re.match(r'(install|setup).*\.r$', f.name.lower()))
            for f in all_files
        )
    # Also check for install_packages.R style names explicitly
    if not has_dep_file:
        has_dep_file = any(
            'install' in f.name.lower() and f.suffix.lower() == '.r'
            for f in all_files
        )
    # Snakemake workflow — Snakefile is the workflow/dependency spec
    if not has_dep_file:
        has_dep_file = any(
            f.name == 'Snakefile' or f.suffix.lower() == '.smk'
            for f in all_files
        )
    # Modern Pluto notebooks embed deps as PLUTO_PROJECT_TOML_CONTENTS — treat as dep file
    if not has_dep_file:
        for f in all_files:
            if f.suffix.lower() == '.jl':
                try:
                    if 'PLUTO_PROJECT_TOML_CONTENTS' in f.read_text(encoding='utf-8', errors='ignore'):
                        has_dep_file = True
                        break
                except Exception:
                    pass
    has_code = bool(code_files)

    # Stata repos bundle packages in ado/ directory — that IS the dependency spec
    has_stata = any(f.suffix.lower() in {'.do', '.ado'} for f in all_files)
    has_ado_dir = any(
        'ado' in f.parts
        for f in all_files
    )
    if has_stata and has_ado_dir:
        has_dep_file = True  # ado/ directory is the Stata package bundle

    # JS/HTML-only repos (web apps) have no pip/conda dependencies
    _non_web_suffixes = {'.py', '.r', '.rmd', '.jl', '.do', '.ado',
                         '.m', '.scala', '.java', '.cpp', '.c', '.rb'}
    _web_suffixes = {'.js', '.html', '.css', '.ts'}
    _code_suffixes = {f.suffix.lower() for f in all_files
                      if f.suffix.lower() in CODE_EXTENSIONS}
    _is_js_only = bool(_code_suffixes & _web_suffixes) and not bool(_code_suffixes & _non_web_suffixes)
    if _is_js_only:
        has_dep_file = True  # JS web apps have no pip/conda deps

    has_draft_only = "requirements_draft.txt" in names_lower and not has_dep_file

    # Check if README contains inline dependency instructions
    # e.g. install.packages(...) or pip install ... — these are informal
    # but valid dependency specs; downgrade from CRITICAL to SIGNIFICANT
    readme_has_inline_deps = False
    _inline_deps_from_readme = False  # True when the source was a README file
    if not has_dep_file:
        for f in all_files:
            if (f.name.lower() in {'readme.md', 'readme.txt', 'readme.rst', 'readme'}
                    and len(f.relative_to(repo_dir).parts) <= 2):
                readme_content = read_file_safe(f).lower()
                if any(pat in readme_content for pat in [
                    'install.packages(', 'pip install', 'conda install',
                    'install_packages(', 'pkg.add(', 'pkg.instantiate('
                ]):
                    readme_has_inline_deps = True
                    _inline_deps_from_readme = True
                    break
        # Also check R source files for library()/require() calls or
        # vector package assignments (packages <- c(...) / pkgs <- c(...))
        # — these represent inline deps that the generator can extract
        if not readme_has_inline_deps:
            _r_inline_pat = re.compile(
                r'(?:library|require)\s*\(|(?:packages?|pkgs?)\s*<-\s*c\s*\(',
                re.IGNORECASE
            )
            for f in all_files:
                if f.suffix.lower() in {'.r', '.rmd', '.qmd'}:
                    if _r_inline_pat.search(read_file_safe(f)):
                        readme_has_inline_deps = True
                        break  # _inline_deps_from_readme stays False

    # Languages with no standard package-manager equivalent of requirements.txt
    # — downgrade [B] from CRITICAL to SIGNIFICANT and give language-specific advice.
    _scripting_only = not any(
        f.suffix.lower() in {'.py', '.r', '.rmd', '.jl'} for f in all_files
    )
    _is_matlab_only = (
        any(f.suffix.lower() == '.m' for f in all_files) and _scripting_only
    )

    # [B] MATLAB suppression: if README already documents both the MATLAB
    # version AND at least one toolbox, the deposit satisfies [B]'s requirement
    # for environment documentation — no finding needed.
    _MATLAB_VERSION_RE = re.compile(r'R\d{4}[ab]\b|matlab\s+\d+\.\d+', re.IGNORECASE)
    _MATLAB_TOOLBOX_RE = re.compile(r'\btoolbox(?:es)?\b', re.IGNORECASE)
    _matlab_readme_ok = False
    if _is_matlab_only:
        _readme_candidates = [
            f for f in all_files
            if f.name.lower() in {'readme.md', 'readme.txt', 'readme.rst', 'readme'}
        ]
        _readme_candidates.sort(key=lambda x: len(x.relative_to(repo_dir).parts))
        for _rf in _readme_candidates:
            _rc = read_file_safe(_rf)
            if _MATLAB_VERSION_RE.search(_rc) and _MATLAB_TOOLBOX_RE.search(_rc):
                _matlab_readme_ok = True
                break

    _is_stata_only = (
        any(f.suffix.lower() == '.do' for f in all_files)
        and not any(f.suffix.lower() in {'.py', '.r', '.rmd', '.jl', '.m'} for f in all_files)
    )
    _is_sas_only = (
        any(f.suffix.lower() == '.sas' for f in all_files)
        and not any(f.suffix.lower() in {'.py', '.r', '.rmd', '.jl', '.m', '.do'} for f in all_files)
    )

    if has_code and not has_dep_file and not has_draft_only and not readme_has_inline_deps:
        if _is_matlab_only:
            if not _matlab_readme_ok:
                findings.append(finding(
                    'B', 'SIGNIFICANT',
                    'No dependency specification found (MATLAB deposit)',
                    'MATLAB has no standard equivalent of requirements.txt. '
                    'Document the MATLAB version and any required toolboxes '
                    'in your README so validators can configure a matching environment.',
                    [f'Code files found: {len(code_files)}',
                     'Recommendation: list MATLAB version and toolboxes in README']
                ))
        elif _is_stata_only:
            findings.append(finding(
                'B', 'SIGNIFICANT',
                'No dependency specification found (Stata deposit)',
                'Stata has no standard equivalent of requirements.txt. '
                'Document the Stata version and any required packages '
                '(ssc install) in your README.',
                [f'Code files found: {len(code_files)}',
                 'Recommendation: list Stata version and ssc-installed packages in README']
            ))
        elif _is_sas_only:
            findings.append(finding(
                'B', 'SIGNIFICANT',
                'No dependency specification found (SAS deposit)',
                'SAS has no pip-style package manager. Document the SAS version, '
                'required SAS products/modules, and any SASLIB paths in your README.',
                [f'Code files found: {len(code_files)}',
                 'Recommendation: list SAS version and required modules in README']
            ))
        else:
            # Suppress entirely for zero-dep readers — nothing to pin.
            if not _is_zero_dep_reader(all_files):
                _is_trivial_helper = (
                    len(code_files) == 1
                    and code_files[0].stat().st_size < 1024
                    and any(kw in code_files[0].stem.lower()
                            for kw in _TRIVIAL_NAMES)
                )
                # Check for inline version comments (e.g. # Package version: 3.5.3)
                # If found, downgrade: the author has documented versions, just not
                # in a standard file. LOW CONFIDENCE with a targeted recommendation.
                _INLINE_VERSION_RE = re.compile(
                    r'(?:import|from)\s+[\w.]+.*#.*(?:version|v)[:\s]+(\d[\d.]+)',
                    re.IGNORECASE
                )
                _has_inline_versions = any(
                    _INLINE_VERSION_RE.search(read_file_safe(f))
                    for f in code_files
                    if f.suffix.lower() == '.py'
                )
                if _has_inline_versions:
                    findings.append(finding(
                        'B', 'LOW CONFIDENCE',
                        'Version numbers found as inline comments — no requirements.txt',
                        'Package versions are documented as inline comments on import '
                        'statements rather than in a requirements.txt or environment.yml. '
                        'This is better than nothing but a dedicated dependency file '
                        'ensures reproducibility.',
                        [f'Code files found: {len(code_files)}',
                         'Recommendation: generate requirements.txt from inline version comments']
                    ))
                else:
                    _b_severity = 'SIGNIFICANT' if _is_trivial_helper else 'CRITICAL'
                    findings.append(finding(
                        'B', _b_severity,
                        'No dependency specification found',
                        'Code files are present but no dependency file was found. '
                        'A requirements_DRAFT.txt will be generated from import '
                        'statements with all versions marked UNKNOWN.',
                        [f'Code files found: {len(code_files)}',
                         'No requirements.txt, environment.yml, renv.lock, '
                         'or equivalent found']
                    ))
    elif has_code and not has_dep_file and not has_draft_only and readme_has_inline_deps:
        if _inline_deps_from_readme:
            _b_inline_title = 'Dependencies documented inline in README but no dependency file found'
            _b_inline_detail = (
                'Install instructions were found in the README but no requirements.txt, '
                'renv.lock, or equivalent file exists. Inline instructions are better '
                'than nothing but a dedicated dependency file ensures reproducibility. '
                'A requirements_DRAFT.txt will be generated.'
            )
        else:
            _b_inline_title = 'No dependency file found — packages detected in code files'
            _b_inline_detail = (
                'No requirements.txt, renv.lock, or equivalent dependency file was found. '
                'Packages were detected from library()/require() calls in the code. '
                'A requirements_DRAFT.txt will be generated from these imports — '
                'add version numbers and rename to requirements.txt before deposit.'
            )
        findings.append(finding(
            'B', 'SIGNIFICANT',
            _b_inline_title,
            _b_inline_detail,
            [f'Code files found: {len(code_files)}',
             'Recommendation: extract install instructions to requirements.txt or renv.lock']
        ))
    elif has_code and has_draft_only:
        # prior run left a requirements_DRAFT.txt — check if versions are pinned
        draft_file = next(f for f in all_files if f.name.lower() == "requirements_draft.txt")
        draft_content = read_file_safe(draft_file)
        has_pinned = any("==" in l for l in draft_content.splitlines() if l.strip() and not l.strip().startswith("#"))
        # Suppress entirely for zero-dep readers — draft is empty and nothing to finalise.
        if not _is_zero_dep_reader(all_files):
            findings.append(finding(
                'B', 'SIGNIFICANT',
                'requirements_DRAFT.txt found from prior run but not yet finalised',
                'A requirements_DRAFT.txt exists but has not been completed and renamed '
                'to requirements.txt. Pin all version numbers and rename before deposit.',
                ['Action: complete version numbers in requirements_DRAFT.txt and rename to requirements.txt']
            ))
    elif has_dep_file:
        # check for unpinned dependencies in requirements.txt
        for f in all_files:
            if f.name.lower() == 'requirements.txt':
                content = read_file_safe(f)
                unpinned = []
                for line in content.splitlines():
                    line = line.strip()
                    if not line or line.startswith('#') or line.startswith('-'):
                        continue
                    # PEP 508 direct URL reference: pkg @ https://...
                    # Treat as pinned if the URL contains a version string
                    # (e.g. en_core_web_sm-3.0.0 or ==3.0.0 in the URL)
                    if ' @ http' in line or '\t@ http' in line:
                        url_part = line.split('@', 1)[-1]
                        if re.search(r'[-/=]\d+\.\d+', url_part):
                            continue  # versioned URL — pinned
                        else:
                            unpinned.append(line)
                            continue
                    if '==' not in line and '>=' not in line \
                            and '<=' not in line and '~=' not in line:
                        unpinned.append(line)
                if unpinned:
                    findings.append(finding(
                        'B', 'SIGNIFICANT',
                        'requirements.txt contains unpinned dependencies',
                        'Package names without exact version numbers '
                        'will install the latest version, which may '
                        'differ from what was used at publication.',
                        [f'Unpinned: {", ".join(unpinned[:10])}']
                        + (['...and more' ]if len(unpinned) > 10 else [])
                    ))
    return findings


def detect_C_absolute_paths(repo_dir, all_files):
    """Failure Mode C: Absolute paths that only work on researcher machine."""
    findings = []
    # Stata installed-package directories (equivalent to site-packages) — skip entirely
    _stata_lib_dirs = {'plus', 'personal', 'stbplus'}
    code_files = [f for f in all_files
                  if f.suffix.lower() in CODE_EXTENSIONS
                  and not _is_minified(f)
                  and not ('ado' in f.parts
                           and any(p in _stata_lib_dirs for p in f.parts))]

    # Also scan Jupyter notebook cell sources
    notebook_sources = []
    import json as _json
    for nb in all_files:
        if nb.suffix.lower() == '.ipynb':
            try:
                data = _json.loads(nb.read_text(encoding='utf-8', errors='ignore'))
                for cell in data.get('cells', []):
                    src = ''.join(cell.get('source', []))
                    if src:
                        notebook_sources.append((nb, src))
            except Exception:
                pass

    abs_pattern = re.compile(
        r'(/Users/[a-zA-Z][a-zA-Z0-9_\-]{1,}/)'
        r'|(/home/[a-zA-Z][a-zA-Z0-9_\-]{1,}/)'
        r'|(/root/[a-zA-Z])'
        r'|(/mnt/[a-zA-Z0-9_\-]{1,}/)'
        r'|([A-Z]:\\[A-Za-z][A-Za-z0-9_. -]{1,}\\)'
        r'|([A-Z]:/[A-Za-z][A-Za-z0-9_. -]{1,}/)'
    )

    # Collect hits grouped by file: {Path: [(line_no, snippet), ...]}
    _hits: dict = {}
    for f in code_files:
        ext = f.suffix.lower()
        content_f = read_file_safe(f)
        for i, line in enumerate(content_f.splitlines(), 1):
            stripped = line.strip()
            # Skip comment lines — per language
            if ext == '.m' and stripped.startswith('%'):        # MATLAB
                continue
            if ext in {'.py', '.r', '.rmd', '.qmd', '.sh', '.bash'} \
                    and stripped.startswith('#'):               # Python / R / shell
                continue
            if ext == '.java' and (stripped.startswith('//') or stripped.startswith('*')):
                continue                                        # Java line/block comment
            if stripped.startswith('#'):                        # catch-all for other # languages
                continue
            if stripped.startswith('"""') or stripped.startswith("'''"):  # Python docstrings
                continue
            # home-relative and env-var paths are not machine-specific hardcodes
            if '~/' in stripped:
                continue  # ~/path expands relative to $HOME, not a fixed machine path
            if re.search(r'\$\{?(?:HOME|USERPROFILE|HOMEPATH)\}?/', stripped):
                continue  # $HOME/path, ${HOME}/path — same issue
            # PATH/library exports and shell sourcing are environment setup, not data paths
            if re.match(r'export\s+\w*(?:PATH|LIB)\s*=', stripped):
                continue
            if re.match(r'(?:source|\.)\s+', stripped):
                continue  # shell source commands expand env vars at runtime
            if abs_pattern.search(line):
                _hits.setdefault(f, []).append((i, stripped[:80]))

    # Collect notebook hits grouped by notebook file
    _nb_hits: dict = {}
    for nb, src in notebook_sources:
        for i, line in enumerate(src.splitlines(), 1):
            stripped = line.strip()
            if stripped.startswith('#'):
                continue
            if abs_pattern.search(line):
                _nb_hits.setdefault(nb, []).append((i, stripped[:80]))

    # Emit a single grouped [C] finding across all affected files and notebooks.
    # proposed_corrections/ still has one file per affected script — the grouped
    # finding just prevents 14 individual [C] lines from dominating the report.
    _path_hits = []
    for f, hits in _hits.items():
        first_line, first_snippet = hits[0]
        _path_hits.append((f.name, first_line, first_snippet))
    for nb, hits in _nb_hits.items():
        first_line, first_snippet = hits[0]
        _path_hits.append((nb.name, first_line, first_snippet))

    if _path_hits:
        total = len(_path_hits)
        files_str = ', '.join(name for name, _, _ in _path_hits[:5])
        if total > 5:
            files_str += f' (and {total - 5} more files)'
        findings.append(finding(
            'C', 'SIGNIFICANT',
            f'Absolute paths detected in {total} file{"s" if total != 1 else ""}',
            'Absolute paths break reproducibility — they only '
            "work on the researcher's machine. Corrected copies "
            "with relative paths will be generated in "
            '/proposed_corrections/.',
            [f'Files: {files_str}',
             f'First occurrence: {_path_hits[0][0]} line {_path_hits[0][1]}: '
             f'{_path_hits[0][2]}']
        ))

    return findings


def detect_D_no_entry_point(repo_dir, all_files):
    """Failure Mode D: No execution order or entry point."""
    findings = []
    names_lower = {f.name.lower() for f in all_files}

    _MAIN_ENTRY_NAMES = {'main.m', 'main.py', 'main.ipynb', 'index.ipynb'}
    has_run_all = (
        any('run_all' in n
            for n in names_lower
            if n.endswith('.sh') or n.endswith('.py') or n.endswith('.do'))
        or bool(names_lower & _MAIN_ENTRY_NAMES)
    )

    has_makefile = 'makefile' in names_lower

    _stata_lib_dirs = {'plus', 'personal', 'stbplus'}
    # Build frontend-dir exclusion set from all_files — immune to zip wrappers
    # and __MACOSX artefacts because all_files is already filtered.
    # A dir is frontend if it has JS/HTML/CSS marker AND no analysis-code ext.
    _dir_exts_d: dict = {}
    for _f in all_files:
        for _anc in _f.parents:
            if _anc == repo_dir:
                break
            _dir_exts_d.setdefault(_anc, set()).add(_f.suffix.lower())
    _frontend_dirs_d = {
        d for d, exts in _dir_exts_d.items()
        if bool(exts & _FRONTEND_MARKERS) and not bool(exts & _ANALYSIS_CODE_EXTS)
    }
    import sys as _sys
    for _fd in sorted(_frontend_dirs_d, key=str):
        print(f'  [FE-DIR] {_fd.relative_to(repo_dir)}: '
              f'{sorted(_dir_exts_d[_fd])}', file=_sys.stderr)
    _researcher_code = [
        f for f in all_files
        if (f.suffix.lower() in CODE_EXTENSIONS or _is_code_txt(f))
        and not _is_minified(f)
        and not ('ado' in f.parts and any(p in _stata_lib_dirs for p in f.parts))
        and not any(fd in f.parents for fd in _frontend_dirs_d)
    ]

    _part_counts = {}
    for _f in _researcher_code:
        _mp = re.search(r'[_\-][Pp]art(\d+)', _f.stem)
        if _mp:
            _part_counts[int(_mp.group(1))] = _part_counts.get(int(_mp.group(1)), 0) + 1
    has_numbered = (
        any(re.match(r'^(\d+)(?:[.\-](\d+))?(?:[_\-\s]|\.(?!\d))', f.name)
            for f in _researcher_code)
        or len(_part_counts) >= 2  # ≥2 distinct _PartN numbers = ordered sequence
    )

    # README-based execution order: if the README references ≥2 actual code
    # files by name, the author has documented the order and [D] should not fire.
    # (QUICKSTART falls back to README order; [D] must honour the same signal.)
    _code_names = {f.name for f in _researcher_code}
    _readme_mentions = 0
    for _rf in all_files:
        if _rf.name.lower() in README_NAMES and len(_rf.relative_to(repo_dir).parts) <= 2:
            _rsrc = read_file_safe(_rf)
            _readme_mentions = sum(1 for n in _code_names if n in _rsrc)
            break
    has_readme_order = _readme_mentions >= 2

    # Numbered directory sequences (01_foo/, 02_bar/) imply execution order
    _dir_numbers = set()
    for _f in all_files:
        for _part in _f.relative_to(repo_dir).parts[:-1]:
            _dm = re.match(r'^(\d{1,3})[_\-\.]', _part)
            if _dm:
                _dir_numbers.add(int(_dm.group(1)))
    has_numbered_dirs = len(_dir_numbers) >= 2

    # README explicitly names a master/entry-point file
    _MASTER_RE = re.compile(
        r'master\s+file|master\s+do.file|execute\s+first|run\s+first',
        re.IGNORECASE
    )
    has_master_phrase = False
    for _f in all_files:
        if _f.name.lower() in README_NAMES and len(_f.relative_to(repo_dir).parts) <= 2:
            if _MASTER_RE.search(read_file_safe(_f)):
                has_master_phrase = True
                break

    code_count = len(_researcher_code)

    if code_count > 1 and not has_run_all \
            and not has_makefile and not has_numbered \
            and not has_numbered_dirs and not has_master_phrase \
            and not has_readme_order:
        findings.append(finding(
            'D', 'SIGNIFICANT',
            'No clear execution entry point or order',
            f'{code_count} code files found but no run_all script, '
            'Makefile, or numbered script sequence detected. '
            'A QUICKSTART_DRAFT.md will be generated.',
            [f'Code files: {code_count}',
             'No run_all.sh, Makefile, or 01_/02_/03_ numbering found']
        ))

    return findings


def detect_N_no_licence(repo_dir, all_files):
    """Failure Mode N: No licence file."""
    findings = []
    names_lower = {f.name.lower() for f in all_files}

    if not names_lower.intersection(LICENCE_NAMES):
        findings.append(finding(
            'N', 'LOW CONFIDENCE',
            'No licence file found',
            'Without a licence, validators have no legal clarity '
            'on whether they can use, reproduce, or share this work. '
            'A LICENCE_DRAFT.txt will be generated. '
            'Note: this is a legal concern, not a reproducibility blocker — '
            'a validator can still attempt reproduction without a licence file.',
            ['No LICENCE, LICENSE, licence.md, or license.txt found']
        ))

    return findings


def detect_Z_no_commit_hash(repo_dir, all_files):
    """Failure Mode Z: No commit hash or version tag in README."""
    findings = []

    # [Z] only makes sense for git-managed deposits. A static journal archive
    # has no git history, so asking for a commit hash is inappropriate.
    if not (repo_dir / '.git').is_dir():
        return findings

    for f in all_files:
        if f.name.lower() in README_NAMES and len(f.relative_to(repo_dir).parts) <= 2:
            content = read_file_safe(f)
            # look for commit hash (40 hex chars) or version tag
            has_hash = bool(re.search(r'\b[0-9a-f]{40}\b', content))
            # Only match version tags that refer to the CODE deposit,
            # not dataset version strings inside parentheses or prose
            has_tag = bool(re.search(
                r'(?:^|\s|commit[:\s]+|tag[:\s]+|release[:\s]+|version[:\s]+)v\d+\.\d+[\d.]*(?:\s|$)',
                content, re.IGNORECASE | re.MULTILINE
            ))
            has_doi = bool(re.search(r'\b10\.\d{4}/', content))
            if not has_hash and not has_tag and not has_doi:
                findings.append(finding(
                    'Z', 'SIGNIFICANT',
                    'No commit hash or version tag in README',
                    'Without a commit hash or version tag, validators '
                    'cannot confirm they have the exact version of code '
                    'used to produce the published results.',
                    [f'Evidence: {f.name} — no 40-char hex hash or '
                     'version tag found']
                ))

    return findings


def detect_BJ_encrypted_files(repo_dir, all_files):
    """Failure Mode BJ: Encrypted or high-entropy data files."""
    findings = []

    for f in all_files:
        if f.suffix.lower() in ENCRYPTED_EXTENSIONS:
            findings.append(finding(
                'BJ', 'CRITICAL',
                f'Encrypted file detected: {f.name}',
                'This file appears to be encrypted and cannot be used '
                'by validators without a decryption key that is not '
                'present in this repository.',
                [f'Evidence: {f.name} has encryption extension '
                 f'{f.suffix}']
            ))

        # check for git-crypt magic bytes in data-like files
        elif f.suffix.lower() in DATA_EXTENSIONS:
            try:
                header = f.read_bytes()[:16]
                if header[:10] == b'\x00GITCRYPT':
                    findings.append(finding(
                        'BJ', 'CRITICAL',
                        f'Git-crypt encrypted file: {f.name}',
                        'This file is encrypted with git-crypt. '
                        'Validators cannot read it without the '
                        'symmetric key.',
                        [f'Evidence: {f.name} contains git-crypt '
                         f'magic bytes']
                    ))
            except Exception:
                pass

    return findings


def detect_BL_git_history_dependency(repo_dir, all_files):
    """Failure Mode BL: Shallow clone / missing git history dependency."""
    findings = []

    for f in all_files:
        if f.name.lower() in {'setup.py', 'setup.cfg', 'pyproject.toml'}:
            content = read_file_safe(f)
            if 'setuptools_scm' in content or 'setuptools-scm' in content:
                findings.append(finding(
                    'BL', 'CRITICAL',
                    f'setuptools_scm detected in {f.name}',
                    'This package uses git history to determine its '
                    'version number. When downloaded as a ZIP from '
                    'Zenodo, Figshare, or GitHub, the .git directory '
                    'is absent and the package will fail to import. '
                    'Pin the version explicitly: __version__ = "1.0.0"',
                    [f'Evidence: {f.name} — setuptools_scm reference']
                ))
            if 'versioneer' in content:
                findings.append(finding(
                    'BL', 'CRITICAL',
                    f'versioneer detected in {f.name}',
                    'versioneer uses git history to determine version '
                    'numbers. ZIP downloads strip the .git directory '
                    'and this will fail immediately.',
                    [f'Evidence: {f.name} — versioneer reference']
                ))

        # check shell scripts and Makefiles for git describe
        if f.suffix.lower() in {'.sh', '.bash', ''} \
                or f.name.lower() == 'makefile':
            content = read_file_safe(f)
            if 'git describe' in content or 'git log' in content:
                findings.append(finding(
                    'BL', 'SIGNIFICANT',
                    f'git describe/log call in {f.name}',
                    'This script calls git commands that require '
                    '.git history. ZIP downloads will not have this '
                    'and the script will fail.',
                    [f'Evidence: {f.name} — git describe or git log']
                ))

    return findings


def detect_BK_system_clock(repo_dir, all_files):
    """Failure Mode BK: System clock dependency."""
    findings = []
    code_files = [f for f in all_files
                  if f.suffix.lower() in CODE_EXTENSIONS]

    clock_in_filename = re.compile(
        r'(datetime\.now|datetime\.today|time\.time)\s*\(\s*\)'
        r'[^\n]{0,50}(f["\']|format|str|%)'
    )
    clock_as_seed = re.compile(
        r'(seed|random)\s*\([^\n]{0,80}(datetime\.now|time\.time)'
    )

    clock_in_logic = re.compile(
        r'(datetime\.now|datetime\.today|time\.time)\s*\(\s*\)'
    )
    clock_in_logic = re.compile(
        r'(datetime\.now|datetime\.today|time\.time)\s*\(\s*\)'
    )
    clock_in_logic = re.compile(
        r'(datetime\.now|datetime\.today|time\.time)\s*\(\s*\)'
    )
    for f in code_files:
        src = read_file_safe(f)
        if clock_in_filename.search(src):
            findings.append(finding(
                'BK', 'SIGNIFICANT',
                f'System clock used in filename generation: {f.name}',
                'Output filenames derived from datetime.now() or '
                'time.time() will differ between runs.',
                [f'Evidence: {f.name} - clock-based filename pattern']
            ))
        if clock_as_seed.search(src):
            findings.append(finding(
                'BK', 'SIGNIFICANT',
                f'System clock used as random seed: {f.name}',
                'Seeds derived from the system clock change every run.',
                [f'Evidence: {f.name} - clock-based seed pattern']
            ))
        elif clock_in_logic.search(src) and not clock_in_filename.search(src):
            findings.append(finding(
                'BK', 'SIGNIFICANT',
                f'System clock used in conditional logic: {f.name}',
                'Code behaviour depends on current date or time. '
                'Results may differ if run in a different month or year.',
                [f'Evidence: {f.name} - clock-based logic pattern']
            ))
    return findings


def detect_W_git_lfs(repo_dir, all_files):
    """Failure Mode W: Git LFS pointer files."""
    findings = []

    for f in all_files:
        if f.suffix.lower() in DATA_EXTENSIONS \
                or f.suffix.lower() in {'.png', '.jpg', '.pdf'}:
            try:
                header = f.read_bytes()[:128].decode('utf-8', errors='ignore')
                if 'version https://git-lfs.github.com' in header:
                    findings.append(finding(
                        'W', 'CRITICAL',
                        f'Git LFS pointer file: {f.name}',
                        'This file is a Git LFS pointer, not the '
                        'actual data. The real file must be retrieved '
                        'using git lfs pull. Validators downloading '
                        'this repository as a ZIP will get the pointer '
                        'file only.',
                        [f'Evidence: {f.name} — Git LFS pointer header']
                    ))
            except Exception:
                pass

    return findings


# ── DETECTOR ADMISSION POLICY ────────────────────────────────────────────────
#
# A detector earns a place in the specification if it predicts validation
# friction — i.e. it identifies a condition that would cause an independent
# validator to spend extra time, fail silently, or be unable to reproduce
# results.
#
# Severity assignment rules:
#   CRITICAL       — validator cannot begin reproduction without resolving this
#   SIGNIFICANT    — validator will likely fail or produce wrong results
#   LOW CONFIDENCE — best practice; may predict friction but evidence is weak
#
# Detectors that are "good housekeeping" but do not predict friction
# (e.g. [BM] CITATION.cff, [BD] contact info) must remain LOW CONFIDENCE
# permanently, regardless of how often they fire.
#
# Before adding a new detector, answer: "Would a validator spend measurably
# more time on a repository where this condition is present?" If no, do not add.
#
# ── main entry point ─────────────────────────────────────────────────────────

def run_simple_detectors(repo_dir, all_files, zip_name=None):
    """Run all simple pattern-matching detectors. Return list of findings."""

    # Clear the per-run content inspection cache so each deposit starts fresh.
    # (The module-level dict persists across calls in long-running processes.)
    _file_inspection_cache.clear()

    print("  [A]  README check...")
    print("  [B]  Dependency check...")
    print("  [C]  Absolute path check...")
    print("  [D]  Entry point check...")
    print("  [N]  Licence check...")
    print("  [Z]  Commit hash check...")
    print("  [W]  Git LFS check...")
    print("  [BJ] Encrypted file check...")
    print("  [BK] System clock check...")
    print("  [BL] Git history dependency check...")
    print("  [F]  Random seed check...")
    print("  [U]  Environment variable check...")
    all_findings = []
    # prior run detection
    prior_report = next((f for f in all_files if f.name.lower() == "cleaning_report.md"), None)
    if prior_report:
        try:
            prior_content = prior_report.read_text(encoding="utf-8", errors="ignore")
            import re as _re
            version_match = _re.search(r"v(\d+\.\d+\.\d+)", prior_content)
            date_match = _re.search(r"(\d{4}-\d{2}-\d{2})", prior_content)
            version_str = version_match.group(0) if version_match else "unknown version"
            date_str = date_match.group(0) if date_match else "unknown date"
            all_findings.append(finding(
                "BQ", "SIGNIFICANT",
                f"Prior ValiChord report detected ({version_str}, {date_str})",
                "A previous ValiChord cleaning report was found in this repository. "
                "This appears to be a re-run. Review prior findings before actioning new ones.",
                [f"Prior report: {prior_report.name} ({version_str}, {date_str})"]
            ))
        except Exception:
            pass
    all_findings += detect_A_no_readme(repo_dir, all_files)
    all_findings += detect_B_no_dependencies(repo_dir, all_files)
    all_findings += detect_C_absolute_paths(repo_dir, all_files)
    all_findings += detect_D_no_entry_point(repo_dir, all_files)
    all_findings += detect_N_no_licence(repo_dir, all_files)
    all_findings += detect_Z_no_commit_hash(repo_dir, all_files)
    all_findings += detect_W_git_lfs(repo_dir, all_files)
    all_findings += detect_BJ_encrypted_files(repo_dir, all_files)
    all_findings += detect_BK_system_clock(repo_dir, all_files)
    all_findings += detect_BL_git_history_dependency(repo_dir, all_files)
    print("  [G]  README adequacy check...")
    all_findings += detect_G_inadequate_readme(repo_dir, all_files)
    print("  [H]  Hardcoded versions check...")
    all_findings += detect_H_hardcoded_versions(repo_dir, all_files)
    print("  [K]  Compute environment check...")
    all_findings += detect_K_compute_environment(repo_dir, all_files)
    print("  [P]  Pre-registration check...")
    all_findings += detect_P_preregistration(repo_dir, all_files)
    print("  [PAP] Pre-analysis plan check...")
    all_findings += detect_PAP_preregistration_document(repo_dir, all_files)
    print("  [V]  Virtual environment check...")
    all_findings += detect_V_virtual_environment(repo_dir, all_files, all_findings)
    print("  [I]  Intermediate files check...")
    all_findings += detect_I_intermediate_files(repo_dir, all_files)
    print("  [J]  Notebook execution order check...")
    all_findings += detect_J_notebook_order(repo_dir, all_files)
    print("  [M]  Python version check...")
    all_findings += detect_M_python_version_conflict(repo_dir, all_files)
    print("  [L]  Missing file references check...")
    all_findings += detect_L_large_files_missing(repo_dir, all_files)
    print("  [O]  Committed outputs check...")
    all_findings += detect_O_output_not_committed(repo_dir, all_files)
    print("  [Q]  Configuration files check...")
    all_findings += detect_Q_config_files(repo_dir, all_files)
    print("  [R]  Statistical assumptions check...")
    all_findings += detect_R_statistical_tests_undocumented(repo_dir, all_files)
    print("  [S]  Software citations check...")
    all_findings += detect_S_software_citations_missing(repo_dir, all_files)
    print("  [T]  Test coverage check...")
    all_findings += detect_T_test_coverage(repo_dir, all_files)
    print("  [X]  Containerisation check...")
    all_findings += detect_X_no_container(repo_dir, all_files)
    print("  [Y]  Data source check...")
    all_findings += detect_Y_data_source_missing(repo_dir, all_files)
    print("  [AA] Figure reproducibility check...")
    all_findings += detect_AA_figure_reproducibility(repo_dir, all_files)
    print("  [AB] Parallel determinism check...")
    all_findings += detect_AB_parallel_no_seed(repo_dir, all_files)
    print("  [AC] Deprecated functions check...")
    all_findings += detect_AC_deprecated_functions(repo_dir, all_files)
    print("  [AD] Gitignore check...")
    all_findings += detect_AD_missing_gitignore(repo_dir, all_files)
    print("  [AE] Mixed languages check...")
    all_findings += detect_AE_mixed_languages(repo_dir, all_files)
    print("  [AF] Output format check...")
    all_findings += detect_AF_output_format_undocumented(repo_dir, all_files)
    print("  [AG] Hardcoded credentials check...")
    all_findings += detect_AG_api_keys_in_code(repo_dir, all_files)
    print("  [AH] Changelog check...")
    all_findings += detect_AH_no_changelog(repo_dir, all_files)
    print("  [AI] Print debugging check...")
    all_findings += detect_AI_print_debugging(repo_dir, all_files)
    print("  [AJ] Magic numbers check...")
    all_findings += detect_AJ_hardcoded_sample_size(repo_dir, all_files)
    print("  [AK] External URLs check...")
    all_findings += detect_AK_external_urls(repo_dir, all_files)
    print("  [AL] Data privacy check...")
    all_findings += detect_AL_data_privacy(repo_dir, all_files)
    print("  [AM] Pipeline automation check...")
    all_findings += detect_AM_makefile_missing(repo_dir, all_files)
    print("  [AN] Commented code check...")
    all_findings += detect_AN_commented_code(repo_dir, all_files)
    print("  [AO] R-specific check...")
    all_findings += detect_AO_r_specific_issues(repo_dir, all_files)
    print("  [AP] Stata-specific check...")
    all_findings += detect_AP_stata_specific(repo_dir, all_files)
    print("  [AQ] Model files check...")
    all_findings += detect_AQ_large_model_files(repo_dir, all_files)
    print("  [AR] Encoding check...")
    all_findings += detect_AR_encoding_issues(repo_dir, all_files)
    print("  [AS] Network calls check...")
    all_findings += detect_AS_network_calls(repo_dir, all_files)
    print("  [AT] Database dependency check...")
    all_findings += detect_AT_database_dependency(repo_dir, all_files)
    print("  [AU] Cloud storage check...")
    all_findings += detect_AU_cloud_storage(repo_dir, all_files)
    print("  [AV] Hardcoded dates check...")
    all_findings += detect_AV_hardcoded_dates(repo_dir, all_files)
    print("  [AW] DOI check...")
    all_findings += detect_AW_missing_doi(repo_dir, all_files, zip_name=zip_name)
    print("  [AX] Container check...")
    all_findings += detect_AX_container_not_tested(repo_dir, all_files)
    print("  [AY] Workflow file check...")
    all_findings += detect_AY_workflow_file(repo_dir, all_files)
    print("  [AZ] Figure format check...")
    all_findings += detect_AZ_figure_format(repo_dir, all_files)
    print("  [BA] Checksums check...")
    all_findings += detect_BA_missing_checksums(repo_dir, all_files)
    print("  [BB] Script permissions check...")
    all_findings += detect_BB_script_permissions(repo_dir, all_files)
    print("  [BC] Line endings check...")
    all_findings += detect_BC_mixed_line_endings(repo_dir, all_files)
    print("  [BD] Contact info check...")
    all_findings += detect_BD_missing_contact(repo_dir, all_files)
    print("  [BE] Compiled files check...")
    all_findings += detect_BE_pyc_files(repo_dir, all_files)
    print("  [BF] Notebook outputs check...")
    all_findings += detect_BF_notebook_outputs_missing(repo_dir, all_files)
    print("  [BG] Funding acknowledgement check...")
    all_findings += detect_BG_missing_acknowledgements(repo_dir, all_files)
    print("  [BH] Archive files check...")
    all_findings += detect_BH_zip_bomb_risk(repo_dir, all_files)
    print("  [BI] Unicode paths check...")
    all_findings += detect_BI_unicode_in_paths(repo_dir, all_files)
    all_findings += detect_BM_citation_cff(repo_dir, all_files)
    all_findings += detect_BP_licence_in_readme_only(repo_dir, all_files)
    all_findings += detect_BR_credentials_exposed(repo_dir, all_files)
    all_findings += detect_BS_archive_code_present(repo_dir, all_files)
    all_findings += detect_BT_spaces_in_filenames(repo_dir, all_files)
    print("  [FD] Duplicate format pairs check...")
    all_findings += detect_FD_duplicate_format_pairs(repo_dir, all_files)
    print("  [BU] Conda channel priority check...")
    all_findings += detect_BU_conda_channel_priority(repo_dir, all_files)
    print("  [BV] Shell error handling check...")
    all_findings += detect_BV_shell_no_set_e(repo_dir, all_files)
    print("  [BW] Empty code files check...")
    all_findings += detect_BW_empty_code_files(repo_dir, all_files)
    print("  [BX] Pluto manifest check...")
    all_findings += detect_BX_pluto_empty_manifest(repo_dir, all_files)
    print("  [BY] Julia manifest check...")
    all_findings += detect_BY_julia_missing_manifest(repo_dir, all_files)
    print("  [BZ] MATLAB v7.3 format check...")
    all_findings += detect_BZ_matlab_v73_format(repo_dir, all_files)
    print("  [CA] README script reference check...")
    all_findings += detect_CA_readme_script_missing(repo_dir, all_files)
    print("  [CB] Snakemake environment isolation check...")
    all_findings += detect_CB_snakemake_no_env_isolation(repo_dir, all_files)
    print("  [CD] Dockerfile build order check...")
    all_findings += detect_CD_dockerfile_run_before_copy(repo_dir, all_files)
    print("  [CC] External tool versions check...")
    all_findings += detect_CC_undocumented_external_tools(repo_dir, all_files)
    print("  [CE] Unpinned GitHub R packages check...")
    all_findings += detect_CE_unpinned_github_packages(repo_dir, all_files)
    print("  [CR] CRLF line endings check...")
    all_findings += detect_CR_crlf_line_endings(repo_dir, all_files)
    print("  [CF] Notebook committed outputs check...")
    all_findings += detect_CF_notebook_outputs_committed(repo_dir, all_files)
    print("  [CG] Unpinned git+ requirements check...")
    all_findings += detect_CG_unpinned_git_requirements(repo_dir, all_files)
    print("  [BN] Codebook reference check...")
    all_findings += detect_BN_codebook_reference_mismatch(repo_dir, all_files)

    print("  [E]  Data documentation check...")
    all_findings += detect_E_missing_data_documentation(repo_dir, all_files)
    all_findings += detect_F_missing_seeds(repo_dir, all_files)
    all_findings += detect_U_environment_variables(repo_dir, all_files)
    print("  [CH] Broken R source() chain check...")
    all_findings += detect_CH_broken_source_chain(repo_dir, all_files)
    print("  [CI] Live data no archive check...")
    all_findings += detect_CI_live_data_no_archive(repo_dir, all_files)
    print("  [CJ] README missing file references check...")
    all_findings += detect_CJ_readme_references_missing_files(repo_dir, all_files)
    print("  [CK] Conflicting READMEs check...")
    all_findings += detect_CK_conflicting_readmes(repo_dir, all_files)
    print("  [CL] Bioconductor version pin check...")
    all_findings += detect_CL_bioconductor_unpinned(repo_dir, all_files)
    print("  [CM] Nextflow container/conda check...")
    all_findings += detect_CM_nextflow_no_container(repo_dir, all_files)
    print("  [CN] Known version conflicts check...")
    all_findings += detect_CN_known_version_conflicts(repo_dir, all_files)
    print("  [CO] MATLAB undocumented functions check...")
    all_findings += detect_CO_matlab_undocumented_functions(repo_dir, all_files)
    print("  [CP] Python 2 syntax check...")
    all_findings += detect_CP_python2_syntax(repo_dir, all_files)
    print("  [CQ] Julia Pkg.add() at runtime check...")
    all_findings += detect_CQ_julia_pkg_add_at_runtime(repo_dir, all_files)
    print("  [CS] Committed model binary (pickle) check...")
    all_findings += detect_CS_committed_model_binary(repo_dir, all_files)
    print("  [CU] Conda unpinned packages check...")
    all_findings += detect_CU_conda_unpinned_packages(repo_dir, all_files)
    print("  [CV] Hardcoded CUDA no fallback check...")
    all_findings += detect_CV_hardcoded_cuda_no_fallback(repo_dir, all_files)
    print("  [CW] Reticulate R/Python coupling check...")
    all_findings += detect_CW_reticulate_coupling(repo_dir, all_files)
    print("  [CX] System-level dependencies check...")
    all_findings += detect_CX_system_dependencies(repo_dir, all_files)
    print("  [CY] Checksum not verified check...")
    all_findings += detect_CY_checksum_not_verified(repo_dir, all_files)
    print("  [CZ] EOL Docker base image check...")
    all_findings += detect_CZ_eol_docker_base_image(repo_dir, all_files)
    print("  [DA] NLP model not in Dockerfile check...")
    all_findings += detect_DA_nlp_model_not_in_dockerfile(repo_dir, all_files)
    print("  [DB] Shiny app interactive verification check...")
    all_findings += detect_DB_shiny_app(repo_dir, all_files)
    print("  [DC] Monorepo independent sub-projects check...")
    all_findings += detect_DC_monorepo_independent_subprojects(repo_dir, all_files)
    print("  [DD] OS-specific commands check...")
    all_findings += detect_DD_os_specific_commands(repo_dir, all_files)
    print("  [DE] PyTorch non-determinism check...")
    all_findings += detect_DE_pytorch_nondeterminism(repo_dir, all_files)
    print("  [DF] External data URL without fetch script check...")
    all_findings += detect_DF_external_data_no_fetch(repo_dir, all_files)
    print("  [DG] Undocumented GUI/manual steps check...")
    all_findings += detect_DG_undocumented_gui_steps(repo_dir, all_files)
    print("  [SP] Specialist/proprietary software check...")
    all_findings += detect_SP_specialist_software(repo_dir, all_files)
    print("  [EP] Data provenance check...")
    all_findings += detect_EP_data_provenance(repo_dir, all_files)
    print("  [DZ] Double-zipped deposit check...")
    all_findings += detect_DZ_double_zipped(repo_dir, all_files)
    print("  [NZ] Nested zip check...")
    all_findings += detect_NZ(repo_dir, all_files)
    print("  [DUP] Duplicate data file check...")
    all_findings += detect_DUP(repo_dir, all_files)
    print("  [3D] 3D mesh no viewer check...")
    all_findings += detect_3D_mesh_no_viewer(repo_dir, all_files)
    print("  [ND] No data files check...")
    all_findings += detect_ND_no_data_files(repo_dir, all_files)
    print("  [FW] Figures in Word format check...")
    all_findings += detect_FW_figures_in_word(repo_dir, all_files)
    print("  [UE] Unicode filenames check...")
    all_findings += detect_UE_unicode_filenames(repo_dir, all_files)
    print("  [NX] No-extension files check...")
    all_findings += detect_NX_no_extension(repo_dir, all_files)
    print("  [NN] Non-sequential numbering check...")
    all_findings += detect_NN_nonsequential_numbering(repo_dir, all_files)
    print("  [IC] Inconsistent extension casing check...")
    all_findings += detect_IC_inconsistent_extension_case(repo_dir, all_files)
    print("  [IC2] Inconsistent filename spacing check...")
    all_findings += detect_IC2_inconsistent_filename_spacing(repo_dir, all_files)
    print("  [TV]  Numbering padding inconsistency check...")
    all_findings += detect_TV_numbering_padding_inconsistency(repo_dir, all_files)
    print("  [FL]  Long filename check...")
    all_findings += detect_FL_long_filenames(repo_dir, all_files)
    print("  [HS]  Human subjects data check...")
    all_findings += detect_HS_human_subjects_data(repo_dir, all_files)
    print("  [DH] Undeclared imports check...")
    all_findings += detect_DH_undeclared_imports(repo_dir, all_files)
    print("  [DI] Variable name mismatch check...")
    all_findings += detect_DI_variable_mismatch(repo_dir, all_files)
    print("  [DJ] Build artefacts check...")
    all_findings += detect_DJ_build_artefacts(repo_dir, all_files)
    print("  [DK] Filename typo check...")
    all_findings += detect_DK_filename_typos(repo_dir, all_files)
    print("  [DL] MATLAB autosave artefacts check...")
    all_findings += detect_DL_matlab_autosave(repo_dir, all_files)

    # [DB] is a more specific form of [G] for Shiny apps — when [DB] fires,
    # its "expected values for specific input combinations" requirement already
    # covers [G]'s "define what successful reproduction looks like" ask.
    if any(f['mode'] == 'DB' for f in all_findings):
        all_findings = [f for f in all_findings if f['mode'] != 'G']

    return all_findings

def detect_F_missing_seeds(repo_dir, all_files):
    """Failure Mode F: Undocumented stochasticity / missing random seeds."""
    findings = []
    code_files = [f for f in all_files
                  if f.suffix.lower() in CODE_EXTENSIONS | {'.ipynb'}]

    rng_imports = {
        'numpy': 'np.random.seed()',
        'random': 'random.seed()',
        'torch': 'torch.manual_seed()',
        'tensorflow': 'tf.random.set_seed()',
        'sklearn': 'random_state= parameter',
        # scipy removed — scipy.stats functions are deterministic; only scipy.stats.distributions random sampling needs seeding
        'lightgbm': 'random_state= parameter',
        'xgboost': 'seed= parameter',
    }

    seed_patterns = re.compile(
        r'(random\.seed|np\.random\.seed|numpy\.random\.seed'
        r'|torch\.manual_seed|tf\.random\.set_seed'
        r'|random_state\s*='
        r'|jax\.random\.PRNGKey|jax\.random\.key'
        r'|set_seed\s*\('
        r'|default_rng\s*\()',
        re.IGNORECASE
    )

    jax_import_pattern = re.compile(r'import jax|from jax')
    jax_key_pattern = re.compile(
        r'jax\.random\.PRNGKey|jax\.random\.key\s*\('
    )

    # Build set of all directories that contain __init__.py (Python packages)
    package_dirs = {
        init_file.parent for init_file in all_files
        if init_file.name == '__init__.py'
    }

    for f in code_files:
        if f.suffix.lower() not in {'.py', '.r', '.rmd', '.jl', '.ipynb'}:
            continue
        # Skip library package internals — files inside a Python package
        # directory or inside a tests/ directory are not analysis entry points
        if f.parent in package_dirs:
            continue
        parts_lower = [p.lower() for p in f.parts]
        if any(p in {'tests', 'test', 'docs', 'doc'} for p in parts_lower):
            continue
        content = read_file_safe(f)
        imported_rngs = []
        for lib, seed_fn in rng_imports.items():
            if re.search(rf'\bimport\s+{lib}\b|from\s+{lib}\s+import'
                         rf'|library\s*\(\s*["\']?{lib}',
                         content, re.IGNORECASE):
                # numpy only stochastic if np.random actually called
                import re as _re
                if lib == 'numpy' and not _re.search(r'np\.random\.|numpy\.random\.', content):
                    continue
                imported_rngs.append((lib, seed_fn))
        if imported_rngs and not seed_patterns.search(content):
            libs = ', '.join(lib for lib, _ in imported_rngs)
            findings.append(finding(
                'F', 'SIGNIFICANT',
                f'No random seed set in {f.name}',
                f'This file imports stochastic libraries ({libs}) '
                f'but no random seed was detected. Results will '
                f'differ between runs.',
                [f'Evidence: {f.name} imports {libs} without seed']
            ))
        if jax_import_pattern.search(content):
            if not jax_key_pattern.search(content):
                findings.append(finding(
                    'F', 'SIGNIFICANT',
                    f'JAX imported without PRNG key management: {f.name}',
                    'JAX uses a separate random number system from numpy. '
                    'np.random.seed() does NOT control JAX randomness. '
                    'Use jax.random.PRNGKey() or jax.random.key().',
                    [f'Evidence: {f.name} imports jax without PRNGKey']
                ))
    return findings


def detect_U_environment_variables(repo_dir, all_files):
    """Failure Mode U: Undocumented environment variables and credentials."""
    findings = []
    code_files = [f for f in all_files
                  if f.suffix.lower() in CODE_EXTENSIONS
                  and not _is_minified(f)
                  and not any(part.lower() in VENDOR_DIRS
                              for part in f.relative_to(repo_dir).parts)]

    credential_patterns = re.compile(
        r'os\.environ\.get\s*\(\s*["\']([^"\']*'
        r'(?:KEY|SECRET|TOKEN|PASSWORD|PASSWD|PWD|AUTH|API_KEY)'
        r'[^"\']*)["\']'
        r'|os\.getenv\s*\(\s*["\']([^"\']*'
        r'(?:KEY|SECRET|TOKEN|PASSWORD|PASSWD|PWD|AUTH|API_KEY)'
        r'[^"\']*)["\']',
        re.IGNORECASE
    )

    config_patterns = re.compile(
        r'os\.environ\.get\s*\(\s*["\']([^"\']+)["\']'
        r'|os\.getenv\s*\(\s*["\']([^"\']+)["\']'
        r'|os\.environ\s*\[\s*["\']([^"\']+)["\']',
        re.IGNORECASE
    )

    has_env_example = any(
        f.name.lower() in {'.env.example', '.env.sample', '.env.template'}
        for f in all_files
    )

    found_credentials = set()
    found_config = set()

    for f in code_files:
        content = read_file_safe(f)
        for match in credential_patterns.finditer(content):
            var_name = match.group(1) or match.group(2)
            if var_name:
                found_credentials.add(var_name.upper())
        for match in config_patterns.finditer(content):
            var_name = (match.group(1) or match.group(2)
                        or match.group(3))
            if var_name:
                found_config.add(var_name.upper())

    found_config -= found_credentials

    if found_credentials:
        findings.append(finding(
            'U', 'CRITICAL',
            'Credential environment variables detected',
            'This repository uses environment variables that appear '
            'to be credentials. Document in .env.example with '
            'placeholder values only.',
            [f'Variables: {", ".join(sorted(found_credentials))}']
        ))

    if found_config and not has_env_example:
        findings.append(finding(
            'U', 'SIGNIFICANT',
            'Environment variables used but no .env.example found',
            'Validators cannot know what variables to set. '
            'A .env.example will be generated.',
            [f'Variables: {", ".join(sorted(list(found_config)[:10]))}']
        ))

    return findings


def detect_E_missing_data_documentation(repo_dir, all_files):
    """Failure Mode E: Data files present but no data documentation."""
    findings = []

    # Use module-level DATA_EXTENSIONS plus a few extras that [E] specifically
    # cares about (.db/.sqlite as data stores; .xml excluded intentionally).
    # Image and 3D-mesh formats are primary data in imaging/biology/
    # palaeontology deposits and must trigger [E] the same way they trigger [BA].
    data_extensions = DATA_EXTENSIONS | {'.db', '.sqlite'} | {
        '.jpg', '.jpeg', '.png', '.tif', '.tiff', '.bmp', '.gif',
        '.webp', '.raw', '.cr2', '.nef', '.dng',
        '.off', '.ply', '.stl', '.obj', '.vtk', '.vtu', '.mesh',
        '.wrl', '.dae', '.fbx', '.glb', '.gltf',
    }

    _model_name_indicators = {'model', 'clf', 'classifier', 'regressor', 'estimator', 'pipeline', 'weights', 'tokenizer', 'vocab', 'checkpoint'}

    def _is_model_artifact(f):
        _model_dirs = {'models', 'model', 'checkpoints', 'saved_model'}
        name_lower = f.name.lower()
        ext = f.suffix.lower()
        in_model_dir = any(part.lower() in _model_dirs for part in f.parts)
        has_model_name = any(ind in name_lower for ind in _model_name_indicators)
        if ext in {'.pkl', '.pickle', '.pt', '.pth', '.onnx', '.safetensors', '.bin'}:
            return has_model_name or in_model_dir
        if ext == '.json':
            return has_model_name or in_model_dir
        return False

    # Build/IDE config files share extensions (.xml) but are not research data.
    _build_config_names = {
        'pom.xml', 'build.gradle', 'build.gradle.kts', 'build.xml',
        'ivy.xml', 'settings.xml', 'settings.gradle', 'settings.gradle.kts',
        'ant.xml', 'maven.xml', 'project.xml', 'assembly.xml',
        '.classpath', '.project', 'gradlew', 'gradlew.bat',
    }

    def _is_ide_config(f):
        name = f.name.lower()
        ext = f.suffix.lower()
        parts_lower = [p.lower() for p in f.parts]
        if ext == '.xml':
            if name.endswith('-style.xml'):        # e.g. intellij-java-google-style.xml
                return True
            if 'checkstyle' in name:               # e.g. checkstyle.xml, checkstyle-config.xml
                return True
            if '.settings' in parts_lower:         # Eclipse .settings/*.xml
                return True
        if ext == '.launch':                       # Eclipse/IntelliJ launch configs
            return True
        return False

    data_files = [
        f for f in all_files
        if (f.suffix.lower() in data_extensions
            or f.suffix.lower() in ARCHIVE_EXTENSIONS)
        and not _is_model_artifact(f)
        and f.name.lower() not in _build_config_names
        and not _is_ide_config(f)
        and not f.name.lower().startswith('readme')
        and not _in_asset_dir(f, repo_dir)
        # Exclude CSV/TSV files that are themselves structured as codebooks —
        # they ARE the documentation, not data that needs to be documented.
        # Mirrors [BA]'s data_files construction exactly.
        and not (f.suffix.lower() in {'.csv', '.tsv'} and _looks_like_codebook(f))
    ]

    if not data_files:
        return findings

    doc_indicators = [
        'codebook', 'data_dictionary', 'data-dictionary',
        'data_readme', 'data-readme',
        'column_description', 'field_description',
    ]
    # 'metadata', 'variables', 'schema' are intentionally excluded here:
    # they appear as substrings in ordinary data filenames (e.g.
    # "site_metadata.csv", "environmental_variables.csv") causing false
    # positives.  Exact-filename coverage is already provided by
    # CODEBOOK_FILENAMES below.

    all_names_lower = [f.name.lower() for f in all_files]
    all_stems_lower = [f.stem.lower() for f in all_files]

    has_data_doc = any(
        any(ind in name for ind in doc_indicators)
        for name in all_names_lower + all_stems_lower
    )
    # Exact codebook filename match
    if not has_data_doc:
        has_data_doc = any(f.name.lower() in CODEBOOK_FILENAMES for f in all_files)
    # Content-based: opaque format inspection via _inspect_file_content
    # Covers .xlsx/.xls, .docx, .pdf, stat files, netCDF, HDF5, SQLite
    _opaque_codebook_finding = None
    if not has_data_doc:
        _INSPECTABLE_E = {
            '.xlsx', '.xls', '.docx', '.pdf',
            '.dta', '.sav', '.zsav', '.sas7bdat', '.xpt',
            '.nc', '.h5', '.hdf5', '.sqlite', '.db', '.sqlite3',
        }
        for f in all_files:
            if f.suffix.lower() in _INSPECTABLE_E:
                res = _inspect_file_content(f)
                if res.has_codebook:
                    has_data_doc = True
                    _opaque_codebook_finding = finding(
                        'E', 'INFO',
                        res.inspection_note or f'Codebook content detected in: {f.name}',
                        f'Content inspection of {f.name} found variable documentation — '
                        '[E] suppressed.',
                        [f'File: {f.name}', f'Format: {f.suffix.lower()}']
                    )
                    break
    # Also treat same-named .txt files as potential data documentation
    # e.g. GDP_FDI_Dataset.txt alongside GDP_FDI_Dataset.csv
    # Minimum size threshold: tiny sidecars (< 500 bytes) are filename labels,
    # not real data documentation, and must not suppress [E].
    if not has_data_doc:
        data_stems = {f.stem.lower() for f in data_files}
        companion_txt = [f for f in all_files
                         if f.suffix.lower() == '.txt'
                         and f.stem.lower() in data_stems
                         and f.stat().st_size >= 500]
        if companion_txt:
            has_data_doc = True

    readme_mentions_data = False
    has_inline_col_doc = False
    # Patterns that indicate inline column/variable documentation in a README.
    # Any line matching one of these patterns counts toward the threshold.
    # Threshold ≥5 to avoid noise from prose that happens to contain backticks.
    _COLUMN_DOC_PATTERNS = [
        re.compile(r'`[a-zA-Z_][a-zA-Z0-9_\s\(\)\.\/\-]*`\s*[:\-—]',
                   re.MULTILINE),                                    # `col_name`: desc
        re.compile(r'^\s*-\s*[A-Za-z_][A-Za-z0-9_\s\(\)\.\/\-]*\s*[\(:–\-]',
                   re.MULTILINE),                                    # - ColumnName: desc
        re.compile(r'^\s*\*\s*[A-Za-z_][A-Za-z0-9_\s\(\)\.\/\-]*\s*[\(:–\-]',
                   re.MULTILINE),                                    # * ColumnName: desc
        re.compile(r'\*\*[a-zA-Z_]\w*\*\*\s*:',
                   re.MULTILINE),                                    # **col_name**: desc
        re.compile(r'^\s*Column\s+\d+\s*:',
                   re.MULTILINE | re.IGNORECASE),                   # Column 1: desc
    ]
    for f in sorted(all_files, key=lambda x: len(x.relative_to(repo_dir).parts)):
        nl = f.name.lower()
        if nl in README_NAMES or ('readme' in nl and f.suffix.lower() in {'.md', '.txt', '.rst', ''}):
            try:
                content = f.read_text(encoding='utf-8', errors='ignore')
                content_lower = content.lower()
                if any(phrase in content_lower for phrase in [
                    'data source', 'dataset', 'data description',
                    'variables', 'data dictionary', 'codebook',
                    'data collection', 'data format', 'column', 'field'
                ]):
                    readme_mentions_data = True
                # ≥5 matches across all patterns = README functions as codebook
                total_matches = sum(
                    len(pat.findall(content)) for pat in _COLUMN_DOC_PATTERNS
                )
                if total_matches >= 5:
                    has_inline_col_doc = True
            except Exception:
                pass
            break  # use only the shallowest README

    # Check ALL README files for explicit variable listings.
    # Patterns that indicate the README itself serves as a codebook:
    #   1. Numbered list with parenthesized variable name: "1. Forest coverage rate (forest_coverage_rate)"
    #   2. Markdown table with column headers: "| col | desc |"
    #   3. Explicit "Variables:" section header
    if not has_data_doc:
        _VARIABLE_SECTION_PATTERNS = [
            re.compile(r'\d+\.\s+\w.*\([\w_]+\)', re.MULTILINE),
            re.compile(r'\|\s*\w+\s*\|\s*\w+.*\|', re.MULTILINE),
            re.compile(r'variables?\s*:\s*\n', re.IGNORECASE | re.MULTILINE),
        ]
        for f in all_files:
            nl = f.name.lower()
            if nl in README_NAMES or ('readme' in nl and f.suffix.lower() in {'.md', '.txt', '.rst', ''}):
                try:
                    content = f.read_text(encoding='utf-8', errors='ignore')
                    if any(pat.search(content) for pat in _VARIABLE_SECTION_PATTERNS):
                        has_data_doc = True
                        break
                except Exception:
                    pass

    if not has_data_doc and not readme_mentions_data:
        data_names = [f.name for f in data_files[:5]]
        extra = f' (and {len(data_files)-5} more)' if len(data_files) > 5 else ''
        findings.append(finding(
            'E', 'SIGNIFICANT',
            f'{len(data_files)} data file(s) present but no data documentation found',
            'Data files are present but no codebook, data dictionary, '
            'or data description was found. Validators cannot assess '
            'whether the data matches what the paper describes.',
            [f'Data files: {", ".join(data_names)}{extra}',
             'Missing: codebook, data dictionary, or README data section']
        ))
    elif data_files and not has_data_doc and not has_inline_col_doc:
        findings.append(finding(
            'E', 'LOW CONFIDENCE',
            'No dedicated data documentation file found',
            'Data files are present but no dedicated codebook or '
            'data dictionary file was found.',
            [f'Data files found: {len(data_files)}']
        ))
    elif _opaque_codebook_finding:
        findings.append(_opaque_codebook_finding)

    return findings


def detect_G_inadequate_readme(repo_dir, all_files):
    """Failure Mode G: README exists but missing critical sections."""
    findings = []
    has_code = any(f.suffix.lower() in CODE_EXTENSIONS or f.name == 'Snakefile' for f in all_files)

    _README_PLAIN = {'readme.md', 'readme.txt', 'readme.rst', 'readme'}
    _README_STEM_PREFIXES = ('readme', 'read me', 'read_me')
    readme_file = None
    readme_is_pdf = False
    # Prefer plain-text READMEs; fall back to PDF/DOCX if none found.
    for f in all_files:
        if f.name.lower() in _README_PLAIN and len(f.relative_to(repo_dir).parts) <= 4:
            readme_file = f
            break
    if not readme_file:
        for f in sorted(all_files, key=lambda x: len(x.relative_to(repo_dir).parts)):
            if (f.stem.lower().startswith(_README_STEM_PREFIXES)
                    and f.suffix.lower() in {'.pdf', '.docx', '.doc', '.pages'}
                    and len(f.relative_to(repo_dir).parts) <= 4):
                readme_file = f
                readme_is_pdf = True
                break

    # Read README content once.  When no README is found, both strings are
    # left empty so that all subsequent keyword checks evaluate to False and
    # every section is flagged as missing — the correct behaviour because a
    # missing README means nothing can possibly be documented there.
    content = ''
    content_lower = ''
    if readme_file:
        try:
            content = readme_file.read_text(encoding='utf-8', errors='ignore')
            content_lower = content.lower()
        except Exception:
            pass

    # ── Data-only deposit: data-quality criteria check ───────────────────────
    # Runs regardless of README presence so it fires even when [A] fires.
    if not has_code:
        has_data = any(f.suffix.lower() in DATA_EXTENSIONS for f in all_files)
        if has_data:
            _data_quality_kws = [
                'row', 'rows', 'record', 'records', 'observation', 'observations',
                'n =', 'n=', 'sample size', 'total of', 'contains',
                'date range', 'collected between', 'period',
                'complete', 'completeness', 'missing', 'excluded',
                'expect', 'should contain', 'threshold', 'coverage',
                'total', 'count', 'entries',
            ]
            _quality_hits = sum(1 for kw in _data_quality_kws if kw in content_lower)
            if _quality_hits < 3:
                findings.append(finding(
                    'G', 'LOW CONFIDENCE',
                    'No data quality criteria documented',
                    'This is a data deposit — validators need to know what '
                    'a correct version of the data should look like. '
                    'Without an expected row count, date range, or completeness '
                    'threshold, there is no way to verify the download is complete.',
                    [
                        'Recommendation: document expected row count per file',
                        'Recommendation: document date range or collection period',
                        'Recommendation: document any exclusion criteria or known missing values',
                    ]
                ))

    # sections we expect in a reproducible research README (code deposits only)
    # Data-only deposits use the data-quality criteria check above instead.
    if has_code:
        required_sections = {
            'installation': [
                'install', 'setup', 'getting started', 'requirements',
                'dependencies', 'environment', 'pip install', 'conda'
            ],
            'execution': [
                'how to run', 'usage', 'running', 'execute', 'run the',
                'to reproduce', 'reproduc', 'quickstart', 'quick start',
                'steps to', 'instructions'
            ],
            'expected outputs': [
                'expected output', 'results', 'figures', 'tables',
                'what to expect', 'output files', 'produces',
                'generates', 'successful reproduction', 'success'
            ],
            'data': [
                'data', 'dataset', 'download', 'source', 'input'
            ],
        }

        missing = []
        for section, keywords in required_sections.items():
            if not any(kw in content_lower for kw in keywords):
                missing.append(section)

        # Post-filter: suppress sub-items where richer content-level evidence is present.
        # 'execution' — description-style narrative can substitute for an explicit section.
        _EXECUTION_INDICATORS = [
            'each section', 'in this order', 'run the script', 'execute',
            'to reproduce', 'step 1', 'step 2', 'first run', 'then run',
            'script is for', 'section of the script', 'analysis are present',
        ]
        if 'execution' in missing:
            if sum(1 for ind in _EXECUTION_INDICATORS if ind in content_lower) >= 2:
                missing.remove('execution')

        # 'expected outputs' — if README describes what each section produces (plots,
        # models, statistics), treat it as documenting expected outputs.
        _OUTPUT_INDICATORS = [
            'plot', 'figure', 'model', 'statistic', 'output', 'produces',
            'saves', 'generates', 'results in', 'table', 'coefficient',
            'regression', 'correlation', 'estimate',
        ]
        if 'expected outputs' in missing:
            if sum(1 for ind in _OUTPUT_INDICATORS if ind in content_lower) >= 2:
                missing.remove('expected outputs')

        if len(missing) >= 3:
            if readme_file:
                _g_body = (
                    f'The README ({readme_file.name}) is missing sections that '
                    'validators need to reproduce the work. Without installation '
                    'instructions, execution steps, and expected outputs, '
                    'validators cannot proceed systematically.'
                )
                _g_readme_line = (
                    f'README: {readme_file.name} '
                    + ('(PDF/binary — sections cannot be auto-checked)'
                       if readme_is_pdf else f'({len(content)} chars)')
                )
            else:
                _g_body = (
                    'No README was found — validators cannot check for required '
                    'sections (installation, execution, expected outputs, data). '
                    'Add a README documenting how to reproduce the analysis.'
                )
                _g_readme_line = 'No README file present'
            findings.append(finding(
                'G', 'LOW CONFIDENCE',
                f'README is missing critical sections: {", ".join(missing)}',
                _g_body,
                [f'Missing sections: {", ".join(missing)}', _g_readme_line]
            ))
        elif len(missing) >= 1:
            findings.append(finding(
                'G', 'LOW CONFIDENCE',
                f'README may be missing sections: {", ".join(missing)}',
                'The README appears to be missing some recommended '
                'sections. This may be intentional if the information '
                'is elsewhere, but validators may struggle to find it.',
                [f'Possibly missing: {", ".join(missing)}']
            ))

        # Code deposit — check for definition of successful reproduction
        success_indicators = [
            'successful reproduction', 'reproduction is successful',
            'expected result', 'should produce', 'should see',
            'tolerance', 'within', 'match', 'identical'
        ]
        has_success_definition = any(
            ind in content_lower for ind in success_indicators
        )
        if not has_success_definition and len(content) > 200:
            findings.append(finding(
                'G', 'SIGNIFICANT',
                'README does not define what successful reproduction looks like',
                'Without a definition of successful reproduction, '
                'validators cannot determine whether their results '
                'match the original. This is the single most important '
                'missing element in most research READMEs. '
                'What should a validator see when they have succeeded?',
                ['Missing: definition of successful reproduction',
                 'Required: expected values, tolerance bands, or '
                 'explicit comparison criteria']
            ))
    # (Data-quality check for data-only deposits is handled at the top of this
    # function so it fires even when there is no README.)

    # If SIGNIFICANT fires, suppress LOW CONFIDENCE to avoid double-reporting [G]
    if any(f['severity'] == 'SIGNIFICANT' for f in findings):
        findings = [f for f in findings if f['severity'] != 'LOW CONFIDENCE']

    return findings


def detect_H_hardcoded_versions(repo_dir, all_files):
    """Failure Mode H: Version numbers hardcoded in code not requirements."""
    findings = []
    code_files = [f for f in all_files
                  if f.suffix.lower() in CODE_EXTENSIONS]

    version_in_code = re.compile(
        r'(pandas|numpy|scipy|sklearn|matplotlib|torch|tensorflow'
        r'|keras|statsmodels|seaborn|plotly|xgboost|lightgbm'
        r'|transformers|datasets|huggingface)[=><!\s]+[\d]+\.[\d]',
        re.IGNORECASE
    )

    for f in code_files:
        content = read_file_safe(f)
        matches = version_in_code.findall(content)
        if matches:
            findings.append(finding(
                'H', 'LOW CONFIDENCE',
                f'Version constraint hardcoded in {f.name}',
                'Version constraints found inside code rather than '
                'in a dependency specification file. This can cause '
                'conflicts and makes dependency management harder. '
                'Move version constraints to requirements.txt or '
                'equivalent.',
                [f'Evidence: {f.name} — {", ".join(set(matches))[:80]}']
            ))
    return findings


def detect_K_compute_environment(repo_dir, all_files):
    """Failure Mode K: Compute environment not documented."""
    findings = []
    # skip for data-only repos
    if not any(f.suffix.lower() in CODE_EXTENSIONS for f in all_files):
        return findings

    # Trivial stdlib-only helpers don't need RAM or runtime documentation.
    _skip_resources = _is_zero_dep_reader(all_files)

    readme_file = None
    for f in all_files:
        if f.name.lower() in {'readme.md', 'readme.txt', 'readme.rst'}:
            readme_file = f
            break

    # When no README exists, content stays empty — all environment indicators
    # evaluate to False, so every missing item is flagged unconditionally.
    content = ''
    if readme_file:
        try:
            content = readme_file.read_text(
                encoding='utf-8', errors='ignore'
            ).lower()
        except Exception:
            pass

    # check for compute environment documentation
    os_indicators = [
        'ubuntu', 'windows', 'macos', 'linux', 'operating system',
        'os:', 'tested on', 'platform'
    ]
    ram_indicators = [
        'ram', 'memory', 'gb', 'gigabyte', 'minimum', 'cores'
    ]
    gpu_indicators = [
        'gpu', 'cuda', 'nvidia', 'a100', 'v100', 'rtx',
        'graphics card', 'accelerator'
    ]
    runtime_indicators = [
        'runtime', 'running time', 'minutes', 'hours',
        'approximately', 'takes', 'estimated'
    ]
    # JDK/Java version: 'jdk', 'openjdk', 'java se', or 'java <digit>'
    _java_ver_pat = re.compile(r'java\s+\d')
    jdk_indicators = ['jdk', 'openjdk', 'java se']

    # R version: "R version 4.4.0", "R 4.4", "R (>= 4.0)", session-info reference
    _r_ver_pat = re.compile(
        r'\br\s+version\s+\d'               # "R version 4.4.0"
        r'|\br\s+\d+\.\d+'                  # "R 4.4" / "R 4.4.0"
        r'|\br\s*\(>=?\s*\d'                # "R (>= 4.0)"
        r'|session[_\- ]?info',             # session-info.txt / sessionInfo()
        re.IGNORECASE
    )

    documented = []
    missing = []

    has_r_version = bool(_r_ver_pat.search(content))
    if has_r_version:
        documented.append('R version')

    if not any(ind in content for ind in os_indicators) and not has_r_version:
        missing.append('operating system')
    if not _skip_resources and not any(ind in content for ind in ram_indicators):
        missing.append('RAM/memory requirements')
    if not _skip_resources and not any(ind in content for ind in runtime_indicators):
        missing.append('estimated runtime')

    # Note documented language/runtime versions so findings acknowledge them
    if any(ind in content for ind in jdk_indicators) or _java_ver_pat.search(content):
        documented.append('JDK version')

    # GPU check only relevant if GPU libraries present
    has_gpu_libs = any(
        f.name.lower() in {'requirements.txt', 'environment.yml'}
        for f in all_files
    )
    if has_gpu_libs:
        code_content = ''
        for f in all_files:
            if f.suffix.lower() == '.py':
                code_content += read_file_safe(f).lower()
        uses_gpu = any(g in code_content for g in [
            'cuda', 'torch.cuda', '.to("cuda")', '.gpu',
            'tf.device', 'jax.devices'
        ])
        if uses_gpu and not any(
            ind in content for ind in gpu_indicators
        ):
            missing.append('GPU specification')

    if not missing:
        return findings

    _doc_note = (f'Documented: {", ".join(documented)}. ' if documented else '')
    _evidence = []
    if documented:
        _evidence.append(f'Already documented: {", ".join(documented)}')
    _evidence.append(f'Missing from README: {", ".join(missing)}')

    if len(missing) >= 2:
        findings.append(finding(
            'K', 'SIGNIFICANT',
            f'Compute environment not documented: {", ".join(missing)}',
            f'{_doc_note}Validators need to know what hardware and software '
            'environment is required to reproduce results. Without '
            'this, they may spend hours on environment issues before '
            'discovering the code requires more RAM or a GPU than '
            'they have available.',
            _evidence
        ))
    else:
        findings.append(finding(
            'K', 'LOW CONFIDENCE',
            f'Compute environment partially documented — missing: {missing[0]}',
            f'{_doc_note}Most compute environment details are present but '
            f'{missing[0]} is not mentioned.',
            _evidence
        ))

    return findings


def detect_P_preregistration(repo_dir, all_files):
    """Failure Mode P: Pre-registration mentioned but no link provided."""
    findings = []

    text_files = [
        f for f in all_files
        if f.suffix.lower() in {'.md', '.txt', '.rst', '.html'}
    ]

    prereg_mentioned = False
    prereg_link = False

    prereg_terms = [
        'pre-registr', 'preregistr', 'registered report',
        'osf.io', 'aspredicted', 'clinicaltrials',
        'protocol registration', 'pre-analysis plan',
        'preanalysis plan'
    ]

    link_pattern = re.compile(
        r'osf\.io/[a-z0-9]+|aspredicted\.org|clinicaltrials\.gov'
        r'|protocols\.io|zenodo\.org|doi\.org',
        re.IGNORECASE
    )

    for f in text_files:
        content = read_file_safe(f).lower()
        if any(term in content for term in prereg_terms):
            prereg_mentioned = True
        if link_pattern.search(content):
            prereg_link = True

    if prereg_mentioned and not prereg_link:
        findings.append(finding(
            'P', 'SIGNIFICANT',
            'Pre-registration mentioned but no link found',
            'The documentation mentions pre-registration or a '
            'registered report but no link to the pre-registration '
            'record was found. Validators cannot verify that the '
            'analysis matches the pre-registered protocol without '
            'this link.',
            ['Pre-registration terms found in documentation',
             'Missing: OSF, AsPredicted, or ClinicalTrials link']
        ))

    return findings


def detect_PAP_preregistration_document(repo_dir, all_files):
    """Positive observation: deposit contains a pre-analysis plan (PAP) or
    pre-registration document.  This is a reproducibility strength and should
    be noted in the report.  Also suppresses [E] for the PAP file itself
    (it is documentation, not a data file requiring a codebook).
    """
    findings = []
    # File-based detection: PDF/DOCX whose name contains pre-registration signals
    _PAP_NAME_RE = re.compile(
        r'pre.?anal|pre.?reg|pap[_\- ]|analysis.plan|pre.?spec|study.?plan',
        re.IGNORECASE
    )
    pap_files = [
        f for f in all_files
        if f.suffix.lower() in {'.pdf', '.docx', '.doc', '.txt', '.md'}
        and _PAP_NAME_RE.search(f.stem)
    ]
    # README-based detection: README mentions pre-registration AND a link
    _PREREG_LINK_RE = re.compile(
        r'osf\.io/[a-z0-9]+|aspredicted\.org|clinicaltrials\.gov'
        r'|protocols\.io|doi\.org/10\.',
        re.IGNORECASE
    )
    readme_has_prereg_link = False
    for f in all_files:
        if f.name.lower() in README_NAMES:
            content = read_file_safe(f)
            if _PREREG_LINK_RE.search(content) and any(
                t in content.lower()
                for t in ('pre-reg', 'preregist', 'pre-analys', 'preanalys',
                          'pap', 'registered report')
            ):
                readme_has_prereg_link = True
                break

    if pap_files or readme_has_prereg_link:
        _evidence = []
        if pap_files:
            _evidence.append(f'Pre-registration document: {", ".join(f.name for f in pap_files[:3])}')
        if readme_has_prereg_link:
            _evidence.append('README documents pre-registration with a persistent link')
        findings.append(finding(
            'PAP', 'LOW CONFIDENCE',
            'Pre-registration or pre-analysis plan present',
            'A pre-registration or pre-analysis plan document was found. '
            'This is a significant reproducibility strength — it allows '
            'validators to verify that the analyses match the pre-registered '
            'protocol and identify any unplanned deviations. '
            'Confirm that the PAP is accessible and linked from the README.',
            _evidence
        ))
    return findings


def detect_V_virtual_environment(repo_dir, all_files, existing_findings=None):
    """Failure Mode V: No virtual environment specification."""
    findings = []
    # Suppress V for Docker repos — requirements.txt is for the build layer, not local install
    if any(f.name == 'Dockerfile' or f.name.startswith('Dockerfile.') for f in all_files):
        return findings
    # Only suppress V if [B] fired for complete absence of deps (not just unpinning)
    if existing_findings:
        b_findings = [f for f in existing_findings if f.get('mode') == 'B']
        if any('no dependency' in f.get('title', '').lower() for f in b_findings):
            return findings

    has_venv_spec = any(
        f.name.lower() in {
            'environment.yml', 'environment.yaml',
            'pipfile', 'pipfile.lock',
            'poetry.lock', 'pyproject.toml',
            '.python-version', 'runtime.txt',
            'conda-lock.yml', 'setup.py', 'setup.cfg'
        }
        for f in all_files
    )

    has_requirements = any(
        f.name.lower() == 'requirements.txt'
        for f in all_files
    )

    _stata_lib_dirs = {'plus', 'personal', 'stbplus'}
    has_python = any(
        f.suffix.lower() == '.py'
        and not ('ado' in f.parts and any(p in _stata_lib_dirs for p in f.parts))
        for f in all_files
    )

    if not has_python:
        return findings

    # Zero-dep readers don't need a virtual environment — suppress entirely
    if _is_zero_dep_reader(all_files):
        return findings

    if not has_venv_spec and not has_requirements:
        findings.append(finding(
            'V', 'SIGNIFICANT',
            'No virtual environment or dependency specification found',
            'Python code is present but no virtual environment '
            'specification (environment.yml, Pipfile, pyproject.toml) '
            'or requirements.txt was found. Validators will be forced '
            'to guess which packages to install and may encounter '
            'version conflicts with their existing Python environment.',
            ['Missing: requirements.txt, environment.yml, or Pipfile']
        ))
    elif has_requirements and not has_venv_spec:
        # check if README mentions virtual environment
        readme_mentions_venv = False
        for f in all_files:
            if f.name.lower() in {'readme.md', 'readme.txt'}:
                content = read_file_safe(f).lower()
                if any(term in content for term in [
                    'venv', 'virtualenv', 'conda', 'virtual environment',
                    'python -m venv', 'conda create'
                ]):
                    readme_mentions_venv = True

        if not readme_mentions_venv:
            findings.append(finding(
                'V', 'LOW CONFIDENCE',
                'requirements.txt present but no virtual environment '
                'setup instructions found',
                'A requirements.txt exists but the README does not '
                'mention creating a virtual environment before '
                'installing. Installing into a global Python '
                'environment risks conflicts and unreproducible '
                'behaviour.',
                ['Recommendation: add venv or conda setup instructions '
                 'to README']
            ))

    return findings


def detect_I_intermediate_files(repo_dir, all_files):
    """Failure Mode I: Intermediate files committed but not regenerable."""
    findings = []

    intermediate_extensions = {
        '.pkl', '.npy', '.npz', '.rds', '.rdata',
        '.feather', '.arrow', '.parquet', '.hdf5', '.h5'
    }

    _model_name_indicators_i = {'model', 'clf', 'classifier', 'regressor', 'estimator', 'pipeline', 'weights', 'tokenizer', 'vocab', 'checkpoint'}

    def _is_model_artifact_i(f):
        _model_dirs = {'models', 'model', 'checkpoints', 'saved_model'}
        name_lower = f.name.lower()
        ext = f.suffix.lower()
        in_model_dir = any(part.lower() in _model_dirs for part in f.parts)
        has_model_name = any(ind in name_lower for ind in _model_name_indicators_i)
        if ext in {'.pkl', '.pickle', '.pt', '.pth', '.onnx', '.safetensors', '.bin'}:
            return has_model_name or in_model_dir
        if ext == '.json':
            return has_model_name or in_model_dir
        return False

    intermediate_files = [
        f for f in all_files
        if f.suffix.lower() in intermediate_extensions and not _is_model_artifact_i(f)
    ]

    # Build set of filenames written by any code file
    _written_fnames = set()
    _write_pat_i = re.compile(
        r'(?:np\.save(?:z(?:_compressed)?)?|to_csv|to_parquet|to_excel'
        r'|savetxt|pickle\.dump|joblib\.dump|torch\.save'
        r'|write\.csv|saveRDS|writeMat)'
        r'\s*\(\s*f?["\']([^"\']+)["\']',
        re.IGNORECASE
    )
    for _cf in all_files:
        if _cf.suffix.lower() in CODE_EXTENSIONS:
            try:
                _csrc = _cf.read_text(encoding='utf-8', errors='ignore')
                for _m in _write_pat_i.finditer(_csrc):
                    _fp = _m.group(1)
                    _fn = _fp.replace('\\', '/').split('/')[-1].lower()
                    if 'savez' in _m.group(0).lower() and not _fn.endswith('.npz'):
                        _fn = _fn + '.npz'
                    _written_fnames.add(_fn)
            except Exception:
                pass
    # Exclude files in data/ that are never written by code (pure raw inputs)
    intermediate_files = [
        f for f in intermediate_files
        if f.name.lower() in _written_fnames
        or not any(part.lower() == 'data' for part in f.parts)
    ]

    if not intermediate_files:
        return findings

    # check if these files are generated by the code
    code_content = ''
    for f in all_files:
        if f.suffix.lower() in CODE_EXTENSIONS:
            code_content += read_file_safe(f)

    untraced = []
    for f in intermediate_files:
        stem = f.stem.lower()
        name = f.name.lower()
        # check if any code writes this file
        if not any(
            ref in code_content.lower()
            for ref in [stem, name, f.suffix.lower()]
        ):
            untraced.append(f.name)

    if intermediate_files and not untraced:
        # files exist and are referenced in code - still flag
        findings.append(finding(
            'I', 'LOW CONFIDENCE',
            f'{len(intermediate_files)} intermediate data file(s) committed',
            'Intermediate files are committed to the repository. '
            'If these are generated by the pipeline, validators '
            'need to know whether to regenerate them or use the '
            'committed versions. Committed intermediates can mask '
            'reproducibility failures if the generation step is skipped.',
            [f'Intermediate files: '
             f'{", ".join(f.name for f in intermediate_files[:5])}',
             'Clarify in README: should validators regenerate these?']
        ))
    elif untraced:
        findings.append(finding(
            'I', 'SIGNIFICANT',
            f'Intermediate files present with no apparent generation code',
            'Intermediate data files are committed but no code that '
            'generates them was found. Validators cannot reproduce '
            'these files from scratch, creating a gap in the '
            'reproducibility chain.',
            [f'Untraced files: {", ".join(untraced[:5])}']
        ))

    return findings


def detect_J_notebook_order(repo_dir, all_files):
    """Failure Mode J: Notebooks with unclear or non-linear execution order."""
    findings = []

    notebooks = [
        f for f in all_files
        if f.suffix.lower() == '.ipynb'
    ]

    if not notebooks:
        return findings

    for nb in notebooks:
        try:
            import json as _json
            content = nb.read_text(encoding='utf-8', errors='ignore')
            data = _json.loads(content)
            cells = data.get('cells', [])

            # check execution counts
            exec_counts = []
            for cell in cells:
                if cell.get('cell_type') == 'code':
                    ec = cell.get('execution_count')
                    if ec is not None:
                        exec_counts.append(ec)

            if not exec_counts:
                # Distinguish "cleared before sharing" (outputs present → [CF]
                # would also fire) from "never run" (no outputs at all).
                has_outputs = any(
                    cell.get('outputs')
                    for cell in cells
                    if cell.get('cell_type') == 'code'
                )
                if has_outputs:
                    findings.append(finding(
                        'J', 'SIGNIFICANT',
                        f'Execution counts cleared before sharing: {nb.name}',
                        'Execution counts were cleared before sharing. '
                        'Cell outputs are present but the original run '
                        'order cannot be verified. Re-run from scratch '
                        '(Kernel > Restart & Run All) to confirm outputs '
                        'are reproducible.',
                        [f'Evidence: {nb.name} — execution counts null, '
                         'cell outputs present']
                    ))
                else:
                    findings.append(finding(
                        'J', 'SIGNIFICANT',
                        f'Notebook has no execution counts: {nb.name}',
                        'This notebook has never been run top-to-bottom '
                        'with saved outputs, or outputs were cleared before '
                        'sharing. Validators cannot verify what the original '
                        'outputs looked like.',
                        [f'Evidence: {nb.name} — all execution counts null']
                    ))
            else:
                # check for non-linear execution
                non_none = [e for e in exec_counts if e is not None]
                if non_none != sorted(non_none):
                    findings.append(finding(
                        'J', 'SIGNIFICANT',
                        f'Notebook cells executed out of order: {nb.name}',
                        'Cell execution counts are not sequential, '
                        'meaning the notebook was not run top-to-bottom. '
                        'Results may depend on a specific non-linear '
                        'execution order that is not documented.',
                        [f'Evidence: {nb.name} — execution order: '
                         f'{non_none[:10]}']
                    ))
        except Exception:
            continue

    return findings


def detect_M_python_version_conflict(repo_dir, all_files):
    """Failure Mode M: Multiple or conflicting Python versions referenced."""
    findings = []
    # skip if no Python files present
    if not any(f.suffix.lower() == '.py' for f in all_files):
        return findings

    version_pattern = re.compile(r'(?<![a-zA-Z])python\s*[=><!\s]+\s*(\d+\.\d+)', re.IGNORECASE)
    versions_found = {}

    check_files = [
        f for f in all_files
        if f.name.lower() in {
            'requirements.txt', 'requirements_extra.txt', 'environment.yml', 'environment.yaml',
            'pipfile', 'pyproject.toml', 'setup.py', 'setup.cfg',
            'runtime.txt', '.python-version', 'readme.md', 'readme.txt'
        }
    ]

    for f in check_files:
        content = read_file_safe(f)
        matches = version_pattern.findall(content)
        if matches:
            versions_found[f.name] = matches

    all_versions = set(
        v for versions in versions_found.values() for v in versions
    )

    if len(all_versions) > 1:
        findings.append(finding(
            'M', 'SIGNIFICANT',
            f'Conflicting Python versions referenced: '
            f'{", ".join(sorted(all_versions))}',
            'Different files specify different Python versions. '
            'This creates ambiguity about which version was used '
            'to produce the published results. Validators will '
            'not know which to install.',
            [f'{fname}: {", ".join(v)}'
             for fname, v in versions_found.items()]
        ))
    elif len(all_versions) == 0:
        findings.append(finding(
            'M', 'LOW CONFIDENCE',
            'Python version not specified anywhere',
            'No Python version requirement was found in any '
            'configuration file. Validators will install their '
            'default Python version which may not match what '
            'was used for the original analysis.',
            ['Recommendation: add python=3.x to environment.yml '
             'or add .python-version file']
        ))

    return findings


def detect_L_large_files_missing(repo_dir, all_files):
    """Failure Mode L: Code references files that appear to be missing."""
    findings = []
    code_files = [f for f in all_files
                  if f.suffix.lower() in CODE_EXTENSIONS]

    all_filenames = {f.name.lower() for f in all_files}
    all_stems = {f.stem.lower() for f in all_files}

    read_pattern = re.compile(
        r'(?:pd\.read_csv|pd\.read_parquet|pd\.read_excel'
        r'|pd\.read_stata|pd\.read_sas|pd\.read_feather'
        r'|xr\.open_dataset|xr\.open_dataarray|netCDF4\.Dataset|nc\.Dataset'
        r'|h5py\.File|rasterio\.open|gdal\.Open'
        r'|np\.load|open|read_csv|read_parquet|loadtxt'
        r'|readRDS|read\.csv|read_dta|haven::read'
        r'|SeqIO\.parse|read\.FASTA|read\.alignment|nib\.load|nibabel\.load|read\.csv|read_csv|gpd\.read_file|fiona\.open|load)'
        r'\s*\(\s*["\']([^"\']+)["\']',
        re.IGNORECASE
    )

    # build set of files generated by the code (intermediate outputs)
    write_pattern = re.compile(
        r'(?:to_csv|to_parquet|to_excel|to_stata|savetxt|save|write_csv|saveRDS|write\.csv'
        r'|write\.table|fwrite|writeMat|csvwrite)'
        r'\s*\(\s*f?["\']([^"\']+)["\']',
        re.IGNORECASE
    )
    # torch.save(obj, filepath) — filename is second argument
    torch_save_pattern = re.compile(
        r'torch\.save\s*\([^,]+,\s*f?["\']([^"\']+\.(?:pt|pth|bin|ckpt))["\']',
        re.IGNORECASE
    )
    # np.savez('path/file', ...) — appends .npz automatically, first arg is path
    savez_pattern = re.compile(
        r'np\.savez(?:_compressed)?\s*\(\s*f?["\']([^"\']+)["\']',
        re.IGNORECASE
    )
    # R-style: write.csv(data, 'filename') — filename is second argument
    write_pattern_r = re.compile(
        r'(?:write\.csv|write\.table|saveRDS|fwrite)'
        r'\s*\([^,]+,\s*f?["\']([^"\']+)["\']',
        re.IGNORECASE
    )
    # also catch filenames assigned to variables then passed to write functions
    varname_pattern = re.compile(
        r'([\w_]+)\s*=\s*["\']([^"\']*\.(?:csv|dta|xlsx|parquet|rds))["\']\s*\n'
        r'.*?\1',
        re.IGNORECASE | re.DOTALL
    )
    generated_files = set()
    for f in code_files:
        content = read_file_safe(f)
        # resolve one level of variable assignment
        var_assign = re.findall(
            r'([A-Z_][A-Z0-9_]*)\s*=\s*["\']([^"\']*\.(?:csv|dta|xlsx|parquet|rds))["\']\s*',
            content
        )
        var_map = {v: p for v, p in var_assign}
        # check for to_csv(VAR) patterns
        for var, path in var_map.items():
            if re.search(r'to_csv\s*\(\s*' + var + r'[,)]', content):
                fname = path.replace('\\', '/').split('/')[-1].lower()
                if fname and '.' in fname:
                    generated_files.add(fname)
        for match in write_pattern.finditer(content):
            filepath = match.group(1)
            fname = filepath.replace('\\', '/').split('/')[-1].lower()
            if fname and '.' in fname:
                generated_files.add(fname)
        for match in write_pattern_r.finditer(content):
            filepath = match.group(1)
            fname = filepath.replace('\\', '/').split('/')[-1].lower()
            if fname and '.' in fname:
                generated_files.add(fname)
        for match in torch_save_pattern.finditer(content):
            filepath = match.group(1)
            fname = filepath.replace('\\', '/').split('/')[-1].lower()
            if fname and '.' in fname:
                generated_files.add(fname)
        for match in savez_pattern.finditer(content):
            filepath = match.group(1)
            if not filepath.endswith('.npz'):
                filepath = filepath + '.npz'
            fname = filepath.replace('\\', '/').split('/')[-1].lower()
            if fname and '.' in fname:
                generated_files.add(fname)

    # Build a line-numbered map of R-script-generated files for output-vs-input annotation.
    # These patterns are NOT added to generated_files so that files written and then
    # re-read within the same script still appear in missing_refs and can be annotated
    # as "will exist after first run" rather than silently suppressed.
    _WRITE_PATS_R_ANNOT = [
        re.compile(r'write\.xlsx\s*\([^,]+,\s*["\']([^"\']+)["\']', re.IGNORECASE),
        re.compile(r'write_xlsx\s*\([^,]+,\s*["\']([^"\']+)["\']', re.IGNORECASE),
        re.compile(r'ggsave\s*\(["\']([^"\']+)["\']', re.IGNORECASE),
        re.compile(r'pdf\s*\(["\']([^"\']+)["\']', re.IGNORECASE),
        re.compile(r'png\s*\(["\']([^"\']+)["\']', re.IGNORECASE),
        re.compile(r'write\.csv\s*\([^,]+,\s*["\']([^"\']+)["\']', re.IGNORECASE),
        re.compile(r'saveRDS\s*\([^,]+,\s*["\']([^"\']+)["\']', re.IGNORECASE),
    ]
    # fname_lower → (script_name, line_number)
    script_output_lines = {}
    for f in code_files:
        if f.suffix.lower() not in {'.r', '.rmd', '.qmd'}:
            continue
        try:
            for line_num, line in enumerate(
                    f.read_text(encoding='utf-8', errors='ignore').splitlines(), 1):
                for pat in _WRITE_PATS_R_ANNOT:
                    m = pat.search(line)
                    if m:
                        fpath = m.group(1).replace('\\', '/').split('/')[-1].lower()
                        if fpath and '.' in fpath and fpath not in script_output_lines:
                            script_output_lines[fpath] = (f.name, line_num)
        except Exception:
            pass
    missing_refs = set()
    # Broad scan of R files: any quoted path with data extension
    r_path_pat = re.compile(
        r'["\']([^"\'\']+\.(?:csv|tsv|rds|rdata|xlsx|dta|sav|txt|gz|zip|nii|mat|fasta|fastq|bam|vcf|bed|bw|bigwig|gff|gtf|geojson|shp|nf|config))["\']',
        re.IGNORECASE
    )
    for f in all_files:
        if f.suffix.lower() in {'.r', '.rmd', '.qmd'}:
            try:
                r_src = f.read_text(encoding='utf-8', errors='ignore')
                for m in r_path_pat.finditer(r_src):
                    fpath = m.group(1).replace('\\', '/').lstrip('./')
                    fname = fpath.split('/')[-1].lower()
                    if fname and not any(af.name.lower() == fname for af in all_files):
                        if fpath not in generated_files and fname not in generated_files:
                            missing_refs.add(fname)
            except Exception:
                pass
    # Also scan shell scripts for output files (redirect > or -o flag)
    shell_write = re.compile(
        r'(?:>\s*|(?:-o|--out(?:put)?)\s+)([\w./\-]+\.(?:txt|csv|tsv|bam|sam|vcf|gz|pdf|png|svg|html))',
        re.IGNORECASE
    )
    for f in all_files:
        if f.suffix.lower() in {'.sh', '.bash'}:
            try:
                sh_content = f.read_text(encoding='utf-8', errors='ignore')
            except Exception:
                sh_content = ''
            for m in shell_write.finditer(sh_content):
                fname = m.group(1).replace('\\', '/').split('/')[-1].lower()
                if fname and '.' in fname:
                    generated_files.add(fname)

    # Scan Nextflow files for params referencing data files
    for f in all_files:
        if f.suffix.lower() == '.nf':
            try:
                nf_src = f.read_text(encoding='utf-8', errors='ignore')
                # params.x = 'data/...' patterns
                for m in re.finditer(r"params\.\w+\s*=\s*'([^']+\.(?:gz|fastq|fasta|bam|vcf|bed|db))", nf_src):
                    fpath = m.group(1)
                    fname = fpath.replace('\\', '/').split('/')[-1].lower()
                    # Also check directory references
                    dname = fpath.replace('\\', '/').split('/')[-1].lower()
                    if not any(af.name.lower() == fname for af in all_files):
                        missing_refs.add(fpath)
                # glob patterns like data/*_{1,2}.fastq.gz
                for m in re.finditer(r"'(data/[^']*\.(?:fastq\.gz|fasta|bam|fastq))'", nf_src):
                    missing_refs.add('FASTQ input data: ' + m.group(1))
            except Exception:
                pass
    # Also scan notebook cell sources for quoted file paths
    import json as _json
    for nb in all_files:
        if nb.suffix.lower() == '.ipynb':
            try:
                nb_data = _json.loads(nb.read_text(encoding='utf-8', errors='ignore'))
                for cell in nb_data.get('cells', []):
                    src = ''.join(cell.get('source', []))
                    for match in read_pattern.finditer(src):
                        fpath = match.group(1)
                        fname = fpath.replace('\\', '/').split('/')[-1].lower()
                        if fname and '.' in fname:
                            # check if file exists
                            if not any(f.name.lower() == fname for f in all_files):
                                missing_refs.add(fname)
                    # Also catch string literals with data file extensions
                    for m in re.finditer(r'["\']([^"\']+\.(?:nii|nii\.gz|npy|npz|mat|csv|tsv|fasta|fastq|gz|bam|vcf))["\']', src, re.IGNORECASE):
                        fpath = m.group(1)
                        fname = fpath.replace('\\', '/').split('/')[-1].lower()
                        if fname and not any(f.name.lower() == fname for f in all_files):
                            missing_refs.add(fname)
            except Exception:
                pass
    for f in code_files:
        content = read_file_safe(f)
        for match in read_pattern.finditer(content):
            filepath = match.group(1)
            fname = filepath.replace('\\', '/').split('/')[-1].lower()
            stem = fname.rsplit('.', 1)[0] if '.' in fname else fname
            if fname and '.' in fname:
                # Exclude model weight files — covered by CS/CV detectors, not L
                _model_exts = {'.pt', '.pth', '.ckpt', '.bin', '.safetensors', '.onnx'}
                if fname.endswith(tuple(_model_exts)):
                    continue
                # Exclude model weight files — covered by CS/CV detectors, not L
                _model_exts = {'.pt', '.pth', '.ckpt', '.bin', '.safetensors', '.onnx'}
                if fname.endswith(tuple(_model_exts)):
                    continue
                if (fname not in all_filenames
                        and fname not in generated_files
                        and stem not in all_stems
                        and not filepath.startswith(('http', 'ftp', '$', '{'))):
                    missing_refs.add(fname)

    def _similar_file(missing_name, candidates):
        """Return a present file that looks like a renamed version of missing_name."""
        _ver_pat = re.compile(r'[_\-]?(final\d*|v\d+|\d{8})$', re.IGNORECASE)
        m_stem = _ver_pat.sub('', re.sub(r'\.[^.]+$', '', missing_name)).lower()
        m_tokens = set(re.split(r'[_\-]', m_stem)) - {''}
        for f in candidates:
            f_stem = _ver_pat.sub('', f.stem).lower()
            f_tokens = set(re.split(r'[_\-]', f_stem)) - {''}
            if m_stem == f_stem or m_stem in f_stem or f_stem in m_stem:
                return f
            # token-set equality handles word-reorder renames
            # (e.g. vsip_coverages_ag → vsip_ag_coverages)
            if len(m_tokens) >= 2 and m_tokens == f_tokens:
                return f
        return None

    if missing_refs:
        # Split into: genuine missing inputs vs script-generated files read back as inputs.
        output_reread   = {mf: script_output_lines[mf]
                           for mf in missing_refs if mf in script_output_lines}
        genuine_missing = {mf for mf in missing_refs if mf not in script_output_lines}

        evidence = []

        if genuine_missing:
            sample_g = sorted(genuine_missing)[:5]
            _ng = len(genuine_missing) - 5
            extra_g = f' (and {_ng} more {"file" if _ng == 1 else "files"})' if len(genuine_missing) > 5 else ''
            evidence.append(
                f'Missing files referenced (inputs not in repository): '
                f'{", ".join(sample_g)}{extra_g}'
            )
            evidence.append('These files must be deposited or download instructions provided.')
            # Similarity hints for genuine missing files only
            _data_candidates = [f for f in all_files if f.suffix.lower() in DATA_EXTENSIONS]
            for mf in sample_g:
                similar = _similar_file(mf, _data_candidates)
                if similar:
                    evidence.append(
                        f'Possible renamed version: `{similar.name}` '
                        f'(referenced as `{mf}`) — verify this is the correct file '
                        f'and update the code path if so'
                    )

        if output_reread:
            evidence.append(
                'Script-generated files read back as inputs '
                '(will exist after first run — fix absolute paths):'
            )
            for mf, (script_name, line_num) in sorted(output_reread.items())[:5]:
                evidence.append(f'- {mf} — written at line {line_num} in {script_name}, read back later')
            if len(output_reread) > 5:
                _n = len(output_reread) - 5
                evidence.append(f'  (and {_n} more script-generated {"file" if _n == 1 else "files"})')
            evidence.append(
                'These will be generated when the script runs — '
                'no action needed except fixing absolute paths ([C]).'
            )

        if not genuine_missing:
            # All missing refs are script outputs; keep as SIGNIFICANT but soften description
            description = (
                'The code reads files that it also writes as intermediate outputs. '
                'These files do not exist in the repository but will be created on '
                'first run. Check for absolute paths ([C]) that would prevent the '
                'script from finding them.'
            )
        else:
            description = (
                'The code attempts to read files that are not present '
                'in the repository. These may be large data files that '
                'were excluded, external downloads, or files that were '
                'accidentally omitted. Validators cannot run the code '
                'without these files.'
            )

        findings.append(finding(
            'L', 'SIGNIFICANT',
            f'Code references {len(missing_refs)} file(s) not found in repository',
            description,
            evidence
        ))

    return findings


def detect_O_output_not_committed(repo_dir, all_files):
    """Failure Mode O: No committed outputs to compare against."""
    findings = []

    output_extensions = {
        '.txt', '.csv', '.xlsx', '.html', '.pdf',
        '.png', '.jpg', '.svg', '.eps', '.tex'
    }

    # look for results/output directories
    result_dir_names = {
        'results', 'output', 'outputs', 'figures',
        'tables', 'plots', 'charts'
    }

    all_dirs = {f.parent.name.lower() for f in all_files}
    has_results_dir = bool(result_dir_names & all_dirs)

    output_files = [
        f for f in all_files
        if f.suffix.lower() in output_extensions
        and f.parent.name.lower() in result_dir_names
    ]

    has_python = any(f.suffix.lower() == '.py' for f in all_files)

    # Dataset archive heuristic: if the repo has at least one large committed
    # data file and very few code files, it is a dataset deposit — the data IS
    # the output. [O] should not fire because there is no computational pipeline
    # to produce outputs from.
    _data_exts = {
        '.json', '.csv', '.tsv', '.parquet', '.db', '.sqlite', '.sqlite3',
        '.feather', '.h5', '.hdf5', '.pkl', '.pickle', '.npy', '.npz',
    }
    _large_data = [
        f for f in all_files
        if f.suffix.lower() in _data_exts and f.stat().st_size > 100_000
    ]
    _py_count = sum(1 for f in all_files if f.suffix.lower() == '.py')
    _is_dataset_archive = len(_large_data) >= 1 and _py_count <= 3

    if has_python and not output_files and not has_results_dir and not _is_dataset_archive:
        findings.append(finding(
            'O', 'SIGNIFICANT',
            'No committed outputs found for comparison',
            'No result files, figures, or tables were found in '
            'standard output directories. Without committed outputs, '
            'validators have no reference to compare their results '
            'against. Even a single representative output file '
            'significantly improves reproducibility verification.',
            ['Missing: results/, output/, or figures/ directory '
             'with committed outputs',
             'Recommendation: commit key tables and figures from '
             'the paper']
        ))
    elif has_python and output_files:
        findings.append(finding(
            'O', 'LOW CONFIDENCE',
            f'{len(output_files)} output file(s) committed — '
            f'verify these match paper',
            'Output files are committed. Validators will compare '
            'their results against these. Ensure these files were '
            'generated by the committed code, not manually edited.',
            [f'Output files: '
             f'{", ".join(f.name for f in output_files[:5])}']
        ))

    return findings


def detect_Q_config_files(repo_dir, all_files):
    """Failure Mode Q: Configuration files missing or undocumented."""
    findings = []
    code_files = [f for f in all_files
                  if f.suffix.lower() in CODE_EXTENSIONS
                  and not _is_minified(f)
                  and not any(part.lower() in VENDOR_DIRS
                              for part in f.relative_to(repo_dir).parts)]

    config_file_pattern = re.compile(
        r'["\']([^"\']+\.(?:yaml|yml|json|toml|ini|cfg|conf))["\']',
        re.IGNORECASE
    )

    all_filenames_lower = {f.name.lower() for f in all_files}
    missing_configs = set()

    for f in code_files:
        content = read_file_safe(f)
        for match in config_file_pattern.finditer(content):
            cfg_ref = match.group(1)
            # Skip f-string / %-format / str.format template placeholders
            # e.g. "{corpus}_{task}.json" or "%s_results.yaml"
            if '{' in cfg_ref or '}' in cfg_ref or '%s' in cfg_ref or '%d' in cfg_ref:
                continue
            cfg_file = cfg_ref.split('/')[-1].lower()
            if cfg_file not in all_filenames_lower:
                missing_configs.add(cfg_file)

    if missing_configs:
        findings.append(finding(
            'Q', 'SIGNIFICANT',
            f'Configuration files referenced but not found: '
            f'{", ".join(sorted(missing_configs)[:5])}',
            'The code references configuration files that are not '
            'present in the repository. Validators cannot run the '
            'code with the same settings used in the original analysis.',
            [f'Missing configs: {", ".join(sorted(missing_configs)[:5])}']
        ))
    return findings


def detect_R_statistical_tests_undocumented(repo_dir, all_files):
    """Failure Mode R: Statistical tests used but assumptions not documented."""
    findings = []
    code_files = [f for f in all_files
                  if f.suffix.lower() in CODE_EXTENSIONS]

    stat_patterns = re.compile(
        r'\b(OLS|WLS|GLS|2SLS|IV|GMM|logit|probit|tobit'
        r'|t\.test|chi\.sq|anova|kruskal|wilcox|mann.whitney'
        r'|LinearRegression|LogisticRegression|statsmodels'
        r'|smf\.ols|sm\.OLS|ivreg|feols|felm'
        r'|fixed.effect|random.effect|panel)\b',
        re.IGNORECASE
    )

    assumption_patterns = re.compile(
        r'\b(heteroskedastic|robust|cluster|bootstrap'
        r'|standard.error|HAC|Newey.West|White'
        r'|vif|multicollin|autocorrelation|serial.correlation'
        r'|hausman|endogeneit)\b',
        re.IGNORECASE
    )

    stat_methods_found = set()
    assumptions_documented = False

    for f in code_files:
        content = read_file_safe(f)
        methods = stat_patterns.findall(content)
        if methods:
            stat_methods_found.update(methods)
        if assumption_patterns.search(content):
            assumptions_documented = True

    if stat_methods_found and not assumptions_documented:
        findings.append(finding(
            'R', 'LOW CONFIDENCE',
            f'Statistical methods detected with no assumption checks found',
            'The code uses statistical methods that have assumptions '
            '(normality, homoskedasticity, independence, etc.) but '
            'no assumption-checking code was detected. Validators '
            'cannot verify that the methods were appropriately applied.',
            [f'Methods detected: '
             f'{", ".join(sorted(stat_methods_found)[:8])}',
             'Recommendation: document assumption checks or '
             'reference where they appear in the paper']
        ))

    return findings


def detect_S_software_citations_missing(repo_dir, all_files):
    """Failure Mode S: Key software used but not cited."""
    findings = []

    major_packages = {
        'numpy', 'pandas', 'scipy', 'matplotlib', 'sklearn',
        'scikit-learn', 'statsmodels', 'torch', 'pytorch',
        'tensorflow', 'keras', 'stata', 'matlab'
    }

    readme_file = None
    for f in all_files:
        if f.name.lower() in {'readme.md', 'readme.txt', 'readme.rst'}:
            readme_file = f
            break

    if not readme_file:
        return findings

    try:
        readme_content = readme_file.read_text(
            encoding='utf-8', errors='ignore'
        ).lower()
    except Exception:
        return findings

    has_citations = any(term in readme_content for term in [
        'citation', 'cite', 'reference', 'bibliography',
        'doi:', 'zenodo', 'joss', 'journal of open source'
    ])

    if not has_citations:
        # check what packages are used
        imports_found = set()
        for f in all_files:
            if f.suffix.lower() == '.py':
                content = read_file_safe(f).lower()
                for pkg in major_packages:
                    if f'import {pkg}' in content or f'from {pkg}' in content:
                        imports_found.add(pkg)

        if imports_found:
            findings.append(finding(
                'S', 'LOW CONFIDENCE',
                'No software citations found in README',
                'Major software packages are used but no citations '
                'or references section was found in the README. '
                'Software citation is increasingly required by '
                'journals and supports reproducibility by '
                'identifying exact software versions.',
                [f'Packages used: {", ".join(sorted(imports_found))}',
                 'Recommendation: add citations for key packages']
            ))

    return findings


def detect_T_test_coverage(repo_dir, all_files):
    """Failure Mode T: No tests present for analysis code."""
    findings = []

    has_python = any(f.suffix.lower() == '.py' for f in all_files)
    if not has_python:
        return findings

    test_indicators = [
        'test_', '_test.py', 'tests/', 'test/', 'spec/',
        'pytest', 'unittest', 'nose'
    ]

    has_tests = any(
        any(ind in f.name.lower() or ind in str(f).lower()
            for ind in test_indicators)
        for f in all_files
    )

    code_files = [
        f for f in all_files
        if f.suffix.lower() == '.py'
        and not any(t in f.name.lower() for t in ['test_', '_test'])
    ]

    if not has_tests and len(code_files) > 3:
        findings.append(finding(
            'T', 'LOW CONFIDENCE',
            'No test files found',
            'No automated tests were found for the analysis code. '
            'Tests are not required for reproducibility but their '
            'absence means there is no automated way to verify '
            'that helper functions produce expected outputs. '
            'Even simple smoke tests significantly improve '
            'validator confidence.',
            [f'Python files without tests: {len(code_files)}',
             'Recommendation: add pytest tests for key functions']
        ))

    return findings


def detect_X_no_container(repo_dir, all_files):
    """Failure Mode X: No containerisation or environment isolation."""
    findings = []

    container_files = {
        'dockerfile', 'docker-compose.yml', 'docker-compose.yaml',
        'singularity', 'singularity.def', 'apptainer.def',
        '.devcontainer', 'devcontainer.json'
    }

    has_container = any(
        f.name.lower() in container_files
        for f in all_files
    )

    has_environment_yml = any(
        f.name.lower() in {'environment.yml', 'environment.yaml'}
        for f in all_files
    )

    has_python = any(f.suffix.lower() == '.py' for f in all_files)

    if has_python and not has_container and not has_environment_yml:
        if _is_zero_dep_reader(all_files):
            return findings
        findings.append(finding(
            'X', 'LOW CONFIDENCE',
            'No containerisation or conda environment found',
            'No Dockerfile, Docker Compose, Singularity, or '
            'conda environment file was found. Without environment '
            'isolation, dependency conflicts between the validator\'s '
            'system and the required packages may prevent reproduction. '
            'A conda environment.yml or Dockerfile is the most '
            'reliable way to ensure environment reproducibility.',
            ['Recommendation: add environment.yml or Dockerfile',
             'Minimum: ensure requirements.txt has pinned versions']
        ))

    return findings


def detect_Y_data_source_missing(repo_dir, all_files):
    """Failure Mode Y: Data files present but no source or provenance."""
    findings = []

    # Use the module-level DATA_EXTENSIONS constant so any future additions
    # are automatically reflected here (avoids the previous local-copy drift).
    _model_name_indicators = {'model', 'clf', 'classifier', 'regressor', 'estimator', 'pipeline', 'weights', 'tokenizer', 'vocab', 'checkpoint'}

    def _is_model_artifact(f):
        _model_dirs = {'models', 'model', 'checkpoints', 'saved_model'}
        name_lower = f.name.lower()
        ext = f.suffix.lower()
        in_model_dir = any(part.lower() in _model_dirs for part in f.parts)
        has_model_name = any(ind in name_lower for ind in _model_name_indicators)
        if ext in {'.pkl', '.pickle', '.pt', '.pth', '.onnx', '.safetensors', '.bin'}:
            return has_model_name or in_model_dir
        if ext in {'.json', '.jsonl', '.ndjson'}:
            return has_model_name or in_model_dir
        return False

    # Include image and 3D-mesh formats as data — these are primary data in
    # imaging/biology/palaeontology deposits and must trigger [Y] the same
    # way they trigger [BA].
    _Y_EXTENDED_EXTS = DATA_EXTENSIONS | ARCHIVE_EXTENSIONS | {
        '.jpg', '.jpeg', '.png', '.tif', '.tiff', '.bmp', '.gif',
        '.webp', '.raw', '.cr2', '.nef', '.dng',
        '.off', '.ply', '.stl', '.obj', '.vtk', '.vtu', '.mesh',
        '.wrl', '.dae', '.fbx', '.glb', '.gltf',
    }
    data_files = [
        f for f in all_files
        if f.suffix.lower() in _Y_EXTENDED_EXTS
        and not _is_model_artifact(f)
        and not f.name.lower().startswith('readme')
        and not _in_asset_dir(f, repo_dir)
    ]

    if not data_files:
        return findings

    # look for source/provenance documentation
    source_indicators = [
        'download', 'source', 'obtain', 'access',
        'available at', 'retrieved from', 'collected from',
        'provided by', 'doi:', 'url:', 'http', 'database',
        'data availability', 'data access'
    ]

    # Broaden README lookup: accept any file containing 'readme' in name
    # (catches prefixed names like 61622524_README_RTMS.txt)
    readme_file = None
    for f in sorted(all_files, key=lambda x: len(x.relative_to(repo_dir).parts)):
        nl = f.name.lower()
        if nl in README_NAMES or (
            'readme' in nl
            and f.suffix.lower() in {'.md', '.txt', '.rst', ''}
        ):
            readme_file = f
            break

    # Patterns that indicate the author has documented restricted/controlled access.
    # These are checked FIRST — access restriction language is itself a form of
    # provenance documentation and downgrades [Y] to LOW CONFIDENCE.
    _ACCESS_RESTRICTION_PATTERNS = [
        r'not publicly available',
        r'data use agreement',
        r'upon request',
        r'contact .{0,40}author',
        r'restricted access',
        r'available on request',
        r'requests? .{0,40}should be directed',
        r'subject to .{0,40}agreement',
        r'provided with permission',
        r'available upon reasonable request',
        r'available from the .{0,40}author',
        r'embargo',
        r'data availability statement',
    ]

    readme_content = None
    if readme_file:
        try:
            readme_content = readme_file.read_text(encoding='utf-8', errors='ignore')
        except Exception:
            pass

    # If no plain-text README was found, try to extract text from .docx/.pdf
    # README-like files so that [Y] can check their content for source indicators.
    if readme_content is None:
        _README_LIKE_Y = {'readme', 'read_me', 'read me', 'methods', 'protocol'}
        for f in sorted(all_files, key=lambda x: len(x.relative_to(repo_dir).parts)):
            if (f.suffix.lower() in {'.docx', '.pdf'}
                    and (f.stem.lower() in _README_LIKE_Y
                         or 'readme' in f.name.lower())
                    and len(f.relative_to(repo_dir).parts) <= 4):
                res = _inspect_file_content(f)
                if res.extracted_text:
                    readme_content = res.extracted_text
                    break

    has_access_restriction = readme_content is not None and any(
        re.search(p, readme_content, re.IGNORECASE)
        for p in _ACCESS_RESTRICTION_PATTERNS
    )
    has_source = readme_content is not None and any(
        ind in readme_content.lower() for ind in source_indicators
    )

    # Priority: access restriction > source documented > neither
    if has_access_restriction:
        findings.append(finding(
            'Y', 'LOW CONFIDENCE',
            'Data files present — access conditions documented in README',
            'Data files are present and the README documents access '
            'conditions (e.g. data not publicly available, available on '
            'request, or subject to a data use agreement). Validators '
            'should confirm that the README includes enough detail for '
            'a reader to understand how to obtain the data.',
            [f'Data files: {", ".join(f.name for f in data_files[:5])}',
             f'README: {readme_file.name if readme_file else "unknown"}',
             'Confirm: access instructions are clear and sufficient']
        ))
    elif not has_source:
        findings.append(finding(
            'Y', 'SIGNIFICANT',
            'Data files present but no data source documented',
            'Data files are present but no information about where '
            'the data came from was found in the README. Validators '
            'cannot verify data provenance, check for updates, or '
            'understand data access restrictions without this '
            'information.',
            [f'Data files: '
             f'{", ".join(f.name for f in data_files[:5])}',
             'Required: data source, URL, DOI, or access instructions']
        ))

    return findings


def detect_AA_figure_reproducibility(repo_dir, all_files):
    """Failure Mode AA: Figures/outputs committed but no figure generation code,
    or committed outputs that validators should verify against a fresh run.

    Recognises:
    - Standard image formats (.png, .jpg, .svg, .eps, .pdf)
    - LaTeX figure/table outputs (.tex) generated by R (TikZ, stargazer, kableExtra)
    - Common output directory names including Out/, output/, results/, tables/
    """
    findings = []

    figure_extensions = {'.png', '.jpg', '.jpeg', '.svg', '.eps', '.pdf',
                         '.tex'}  # .tex covers TikZ figures and stargazer/kableExtra tables
    # 'images' intentionally excluded — it's an asset input dir (stimuli, icons),
    # not an output dir.  Researchers put generated figures in 'figures'/'results'.
    _FIGURE_DIRS = {
        'figures', 'figure', 'figs', 'fig',
        'plots', 'plot',
        'results', 'result',
        'out', 'output', 'outputs',
        'tables', 'table', 'tex',
    }
    figure_files = [
        f for f in all_files
        if f.suffix.lower() in figure_extensions
        and f.parent.name.lower() in _FIGURE_DIRS
        and not any(part.lower() in VENDOR_DIRS
                    for part in f.relative_to(repo_dir).parts)
        and not _in_asset_dir(f, repo_dir)
    ]
    # Fallback: if no named-dir match, look for any non-root directory
    # with ≥10 figure-extension files (catches Out/, Results_v2/, etc.)
    if not figure_files:
        _dir_counts: dict = {}
        for f in all_files:
            if (f.suffix.lower() in figure_extensions
                    and not _in_asset_dir(f, repo_dir)):
                _dir_counts[f.parent] = _dir_counts.get(f.parent, 0) + 1
        for _d, _n in _dir_counts.items():
            if _n >= 10 and _d != repo_dir:
                figure_files += [
                    f for f in all_files
                    if f.parent == _d and f.suffix.lower() in figure_extensions
                    and not _in_asset_dir(f, repo_dir)
                ]

    if not figure_files:
        return findings

    # look for figure/output generation code — Python, R, and LaTeX output
    plot_patterns = re.compile(
        r'(plt\.|ggplot|plot\(|savefig|ggsave|matplotlib'
        r'|seaborn|plotly|bokeh|altair'
        r'|stargazer|kableExtra|tikz|pdf\(|png\(|svg\(|eps\('
        r'|sink\(|write\.csv|write_csv|fwrite)',
        re.IGNORECASE
    )

    has_plot_code = False
    for f in all_files:
        if f.suffix.lower() in CODE_EXTENSIONS:
            content = read_file_safe(f)
            if plot_patterns.search(content):
                has_plot_code = True
                break

    if figure_files and not has_plot_code:
        findings.append(finding(
            'AA', 'SIGNIFICANT',
            f'{len(figure_files)} figure(s) committed but no figure generation code found',
            'Figure files are committed but no code that generates '
            'figures was detected. Validators cannot reproduce the '
            'figures from scratch. If figures are generated by the '
            'analysis scripts, ensure the plotting code is included.',
            [f'Figures: {", ".join(f.name for f in figure_files[:5])}']
        ))
    elif figure_files and has_plot_code:
        findings.append(finding(
            'AA', 'LOW CONFIDENCE',
            f'{len(figure_files)} figure(s) committed — verify generation code produces matching output',
            'Figures are committed and plotting code exists. '
            'Validators should verify that running the code '
            'reproduces figures that match the committed versions '
            'and the published paper.',
            [f'Figures to verify: '
             f'{", ".join(f.name for f in figure_files[:5])}']
        ))

    return findings


def detect_AB_parallel_no_seed(repo_dir, all_files):
    """Failure Mode AB: Parallelisation without determinism controls."""
    findings = []
    _stata_lib_dirs = {'plus', 'personal', 'stbplus'}
    code_files = [f for f in all_files
                  if f.suffix.lower() in CODE_EXTENSIONS
                  and not ('ado' in f.parts
                           and any(p in _stata_lib_dirs for p in f.parts))]

    parallel_patterns = re.compile(
        r'(multiprocessing|concurrent\.futures|joblib|dask'
        r'|ray\.|(?<![Mm]ax)(?<![Aa]vg)(?<![Mm]in)\bPool\b|ProcessPool|ThreadPool'
        r'|n_jobs\s*=|parallel\s*=\s*True'
        r'|mp\.Pool|futures\.ProcessPoolExecutor'
        r'|DataLoader[^)]*num_workers\s*=\s*[1-9]'
        # R parallel libraries
        r'|mclapply|parLapply|parSapply|mcmapply|registerDoParallel|%dopar%'
        r'|library\s*\(\s*["\']?(?:parallel|foreach|doParallel)\b'
        # Java / Scala
        r'|ForkJoinPool|ExecutorService|parallelStream|newFixedThreadPool|newCachedThreadPool'
        # Shell / Make
        r'|make\s+(?:-j|--jobs)|xargs\s+-P|\bparallel\s+.*:::)',
        re.IGNORECASE
    )

    determinism_patterns = re.compile(
        r'(worker_init_fn|pl\.seed_everything'
        r'|torch\.use_deterministic_algorithms'
        r'|PYTHONHASHSEED|random_state\s*=\s*\d'
        r'|initializer\s*='
        # R
        r'|set\.seed\s*\(|clusterSetRNGStream)',
        re.IGNORECASE
    )

    # JS is single-threaded — async/Promise are not non-deterministic parallelism
    _ab_code_suffixes = {f.suffix.lower() for f in code_files}
    _non_web = {'.py', '.r', '.rmd', '.jl', '.do', '.m', '.scala', '.java', '.cpp', '.c'}
    if not (_ab_code_suffixes & _non_web):
        return findings  # JS/web-only repo

    uses_parallel = False
    has_determinism = False
    parallel_langs: set[str] = set()   # extensions that triggered parallel match

    # Include Makefiles (no suffix or .mk) alongside code_files
    _makefile_names = {'makefile', 'gnumakefile', 'bsdmakefile'}
    _makefile_files = [
        f for f in all_files
        if f.name.lower() in _makefile_names or f.suffix.lower() == '.mk'
    ]

    for f in list(code_files) + _makefile_files:
        content = read_file_safe(f)
        if parallel_patterns.search(content):
            uses_parallel = True
            ext = f.suffix.lower() if f.suffix else '.mk'
            parallel_langs.add(ext)
        if determinism_patterns.search(content):
            has_determinism = True

    if uses_parallel and not has_determinism:
        # Build language-specific guidance so the recommendation is actionable
        # for the actual language(s) used — not just Python.
        _guidance = []
        _lang_names = []

        if '.java' in parallel_langs or '.scala' in parallel_langs:
            _lang = 'Java' if '.java' in parallel_langs else 'Scala'
            _lang_names.append(_lang)
            _guidance.append(
                f'{_lang}: use ForkJoinPool with a fixed parallelism level; '
                'ensure tasks produce results in a deterministic order '
                '(e.g. collect into a sorted structure rather than relying '
                'on thread-scheduling order); seed any per-thread RNG with '
                'a fixed value (new Random(42))'
            )
        if '.py' in parallel_langs:
            _lang_names.append('Python')
            _guidance.append(
                'Python: set PYTHONHASHSEED=0 and use worker_init_fn '
                'to seed each worker process'
            )
        if {'.r', '.rmd'} & parallel_langs:
            _lang_names.append('R')
            _guidance.append(
                'R: call set.seed() before parallel blocks; use '
                'parallel::clusterSetRNGStream() for cluster-based parallelism'
            )
        if {'.cpp', '.c'} & parallel_langs:
            _lang_names.append('C/C++')
            _guidance.append(
                'C/C++: seed each thread\'s RNG independently with a '
                'fixed, deterministic value'
            )
        if {'.sh', '.bash', '.mk'} & parallel_langs:
            _lang_names.append('Shell/Make')
            _guidance.append(
                'Shell/Make: document the exact number of parallel jobs used '
                '(e.g. make -j4, not make -j) and ensure all parallel tasks '
                'are independent with no shared output files or race conditions'
            )
        if not _guidance:
            _guidance = [
                'Set a fixed seed for each parallel worker and ensure '
                'task results are collected in a deterministic order'
            ]

        _lang_str = '/'.join(_lang_names) if _lang_names else 'parallel'
        findings.append(finding(
            'AB', 'SIGNIFICANT',
            f'Parallelisation used without determinism controls ({_lang_str})',
            'The code uses parallel processing but no determinism controls '
            'were found. Parallel execution order is non-deterministic by '
            'default — results may vary between runs depending on thread '
            'scheduling.',
            ['Parallel patterns detected without determinism controls',
             *_guidance]
        ))

    return findings


def detect_AC_deprecated_functions(repo_dir, all_files):
    """Failure Mode AC: Use of deprecated functions likely to break."""
    findings = []
    code_files = [f for f in all_files
                  if f.suffix.lower() == '.py'
                  and not _is_minified(f)
                  and not any(part.lower() in VENDOR_DIRS
                              for part in f.relative_to(repo_dir).parts)]

    deprecated = {
        'np.bool': 'np.bool_',
        'np.int': 'np.int_',
        'np.float': 'np.float64',
        'np.complex': 'np.complex128',
        'np.object': 'np.object_',
        'np.str': 'np.str_',
        'sklearn.cross_validation': 'sklearn.model_selection',
        'from sklearn.externals': 'install joblib directly',
        'pd.Panel': 'pd.DataFrame (Panel removed)',
        'DataFrame.ix[': 'DataFrame.loc[ or .iloc[',
        'tensorflow.compat.v1': 'TensorFlow 2.x API',
    }

    for f in code_files:
        content = read_file_safe(f)
        found = []
        for old, new in deprecated.items():
            if old in content:
                found.append(f'{old} → {new}')
        if found:
            findings.append(finding(
                'AC', 'SIGNIFICANT',
                f'Deprecated functions detected in {f.name}',
                'This file uses functions that have been removed or '
                'deprecated in recent package versions. Running this '
                'code with current package versions will likely '
                'produce errors.',
                [f'Deprecated: {d}' for d in found[:5]]
            ))

    return findings


def detect_AD_missing_gitignore(repo_dir, all_files):
    """Failure Mode AD: No .gitignore — sensitive or junk files may be committed."""
    findings = []

    # Only meaningful for git repositories. A static archive has no git history
    # and no need for a .gitignore.
    if not (repo_dir / '.git').is_dir():
        return findings

    has_gitignore = any(
        f.name == '.gitignore' for f in all_files
    )

    has_python = any(f.suffix.lower() == '.py' for f in all_files)

    if has_python and not has_gitignore:
        # check for files that should be ignored
        junk_files = [
            f for f in all_files
            if f.suffix.lower() in {'.pyc', '.pyo'}
            or f.name in {'.DS_Store', 'Thumbs.db', 'desktop.ini'}
            or '__pycache__' in str(f)
        ]

        if junk_files:
            findings.append(finding(
                'AD', 'SIGNIFICANT',
                'No .gitignore and junk files detected in repository',
                'No .gitignore file was found and system or compiled '
                'files are present in the repository. These files '
                'bloat the repository, may contain system-specific '
                'paths, and suggest the repository was not cleaned '
                'before sharing.',
                [f'Junk files found: '
                 f'{", ".join(f.name for f in junk_files[:5])}']
            ))
        else:
            findings.append(finding(
                'AD', 'LOW CONFIDENCE',
                'No .gitignore file found',
                'No .gitignore file was found. Without one, compiled '
                'files, credentials, and system files may be '
                'accidentally committed in future.',
                ['Recommendation: add a .gitignore file']
            ))

    return findings


def detect_AE_mixed_languages(repo_dir, all_files):
    """Failure Mode AE: Multiple languages used without integration docs."""
    findings = []

    language_extensions = {
        'Python': {'.py'},
        'R': {'.r', '.rmd', '.qmd'},
        'Julia': {'.jl'},
        'Stata': {'.do', '.ado'},
        'MATLAB': {'.m', '.mlx'},
        'Shell': {'.sh', '.bash'},
        'SQL': {'.sql'},
    }

    _stata_lib_dirs = {'plus', 'personal', 'stbplus'}
    languages_found = {}
    for lang, exts in language_extensions.items():
        files = [
            f for f in all_files
            if f.suffix.lower() in exts
            and not ('ado' in f.parts and any(p in _stata_lib_dirs for p in f.parts))
        ]
        if files:
            languages_found[lang] = len(files)

    if len(languages_found) >= 3:
        readme_file = None
        for f in all_files:
            if f.name.lower() in {'readme.md', 'readme.txt'}:
                readme_file = f
                break

        integration_documented = False
        if readme_file:
            try:
                content = readme_file.read_text(
                    encoding='utf-8', errors='ignore'
                ).lower()
                if any(lang.lower() in content
                       for lang in languages_found):
                    integration_documented = True
            except Exception:
                pass

        if not integration_documented:
            langs = ', '.join(
                f'{l} ({n} files)'
                for l, n in languages_found.items()
            )
            findings.append(finding(
                'AE', 'SIGNIFICANT',
                f'Multiple languages used without integration documentation',
                'This repository uses multiple programming languages '
                'but the README does not explain how they fit together. '
                'Validators need to know the execution order across '
                'languages and any data handoffs between them.',
                [f'Languages: {langs}',
                 'Required: explain how languages interact in README']
            ))

    return findings


def detect_AF_output_format_undocumented(repo_dir, all_files):
    """Failure Mode AF: Output format not documented."""
    findings = []

    code_files = [f for f in all_files
                  if f.suffix.lower() in CODE_EXTENSIONS]

    write_patterns = re.compile(
        r'(to_csv|to_excel|to_parquet|to_stata|to_latex'
        r'|savefig|to_html|write_csv|fwrite|write\.csv'
        r'|saveRDS|save\.image|np\.save|pickle\.dump'
        r'|\.write\s*\()',
        re.IGNORECASE
    )

    has_write_code = False
    for f in code_files:
        content = read_file_safe(f)
        if write_patterns.search(content):
            has_write_code = True
            break

    if not has_write_code:
        return findings

    readme_file = None
    for f in all_files:
        if f.name.lower() in {'readme.md', 'readme.txt'}:
            readme_file = f
            break

    if readme_file:
        try:
            content = readme_file.read_text(
                encoding='utf-8', errors='ignore'
            ).lower()
            output_documented = any(term in content for term in [
                'output', 'result', 'produces', 'generates',
                'will create', 'will produce', 'expected'
            ])
            if not output_documented:
                findings.append(finding(
                    'AF', 'LOW CONFIDENCE',
                    'Code writes output files but outputs not documented in README',
                    'The code generates output files but the README '
                    'does not describe what outputs to expect. '
                    'Validators cannot verify successful completion '
                    'without knowing what files should be produced.',
                    ['Recommendation: list expected output files '
                     'in README']
                ))
        except Exception:
            pass

    return findings


def detect_AG_api_keys_in_code(repo_dir, all_files):
    """Failure Mode AG: API keys or tokens hardcoded in source files."""
    findings = []
    code_files = [f for f in all_files
                  if f.suffix.lower() in CODE_EXTENSIONS
                  and not _is_minified(f)
                  and not any(part.lower() in VENDOR_DIRS
                              for part in f.relative_to(repo_dir).parts)]

    key_patterns = re.compile(
        r'[A-Z_]*(?:KEY|SECRET|TOKEN|PASSWORD|AUTH|CREDENTIAL|API)[A-Z_]*'
        r'\s*=\s*["\'][a-zA-Z0-9_\-]{16,}["\']',
        re.IGNORECASE
    )

    # Matches public/private/protected static final String constants.
    # Value group captures the string content (8–200 chars).
    _java_const_re = re.compile(
        r'(?:public|private|protected)\s+static\s+final\s+String\s+(\w+)\s*=\s*["\']([^"\']{8,200})["\']'
    )

    def _java_value_is_safe(value: str) -> bool:
        """Return True if the string value looks like a descriptor, not a credential."""
        # No digits at all → human-readable name (e.g. "DefaultAuthenticationKey")
        if not re.search(r'\d', value):
            return True
        # Dotted qualified name WITH at least one dot: letters, digits, dots, hyphens only.
        # Requiring a dot prevents bare alphanumeric strings with digits (e.g. AWS keys)
        # from being falsely suppressed.
        # e.g. "weka.core.DoNotLoadIfEnvVarNotSet", "log4j2.appender.A1"
        if '.' in value and re.fullmatch(r'[A-Za-z][A-Za-z0-9.\-]*', value):
            return True
        return False

    def _java_match_value(m: str):
        """Extract the string literal value from a key_patterns match."""
        vm = re.search(r'["\']([^"\']+)["\']', m)
        return vm.group(1) if vm else ''

    for f in code_files:
        content = read_file_safe(f)
        matches = key_patterns.findall(content)
        # Java-specific: suppress static final String constants whose values
        # are human-readable descriptive strings, not tokens/hashes/UUIDs.
        if matches and f.suffix.lower() == '.java':
            # Pass 1 — declaration-based: scan every line for a static final
            # String declaration; if the value looks safe, note the var name.
            safe_vars: set = set()
            for line in content.splitlines():
                cm = _java_const_re.search(line)
                if cm and _java_value_is_safe(cm.group(2)):
                    safe_vars.add(cm.group(1))
            # Pass 2 — value-direct: even if the declaration wasn't caught
            # (multi-line decl, inherited constant, non-static-final local),
            # suppress matches whose own value looks safe.  This closes the
            # gap between declaration-scanning and actual match sites.
            matches = [
                m for m in matches
                if m.split('=')[0].strip().split('\n')[-1].strip() not in safe_vars
                and not _java_value_is_safe(_java_match_value(m))
            ]
        if matches:
            # Extract variable names from matches for evidence
            var_names = []
            for m in matches:
                var = m.split('=')[0].strip().split('\n')[-1].strip()
                if var and var not in var_names:
                    var_names.append(var)
            if not var_names:
                continue  # parsing yielded no variable names — suppress vacuous finding
            evidence_lines = [f'Hardcoded credential: {v}' for v in var_names[:5]]
            evidence_lines.append('Action required: rotate these credentials immediately if real')
            findings.append(finding(
                'AG', 'CRITICAL',
                f'Possible hardcoded credentials in {f.name}: {", ".join(var_names[:3])}',
                'What appears to be an API key or token is hardcoded '
                'in source code. If real, this is a security issue — '
                'credentials committed to a repository should be '
                'considered compromised. Replace with environment '
                'variables immediately.',
                evidence_lines
            ))

    return findings


def detect_AH_no_changelog(repo_dir, all_files):
    """Failure Mode AH: No changelog or version history."""
    findings = []

    changelog_names = {
        'changelog', 'changelog.md', 'changelog.txt',
        'changes', 'changes.md', 'history.md',
        'news.md', 'releases.md'
    }

    has_changelog = any(
        f.name.lower() in changelog_names
        for f in all_files
    )

    has_readme = any(
        f.name.lower() in {'readme.md', 'readme.txt'}
        for f in all_files
    )

    # only flag if there's a substantial codebase
    py_files = [f for f in all_files if f.suffix.lower() == '.py']

    if len(py_files) > 5 and not has_changelog and has_readme:
        findings.append(finding(
            'AH', 'LOW CONFIDENCE',
            'No changelog or version history found',
            'No changelog file was found. For research code, a '
            'changelog helps validators understand what changed '
            'between versions and whether the committed code matches '
            'the version used to generate the published results.',
            ['Recommendation: add CHANGELOG.md noting the version '
             'used for publication']
        ))

    return findings


def detect_AI_print_debugging(repo_dir, all_files):
    """Failure Mode AI: Excessive print debugging suggests unfinished code."""
    findings = []
    code_files = [f for f in all_files
                  if f.suffix.lower() in {'.py', '.ipynb'}
                  and not _is_minified(f)
                  and not any(part.lower() in VENDOR_DIRS
                              for part in f.relative_to(repo_dir).parts)]

    for f in code_files:
        content = read_file_safe(f)
        lines = content.splitlines()

        print_count = sum(
            1 for line in lines
            if re.search(r'^\s*print\s*\(', line)
            and not re.search(r'#.*print', line)
        )

        total_lines = len([l for l in lines if l.strip()])

        if total_lines > 0 and print_count / total_lines > 0.1:
            findings.append(finding(
                'AI', 'LOW CONFIDENCE',
                f'High density of print statements in {f.name}',
                f'{print_count} print statements in {total_lines} '
                f'lines of code suggests debugging output was not '
                f'cleaned up before publication. This does not affect '
                f'reproducibility but suggests the code may not be '
                f'in its final form.',
                [f'Evidence: {print_count} prints in {total_lines} '
                 f'non-blank lines ({print_count*100//total_lines}%)']
            ))

    return findings


def detect_AJ_hardcoded_sample_size(repo_dir, all_files):
    """Failure Mode AJ: Sample sizes or thresholds hardcoded without explanation."""
    findings = []
    code_files = [f for f in all_files
                  if f.suffix.lower() in {'.py', '.r', '.rmd'}
                  and not _is_minified(f)
                  and not any(part.lower() in VENDOR_DIRS
                              for part in f.relative_to(repo_dir).parts)]

    magic_number_pattern = re.compile(
        r'(?:head|sample|nrow|iloc|[:]\s*)\s*\(\s*(\d{3,})\s*\)'
        r'|n\s*=\s*(\d{3,})\b'
        r'|threshold\s*=\s*(0\.\d+)'
        r'|cutoff\s*=\s*(0\.\d+)',
        re.IGNORECASE
    )

    for f in code_files:
        content = read_file_safe(f)
        matches = magic_number_pattern.findall(content)
        flat = [m for group in matches for m in group if m]

        if len(flat) >= 3:
            findings.append(finding(
                'AJ', 'LOW CONFIDENCE',
                f'Multiple hardcoded numerical thresholds in {f.name}',
                'Several hardcoded numbers that appear to be sample '
                'sizes, thresholds, or cutoffs were found without '
                'explanatory comments. Validators cannot determine '
                'if these match the values described in the paper '
                'without documentation.',
                [f'Values found: {", ".join(sorted(set(flat))[:8])}',
                 'Recommendation: add comments explaining each value']
            ))

    return findings


def _strip_comment_lines(content: str, suffix: str) -> str:
    """Replace comment-only lines with blank lines before URL scanning.

    Prevents URLs in boilerplate template comments (e.g. the Shiny
    '# https://shiny.posit.co/' header) from triggering [AK].
    Only whole-line comments are stripped; inline trailing comments are kept.
    """
    ext = suffix.lower()
    if ext in {'.r', '.rmd', '.qmd', '.py', '.sh', '.jl',
               '.yaml', '.yml', '.rb', '.pl', '.perl'}:
        pfx = '#'
    elif ext == '.m':
        pfx = '%'
    else:
        return content
    out = []
    for line in content.splitlines():
        out.append('' if line.lstrip().startswith(pfx) else line)
    return '\n'.join(out)


def detect_AK_external_urls(repo_dir, all_files):
    """Failure Mode AK: External URLs that may become unavailable."""
    findings = []
    code_files = [f for f in all_files
                  if f.suffix.lower() in CODE_EXTENSIONS | {'.md', '.txt'}
                  and not _is_minified(f)
                  and not any(part.lower() in VENDOR_DIRS
                              for part in f.relative_to(repo_dir).parts)]

    url_pattern = re.compile(
        r'https?://(?!github\.com|zenodo\.org|doi\.org|arxiv\.org'
        r'|pypi\.org|anaconda\.org|conda\.io'
        # Author identity / social-profile domains — stable, not data links
        r'|linkedin\.com|twitter\.com|(?:www\.)?x\.com'
        r'|orcid\.org|researchgate\.net|academia\.edu'
        # Licence reference domains — stable, not external data dependencies
        r'|creativecommons\.org|gnu\.org|apache\.org|opensource\.org|spdx\.org'
        r')[^\s\'")\]>]+',
        re.IGNORECASE
    )

    # Collect URLs; deduplicate by domain for deterministic, compact output.
    # Sort files by name so the same domain always maps to the same
    # representative URL regardless of rglob traversal order.
    _domain_pat = re.compile(r'https?://(?:www\.)?([^/\s\'")\]>]+)')
    # Valid domain: starts with alphanumeric, has at least one dot-separated
    # label of 2+ letters.  Rejects artefacts like '....' or bare hostnames.
    _VALID_DOMAIN_RE = re.compile(
        r'^[a-zA-Z0-9]([a-zA-Z0-9\-]{0,61}[a-zA-Z0-9])?'
        r'(\.[a-zA-Z]{2,})+$'
    )
    _domain_to_url: dict = {}   # domain -> first URL seen
    for f in sorted(code_files, key=lambda x: x.name):
        content = _strip_comment_lines(read_file_safe(f), f.suffix)
        for url in url_pattern.findall(content):
            m = _domain_pat.match(url)
            if m:
                domain = m.group(1).lower().strip()
                if domain and domain not in _domain_to_url and _VALID_DOMAIN_RE.match(domain):
                    _domain_to_url[domain] = url

    urls_found = set(_domain_to_url.values())

    # Guard: > 20 unique domains almost always means a minified file was scanned
    # despite the exclusion filters — suppress entirely with a LOW CONFIDENCE note.
    if len(_domain_to_url) > 20:
        findings.append(finding(
            'AK', 'LOW CONFIDENCE',
            'URL detection skipped — unusually high domain count',
            f'{len(_domain_to_url)} unique domains were detected after vendor/minified '
            'exclusions. This strongly suggests that a bundled or generated file was '
            'scanned. Manual review recommended — check for framework JS not in a '
            'standard vendor directory.',
            [f'Domain count: {len(_domain_to_url)}',
             'Likely cause: minified JavaScript or generated frontend code not '
             'in a recognised vendor directory']
        ))
        return findings

    # Colab escalation: if Colab links are the primary analysis and there is no
    # local code, the repository is irreproducible if those links go dead.
    _colab_urls = sorted(
        u for u in urls_found if 'colab.research.google.com' in u.lower()
    )
    _has_local_code = any(
        f.suffix.lower() in CODE_EXTENSIONS | {'.ipynb'} for f in all_files
    )
    if _colab_urls and not _has_local_code:
        findings.append(finding(
            'AK', 'SIGNIFICANT',
            'Primary analysis hosted on Google Colab — no local code found',
            'The repository links to Google Colab notebook(s) but contains no '
            'local code files. If these links go offline or the notebooks are '
            'made private, the analysis cannot be reproduced by definition. '
            'Download the notebooks from Colab and commit them to the repository.',
            [f'Colab link: {u}' for u in _colab_urls[:3]] +
            ['Fix: in Colab — File > Download > Download .ipynb — '
             'then commit the .ipynb file alongside your data']
        ))
        urls_found -= set(_colab_urls)   # don't double-report at LOW CONFIDENCE

    # DataCamp escalation: DataCamp workspaces can be private or expire, making
    # them even more fragile than Colab links.
    _datacamp_urls = sorted(
        u for u in urls_found if 'datacamp.com/workspace' in u.lower()
    )
    if _datacamp_urls and not _has_local_code:
        findings.append(finding(
            'AK', 'SIGNIFICANT',
            'Primary analysis hosted on DataCamp workspace — no local code found',
            'The repository links to a DataCamp workspace but contains no local '
            'code files. DataCamp workspaces can be set to private or deleted, '
            'making them more fragile than other cloud notebooks. If the link '
            'goes offline the analysis cannot be reproduced. Export the notebook '
            'and commit it to the repository.',
            [f'DataCamp link: {u}' for u in _datacamp_urls[:3]] +
            ['Fix: in DataCamp workspace — File > Export as Notebook (.ipynb) — '
             'then commit the .ipynb file to the repository']
        ))
        urls_found -= set(_datacamp_urls)   # don't double-report at LOW CONFIDENCE

    if urls_found:
        # Show up to 10 unique domains, sorted, with "and N more" if needed.
        MAX_DOMAINS = 10
        _reported_urls = set(_colab_urls) | set(_datacamp_urls)
        _remaining_domains = sorted(
            d for d, u in _domain_to_url.items()
            if u not in _reported_urls and d.strip()
        )
        if not _remaining_domains:
            return findings
        shown = _remaining_domains[:MAX_DOMAINS]
        rest = len(_remaining_domains) - len(shown)
        evidence_str = ', '.join(shown)
        if rest > 0:
            evidence_str += f' — and {rest} more domain{"s" if rest != 1 else ""}'
        findings.append(finding(
            'AK', 'LOW CONFIDENCE',
            f'External URLs detected — may become unavailable',
            'The code or documentation references external URLs. '
            'If these URLs go offline, validators will be unable '
            'to access required resources. Use DOIs or archived '
            'URLs where possible.',
            [f'Domains: {evidence_str}']
        ))

    return findings


def detect_AL_data_privacy(repo_dir, all_files):
    """Failure Mode AL: Potential personal or sensitive data indicators."""
    findings = []

    sensitive_patterns = re.compile(
        r'\b(ssn|social.security|date.of.birth|dob'
        r'|phone.number|email.address|home.address'
        r'|medical.record|patient.id|participant.id'
        r'|subject.id|ip.address|passport'
        r'|national.insurance|nin\b|nhs.number)\b',
        re.IGNORECASE
    )

    data_files = [
        f for f in all_files
        if f.suffix.lower() in {'.csv', '.tsv', '.xlsx', '.xls'}
    ]

    # Directory/filename patterns that indicate model outputs, not raw personal data
    model_output_indicators = {
        'coef', 'coefs', 'coefficient', 'coefficients',
        'prediction', 'predictions', 'output', 'outputs',
        'result', 'results', 'fitted', 'estimates'
    }

    flagged = []
    for f in data_files:
        # Skip files that are clearly model outputs not raw personal data
        name_lower = f.stem.lower()
        parent_lower = f.parent.name.lower()
        if any(ind in name_lower for ind in model_output_indicators):
            continue
        if any(ind in parent_lower for ind in model_output_indicators):
            continue
        content = read_file_safe(f)
        if sensitive_patterns.search(content[:2000]):
            flagged.append(f.name)

    if flagged:
        findings.append(finding(
            'AL', 'SIGNIFICANT',
            f'Potential sensitive data indicators in: '
            f'{", ".join(flagged[:3])}',
            'Data files contain column names or values that suggest '
            'personally identifiable or sensitive information. '
            'Verify that data sharing complies with IRB approval, '
            'GDPR, and journal data sharing policies before '
            'publishing this repository.',
            [f'Files with sensitive indicators: {", ".join(flagged)}',
             'Required: data anonymisation or access restriction '
             'documentation']
        ))

    return findings


def detect_AM_makefile_missing(repo_dir, all_files):
    """Failure Mode AM: Complex pipeline with no automation."""
    findings = []

    pipeline_indicators = [
        f for f in all_files
        if f.suffix.lower() == '.py'
        and re.match(r'^\d+_', f.name)
    ]

    has_automation = any(
        f.name.lower() in {
            'makefile', 'dodo.py', 'snakefile',
            'workflow.py', 'pipeline.py', 'run_all.py',
            'run_all.sh', 'main.py', 'reproduce.py',
            'reproduce.sh'
        }
        for f in all_files
    )

    if len(pipeline_indicators) >= 4 and not has_automation:
        findings.append(finding(
            'AM', 'SIGNIFICANT',
            f'{len(pipeline_indicators)} numbered scripts with no pipeline automation',
            'The repository has multiple numbered scripts suggesting '
            'a sequential pipeline, but no automation file '
            '(Makefile, Snakefile, run_all.py) was found. '
            'Validators must manually execute each script in order. '
            'A single entry point that runs the full pipeline '
            'significantly improves reproducibility.',
            [f'Scripts: '
             f'{", ".join(f.name for f in pipeline_indicators[:6])}',
             'Recommendation: add run_all.py or Makefile']
        ))

    return findings


def detect_AN_commented_code(repo_dir, all_files):
    """Failure Mode AN: Large blocks of commented-out code."""
    findings = []
    code_files = [f for f in all_files
                  if f.suffix.lower() == '.py']

    for f in code_files:
        content = read_file_safe(f)
        lines = content.splitlines()

        commented = sum(
            1 for line in lines
            if line.strip().startswith('#')
            and len(line.strip()) > 5
            and not line.strip().startswith('#!/')
        )
        total = len([l for l in lines if l.strip()])

        if total > 20 and commented / total > 0.25:
            findings.append(finding(
                'AN', 'LOW CONFIDENCE',
                f'High proportion of commented code in {f.name}',
                f'{commented} of {total} non-blank lines are comments '
                f'({commented*100//total}%). Large blocks of commented '
                f'code suggest earlier versions of the analysis may '
                f'be present. This is not a reproducibility error but '
                f'may indicate the committed code is not the final '
                f'version.',
                [f'Evidence: {commented*100//total}% commented lines '
                 f'in {f.name}']
            ))

    return findings


_SESSION_INFO_PAT = re.compile(
    r'sessionInfo\s*\('                     # base R: sessionInfo()
    r'|sessioninfo::session_info\s*\('      # sessioninfo pkg (lowercase)
    r'|devtools::session_info\s*\('         # devtools wrapper
)


def detect_AO_r_specific_issues(repo_dir, all_files):
    findings = []
    # Iterate all_files directly — same pattern as [E] and [BA].
    # Every detector that needs recursive traversal should use all_files,
    # not its own rglob: all_files is built from repo_dir.rglob('*') in
    # both entry points and is the single shared recursive file scanner.
    r_files = [f for f in all_files if f.suffix.lower() in {'.r', '.rmd', '.qmd'}]
    if not r_files:
        return findings
    has_renv = any(f.name.lower() in {'renv.lock', 'packrat.lock'} for f in all_files)
    session_info_files = {'session_info.txt', 'session_info.log', 'sessioninfo.txt',
                          'r_session_info.txt', 'session-info.txt'}
    has_session_info = (
        any(_SESSION_INFO_PAT.search(read_file_safe(f)) for f in r_files) or
        any(f.name.lower() in session_info_files for f in all_files)
    )
    if not has_renv:
        findings.append(finding('AO', 'SIGNIFICANT',
            'R code present but no renv.lock found',
            'Without renv.lock validators cannot install exact package versions.',
            ['Missing: renv.lock', 'Run renv::init() and renv::snapshot()']))
    if not has_session_info:
        _scanned = ', '.join(f.name for f in r_files[:5])
        if len(r_files) > 5:
            _scanned += f' (+ {len(r_files) - 5} more)'
        findings.append(finding('BN', 'LOW CONFIDENCE',
            'No sessionInfo() call found in R scripts',
            'sessionInfo() documents exact R and package versions used.',
            [f'Scanned {len(r_files)} R file(s): {_scanned}',
             'Recommendation: add sessionInfo() at end of main script']))
    return findings

def detect_AP_stata_specific(repo_dir, all_files):
    findings = []
    stata_files = [f for f in all_files if f.suffix.lower() in {'.do', '.ado'}]
    if not stata_files:
        return findings
    has_version = any(
        re.search(r'version\s+\d+', read_file_safe(f), re.MULTILINE)
        for f in stata_files
    )
    if not has_version:
        findings.append(finding('AP', 'SIGNIFICANT',
            'Stata do-files missing version declaration',
            'Without version declaration Stata behaviour differs between versions.',
            ['Missing: version XX at top of do-files']))
    return findings

def detect_AQ_large_model_files(repo_dir, all_files):
    return []

def detect_AR_encoding_issues(repo_dir, all_files):
    findings = []
    py_files = [f for f in all_files if f.suffix.lower() in {'.py', '.ipynb'}]
    bad = []
    for f in py_files:
        content = read_file_safe(f)
        if re.search(r'open\s*\(', content) and 'encoding=' not in content:
            bad.append(f.name)
    if len(bad) >= 2:
        bad_files = ','.join(bad[:5])
        findings.append(finding('AR', 'LOW CONFIDENCE',
            f'open() without encoding in {len(bad)} files',
            'open() without encoding behaves differently on Windows vs Linux/Mac.',
            [f'Files: {bad_files}', 'Fix: add encoding="utf-8"']))
    return findings


def detect_AS_network_calls(repo_dir, all_files):
    findings = []
    code_files = [f for f in all_files
                  if f.suffix.lower() in CODE_EXTENSIONS
                  and not _is_minified(f)
                  and not any(part.lower() in VENDOR_DIRS
                              for part in f.relative_to(repo_dir).parts)]
    # urllib.parse is string manipulation, not a network call — exclude it.
    # Only match actual I/O operations: HTTP clients and socket connections.
    net_pattern = re.compile(
        r'(requests\.'
        r'|urllib\.request'
        r'|urllib\.urlopen'
        r'|http\.client'
        r'|socket\.connect'
        r'|wget\.'
        r'|httpx\.'
        r'|aiohttp\.)',
        re.IGNORECASE
    )
    files_with_network = []
    for f in code_files:
        content = read_file_safe(f)
        if net_pattern.search(content):
            files_with_network.append(f.name)
    if files_with_network:
        findings.append(finding('AS', 'SIGNIFICANT',
            f'Network calls detected in {len(files_with_network)} file(s)',
            'Code makes network requests at runtime. These will fail without internet access or if remote resources move. Validators in restricted environments cannot reproduce results.',
            [f'Files: {", ".join(files_with_network[:5])}',
             'Recommendation: document all external dependencies and provide offline fallback']))
    return findings


def detect_AT_database_dependency(repo_dir, all_files):
    findings = []
    code_files = [f for f in all_files
                  if f.suffix.lower() in CODE_EXTENSIONS
                  and not any(part.lower() in VENDOR_DIRS
                              for part in f.relative_to(repo_dir).parts)]
    db_pattern = re.compile(r'(psycopg2|pymysql|sqlalchemy|sqlite3\.connect|pymongo|cx_Oracle|pyodbc|ibm_db|snowflake\.connector)', re.IGNORECASE)
    db_files = []
    for f in code_files:
        content = read_file_safe(f)
        if db_pattern.search(content):
            db_files.append(f.name)
    if db_files:
        findings.append(finding('AT', 'SIGNIFICANT',
            f'Database connections detected in {len(db_files)} file(s)',
            'Code connects to external databases. Validators cannot reproduce results without access to these databases. Document connection requirements and provide sample data or database dumps.',
            [f'Files with DB connections: {", ".join(db_files[:5])}',
             'Required: connection documentation or sample data export']))
    return findings


def detect_AU_cloud_storage(repo_dir, all_files):
    findings = []
    code_files = [f for f in all_files
                  if f.suffix.lower() in CODE_EXTENSIONS
                  and not _is_minified(f)
                  and not any(part.lower() in VENDOR_DIRS
                              for part in f.relative_to(repo_dir).parts)]
    cloud_pattern = re.compile(r'(boto3|s3fs|gcsfs|azure\.storage|google\.cloud\.storage|gs://|s3://|azure://)', re.IGNORECASE)
    cloud_files = []
    for f in code_files:
        content = read_file_safe(f)
        if cloud_pattern.search(content):
            cloud_files.append(f.name)
    if cloud_files:
        findings.append(finding('AU', 'SIGNIFICANT',
            f'Cloud storage access detected in {len(cloud_files)} file(s)',
            'Code reads from or writes to cloud storage (S3, GCS, Azure). Validators require cloud credentials and access permissions to reproduce results.',
            [f'Files: {", ".join(cloud_files[:5])}',
             'Required: document storage buckets, access method, and credentials process']))
    return findings


_DRYAD_FILENAME_RE = re.compile(
    r'doi_(\d+)_(\d+)_([^_].+?)(?:__v\d+)?(?:\.zip)?$', re.IGNORECASE
)
_DOI_RE = re.compile(r'\b10\.\d{4,9}/\S+', re.IGNORECASE)


def _extract_doi_from_filename(zip_name: str):
    """Return DOI string extracted from a Dryad-style filename, or None.

    Dryad format: doi_10_5061_dryad_88r38__v20171123.zip
    Reconstructs: 10.5061/dryad.88r38
    """
    m = _DRYAD_FILENAME_RE.search(zip_name)
    if m:
        prefix = m.group(1)       # 10
        registrar = m.group(2)    # 5061
        suffix = m.group(3)       # dryad_88r38
        suffix_clean = suffix.replace('_', '.')  # dryad.88r38
        return f'{prefix}.{registrar}/{suffix_clean}'
    return None


# Regexes for context-aware deposit-DOI detection used by _has_deposit_doi().
_AW_REF_HEADING_RE = re.compile(
    r'^(?:#+\s*)?(?:references?|bibliography|works?\s+cited|further\s+reading)\s*:?\s*$',
    re.IGNORECASE,
)
# A numbered citation line: [1], 1., 1) or (1) at the start.
_AW_CITE_LINE_RE = re.compile(r'^\s*(?:\[\d+\]|\d+[.)]\s+|\(\d+\)\s+)')


def _has_deposit_doi(content: str) -> bool:
    """Return True if *content* contains a DOI that identifies the deposit itself.

    Ignores DOIs that appear inside bibliography / reference sections or on
    numbered-citation lines — those identify cited papers, not the deposit.

    Suppression signals accepted:
    - Zenodo badge URL            (zenodo.org/badge/)
    - Zenodo-prefixed DOI         (10.5281/zenodo.<digits>)
    - Zenodo DOI redirect         (zenodo.org/doi/)
    - Any DOI / doi.org link that is NOT inside a reference section and NOT
      on a numbered-citation line.
    """
    lower = content.lower()

    # Fast-path: unambiguous deposit-level Zenodo patterns.
    if ('zenodo.org/badge/' in lower
            or 'zenodo.org/doi/' in lower
            or re.search(r'10\.5281/zenodo\.\d', lower)):
        return True

    lines = content.splitlines()
    in_ref_section = False

    for i, line in enumerate(lines):
        stripped = line.strip()
        stripped_lower = stripped.lower()

        # Enter reference section on a matching heading.
        if _AW_REF_HEADING_RE.match(stripped):
            in_ref_section = True
            continue

        # Any other markdown heading resets reference-section state.
        if re.match(r'^#+\s+\S', stripped) and not _AW_REF_HEADING_RE.match(stripped):
            in_ref_section = False

        # Does this line mention a DOI or Zenodo at all?
        has_doi_here = (
            'doi.org' in stripped_lower
            or 'zenodo' in stripped_lower
            or bool(_DOI_RE.search(line))
        )
        if not has_doi_here:
            continue

        # Skip DOIs inside a reference section.
        if in_ref_section:
            continue

        # Skip numbered citation lines outside a formal reference section
        # (e.g. inline "[1] Smith et al. https://doi.org/...").
        if _AW_CITE_LINE_RE.match(stripped):
            continue

        # Check the preceding few lines for informal reference markers
        # (e.g. a "References:" label without a heading character).
        ctx_before = '\n'.join(lines[max(0, i - 4):i]).lower()
        if re.search(r'\breferences?\s*:|\bbibliograph|\bworks?\s+cited\b', ctx_before):
            continue

        # DOI appears in a non-reference context — treat as deposit identifier.
        return True

    return False


def detect_AW_missing_doi(repo_dir, all_files, zip_name=None):
    findings = []
    # CITATION.cff is itself a persistent identifier mechanism — suppress [AW]
    # when one is present (same logic as [BM]).
    if any(f.name.lower() == 'citation.cff' for f in all_files):
        return findings
    # DOI encoded in the deposit filename (Dryad and similar archives).
    if zip_name:
        doi_from_filename = _extract_doi_from_filename(zip_name)
        if doi_from_filename:
            return findings
    # Scan text files for any DOI pattern (doi:, doi.org, bare 10.XXXX/...).
    # Exclude ValiChord-generated output files — they contain "zenodo", "doi:",
    # and real DOI strings injected by the tool itself, which would cause false
    # suppression on a re-run of a ValiChord output zip.
    _GENERATED_NAMES = {
        'ASSESSMENT.md', 'CLEANING_REPORT.md',
    }
    text_files = [
        f for f in all_files
        if f.suffix.lower() in {'.md', '.txt', '.rst'}
        and f.name not in _GENERATED_NAMES
        and not (f.name.endswith('_DRAFT.md') or f.name.endswith('_DRAFT.txt'))
    ]
    has_doi = False
    for f in text_files:
        if _has_deposit_doi(read_file_safe(f)):
            has_doi = True
            break
    if not has_doi:
        # Platform-aware recommendation based on zip filename.
        _zn = (zip_name or '').lower()
        if re.search(r'dataverse_files', _zn, re.IGNORECASE):
            _doi_rec = (
                'This appears to be a Harvard Dataverse deposit — Dataverse '
                'assigns a DOI automatically (format: 10.7910/DVN/XXXXXX). '
                'Add your existing Dataverse DOI to the README rather than '
                'creating a new one elsewhere.'
            )
        elif re.search(r'zenodo', _zn):
            _doi_rec = (
                'This appears to be a Zenodo deposit — Zenodo assigns a DOI '
                'automatically (format: 10.5281/zenodo.XXXXXXX). '
                'Add your Zenodo DOI to the README.'
            )
        elif re.search(r'figshare', _zn):
            _doi_rec = (
                'This appears to be a Figshare deposit — Figshare assigns a '
                'DOI automatically. Add your Figshare DOI to the README.'
            )
        elif re.search(r'dryad|datadryad', _zn):
            _doi_rec = (
                'This appears to be a Dryad deposit — Dryad assigns a DOI '
                'automatically (format: 10.5061/dryad.XXXXXXX). '
                'Add your Dryad DOI to the README.'
            )
        elif re.search(r'osfstorage|(?<![a-z])osf(?![a-z])', _zn):
            _doi_rec = (
                'This appears to be an OSF deposit. OSF projects have a persistent URL '
                'in the format osf.io/XXXXX — find yours on the OSF project page and '
                'add it to your README. You can also register a DOI for your OSF '
                'project via the OSF admin panel (under "Identifiers").'
            )
        elif re.search(r'mendeley', _zn):
            _doi_rec = (
                'This appears to be a Mendeley Data deposit — Mendeley Data '
                'assigns a DOI automatically. Add your Mendeley Data DOI to '
                'the README.'
            )
        else:
            _doi_rec = 'Recommendation: deposit on Zenodo to get a DOI'
        findings.append(finding('AW', 'LOW CONFIDENCE',
            'No DOI or persistent identifier found in documentation',
            'No DOI, Zenodo link, or other persistent identifier was found. A DOI ensures the repository remains citable and accessible long-term.',
            [_doi_rec,
             'Add DOI badge to README']))
    return findings


def detect_AX_container_not_tested(repo_dir, all_files):
    findings = []
    has_dockerfile = any(f.name.lower() == 'dockerfile' for f in all_files)
    if not has_dockerfile:
        return findings
    dockerfile = next(f for f in all_files if f.name.lower() == 'dockerfile')
    content = read_file_safe(dockerfile)
    issues = []
    if 'COPY . .' in content or 'COPY ./' in content:
        if 'WORKDIR' not in content:
            issues.append('No WORKDIR set before COPY')
    if 'latest' in content.lower():
        import re as _reax
        _from_m = _reax.search(r'^FROM\s+(\S+)', content, _reax.IGNORECASE | _reax.MULTILINE)
        _image = _from_m.group(1) if _from_m else 'python:latest'
        _base = _image.split(':')[0]
        _hub_url = f'https://hub.docker.com/_/{_base.split("/")[-1]}' if '/' not in _base else f'https://hub.docker.com/r/{_base}'
        _suggested = _base + ':3.11-slim' if 'python' in _base.lower() else _base + ':<version>'
        issues.append(
            f'Base image uses :latest tag ({_image}) — resolves to a different version over time.\n'
            f'Fix: replace with a pinned version, e.g. FROM {_suggested}\n'
            f'See available tags: {_hub_url}\n'
            f'Note: validators pulling this image in future may get different results.'
        )
    if 'RUN pip install' in content and 'requirements' not in content.lower():
        issues.append('pip install without requirements file — not reproducible')
    if issues:
        findings.append(finding('AX', 'SIGNIFICANT',
            'Dockerfile has reproducibility issues',
            'The Dockerfile contains patterns that may cause different builds on different runs.',
            issues))
    return findings


def detect_AY_workflow_file(repo_dir, all_files):
    findings = []
    has_python = any(f.suffix.lower() == '.py' for f in all_files)
    if not has_python:
        return findings
    ci_files = [f for f in all_files if f.suffix.lower() in {'.yml', '.yaml'}
                and any(ci in str(f).lower() for ci in ['github', 'gitlab', 'circle', 'travis', 'actions'])]
    if ci_files:
        # deduplicate by filename — same file at different depths counts once
        _seen = set()
        _unique_ci = []
        for _f in ci_files:
            if _f.name not in _seen:
                _seen.add(_f.name)
                _unique_ci.append(_f)
        findings.append(finding('AY', 'LOW CONFIDENCE',
            f'CI/CD workflow file(s) found — verify they test reproducibility',
            'Continuous integration workflows are present. Ensure they test that the full analysis pipeline runs successfully, not just code style checks.',
            [f'Workflow files: {", ".join(f.name for f in _unique_ci[:5])}']))
    return findings


def detect_AZ_figure_format(repo_dir, all_files):
    findings = []
    code_files = [f for f in all_files if f.suffix.lower() == '.py']
    bitmap_save = re.compile(r'savefig\s*\([^)]*\.(png|jpg|jpeg)[^)]*\)', re.IGNORECASE)
    vector_save = re.compile(r'savefig\s*\([^)]*\.(svg|eps|pdf)[^)]*\)', re.IGNORECASE)
    saves_bitmap = False
    saves_vector = False
    for f in code_files:
        content = read_file_safe(f)
        if bitmap_save.search(content):
            saves_bitmap = True
        if vector_save.search(content):
            saves_vector = True
    if saves_bitmap and not saves_vector:
        findings.append(finding('AZ', 'LOW CONFIDENCE',
            'Figures saved as bitmap only (PNG/JPG) — consider vector format',
            'Figures are saved as bitmap images. Vector formats (SVG, EPS, PDF) scale without quality loss and are preferred by journals. Bitmap figures may appear different at different resolutions.',
            ['Recommendation: save figures as SVG or PDF in addition to PNG']))
    return findings


def detect_BA_missing_checksums(repo_dir, all_files):
    findings = []
    # Use DATA_EXTENSIONS as the base — this is the same set used by the LOG's
    # data_file_count, keeping the two counts aligned.
    # Added: .tif/.tiff (microscopy/GIS raster data) and ML model formats.
    # Excluded: .txt (too many false positives — captions, requirements, etc.),
    # archive formats (archives are identified separately), and common figure
    # formats (.png, .jpg, etc.) which are supplementary figures, not data.
    _CHECKSUM_WORTHY = DATA_EXTENSIONS | {
        '.tiff', '.tif',
        '.pt', '.pth', '.onnx', '.pb', '.bin', '.safetensors', '.ckpt',
    }
    data_files = [f for f in all_files
                  if f.suffix.lower() in _CHECKSUM_WORTHY
                  and not f.name.lower().startswith('readme')
                  and f.name.lower() not in CODEBOOK_FILENAMES
                  and not _in_asset_dir(f, repo_dir)]
    if not data_files:
        return findings

    # A dedicated checksum file (md5sums.txt, SHA256SUMS, checksums.md, etc.)
    has_checksums = any(
        'checksum' in f.name.lower() or 'md5' in f.name.lower()
        or 'sha256' in f.name.lower() or 'sha1' in f.name.lower()
        or 'sha512' in f.name.lower()
        for f in all_files
    )

    # README must *explicitly document* checksums — not merely mention a hash
    # algorithm in passing (e.g. "we used SHA-256 for signing key material").
    # Accept:
    #   • A section heading: ## Checksums / ## File Integrity / ## MD5 / etc.
    #   • An inline filename+hexhash entry: body_temps.csv: d41d8cd9...
    #   • The word "checksum" anywhere (unambiguous — not an algorithm name)
    _CHECKSUM_HEADING_RE = re.compile(
        r'(?:^|\n)#{1,4}\s*(?:checksum|file\s+integrity|md5\b|sha-?256|sha-?1\b|sha-?512)',
        re.IGNORECASE | re.MULTILINE,
    )
    _CHECKSUM_INLINE_RE = re.compile(
        r'[\w./\-]+\.\w{2,6}\s*[:\|]\s*[0-9a-f]{32,64}\b',
        re.IGNORECASE,
    )
    readme_has_checksums = False
    for f in all_files:
        if f.name.lower() in {'readme.md', 'readme.txt', 'readme.rst'}:
            content = read_file_safe(f)
            if ('checksum' in content.lower()
                    or _CHECKSUM_HEADING_RE.search(content)
                    or _CHECKSUM_INLINE_RE.search(content)):
                readme_has_checksums = True
                break

    if not has_checksums and not readme_has_checksums:
        _scope_note = (
            ' For large deposits, checksums for every file is impractical — '
            'prioritise the primary input data files.'
            if len(data_files) > 20 else ''
        )
        findings.append(finding('BA', 'LOW CONFIDENCE',
            f'{len(data_files)} data files with no checksums documented',
            'No file checksums were found. Checksums allow validators to verify '
            'they have identical copies of the data files, ruling out download '
            'corruption as a source of discrepancy.' + _scope_note,
            ['Recommendation: add MD5 or SHA256 checksums to README for key data files']))
    return findings


def detect_BB_script_permissions(repo_dir, all_files):
    findings = []
    shell_files = [f for f in all_files
                   if f.suffix.lower() in {'.sh', '.bash'}
                   and not any(part.lower() in VENDOR_DIRS
                               for part in f.relative_to(repo_dir).parts)]
    if not shell_files:
        return findings
    import stat as _stat
    non_executable = []
    for f in shell_files:
        try:
            mode = f.stat().st_mode
            if not (mode & _stat.S_IXUSR):
                non_executable.append(f.name)
        except Exception:
            pass
    # deduplicate by filename — same script at multiple depths counts once
    _seen: set = set()
    _unique_ne: list = []
    for _n in non_executable:
        if _n not in _seen:
            _seen.add(_n)
            _unique_ne.append(_n)
    non_executable = _unique_ne
    if non_executable:
        findings.append(finding('BB', 'SIGNIFICANT',
            f'Shell scripts not marked executable: {", ".join(non_executable[:5])}',
            'Shell scripts exist but are not marked executable. Validators running these scripts will get permission denied errors.',
            [f'Fix: chmod +x {" ".join(non_executable[:5])}']))
    return findings


def detect_BC_mixed_line_endings(repo_dir, all_files):
    return []

def detect_BD_missing_contact(repo_dir, all_files):
    findings = []
    # Collect all README-like files (any depth, any case variant such as ReadMe.txt)
    readme_files = [
        f for f in all_files
        if f.name.lower() in README_NAMES
        or ('readme' in f.name.lower() and f.suffix.lower() in {'.md', '.txt', '.rst', ''})
    ]
    # Suppression requires a genuinely reachable contact mechanism — not just
    # the word "contact" or an author name.  Accepted mechanisms:
    #   1. Email address  (user@domain.tld)
    #   2. ORCiD URL      (orcid.org/XXXX-XXXX-XXXX-XXXX)
    #   3. GitHub issues  (github.com/owner/repo/issues)  — acceptable route for support
    _email_re = re.compile(r'[\w.+-]+@[\w-]+\.[a-z]{2,}', re.IGNORECASE)
    _orcid_re = re.compile(r'orcid\.org/\d{4}-\d{4}-\d{4}-\d{3}[\dX]', re.IGNORECASE)
    _gh_issues_re = re.compile(r'github\.com/[^/\s]+/[^/\s]+/issues', re.IGNORECASE)
    has_contact = False
    for readme_file in readme_files:
        content = read_file_safe(readme_file)
        if (_email_re.search(content)
                or _orcid_re.search(content)
                or _gh_issues_re.search(content)):
            has_contact = True
            break
    if not has_contact:
        # Check CITATION.cff as a fallback source of contact info
        citation_path = next((f for f in all_files if f.name == 'CITATION.cff'), None)
        if citation_path:
            try:
                cff_content = citation_path.read_text(errors='ignore')
                if 'email:' in cff_content:
                    return findings  # contact info found in CITATION.cff
            except Exception:
                pass
        findings.append(finding('BD', 'LOW CONFIDENCE',
            'No contact information found in README',
            'No author contact information was found. Validators who encounter problems have no way to reach the researcher for clarification.',
            ['Recommendation: add author name and contact email to README']))
    return findings


def detect_BE_pyc_files(repo_dir, all_files):
    findings = []
    pyc_files = [f for f in all_files if f.suffix.lower() in {'.pyc', '.pyo'} or '__pycache__' in str(f)]
    if pyc_files:
        findings.append(finding('BE', 'SIGNIFICANT',
            f'{len(pyc_files)} compiled Python file(s) committed',
            'Compiled .pyc files are committed. These are system-specific and will cause import errors on different Python versions or operating systems. Add *.pyc and __pycache__/ to .gitignore.',
            [f'Files: {", ".join(f.name for f in pyc_files[:5])}',
             'Fix: git rm --cached **/*.pyc and add to .gitignore']))
    return findings


def detect_BF_notebook_outputs_missing(repo_dir, all_files):
    findings = []
    notebooks = [f for f in all_files if f.suffix.lower() == '.ipynb']
    if not notebooks:
        return findings
    import json as _json
    for nb in notebooks:
        try:
            data = _json.loads(nb.read_text(encoding='utf-8', errors='ignore'))
            cells = data.get('cells', [])
            has_outputs = any(
                cell.get('outputs') for cell in cells
                if cell.get('cell_type') == 'code'
            )
            if not has_outputs:
                findings.append(finding('BF', 'SIGNIFICANT',
                    f'Notebook has no saved outputs: {nb.name}',
                    'This notebook has no saved cell outputs. Validators cannot see what the original results looked like without running the notebook themselves.',
                    [f'Evidence: {nb.name} — all output cells empty']))
        except Exception:
            continue
    return findings


def detect_BG_missing_acknowledgements(repo_dir, all_files):
    findings = []
    readme_file = None
    for f in all_files:
        if f.name.lower() in {'readme.md', 'readme.txt'}:
            readme_file = f
            break
    if not readme_file:
        return findings
    content = read_file_safe(readme_file).lower()
    code_files = [f for f in all_files if f.suffix.lower() == '.py']
    if len(code_files) > 5:
        has_funding = any(term in content for term in ['grant', 'funded', 'funding', 'acknowledge', 'nsf', 'nih', 'esrc', 'ukri', 'erc', 'support'])
        if not has_funding:
            findings.append(finding('BG', 'LOW CONFIDENCE',
                'No funding acknowledgement found',
                'No funding acknowledgement was found. Most funders require acknowledgement in associated code repositories.',
                ['Recommendation: add funding source to README']))
    return findings


def detect_BH_zip_bomb_risk(repo_dir, all_files):
    findings = []
    zip_files = [
        f for f in all_files
        if f.suffix.lower() in {'.zip', '.gz', '.tar', '.bz2', '.7z'}
        and not _is_single_file_compressed(f)
    ]
    if zip_files:
        findings.append(finding('BH', 'LOW CONFIDENCE',
            f'{len(zip_files)} compressed archive(s) committed',
            'Compressed archives are committed. Validators need to know what these contain and whether to extract them as part of the pipeline.',
            [f'Archives: {", ".join(f.name for f in zip_files[:5])}',
             'Document: should validators extract these, and what do they contain?']))
    return findings


def detect_BI_unicode_in_paths(repo_dir, all_files):
    return []

def detect_AV_hardcoded_dates(repo_dir, all_files):
    return []


def detect_BM_citation_cff(repo_dir, all_files):
    """Check CITATION.cff exists and has required fields."""
    findings = []
    cff_files = [f for f in all_files if f.name.lower() == 'citation.cff']
    if not cff_files:
        findings.append(finding(
            'BM', 'LOW CONFIDENCE',
            'No CITATION.cff found',
            'A CITATION.cff file makes your repository directly citable '
            'and is increasingly expected by journals and data archives.',
            ['Recommendation: create CITATION.cff — see https://citation-file-format.github.io/']
        ))
    # CITATION.cff present — INVENTORY already notes it; no [BM] finding needed.
    return findings


def detect_BN_codebook_reference_mismatch(repo_dir, all_files):
    """Check if any README references a codebook file that doesn't exist.

    Scans ALL readme files (not just the shallowest) so that sub-directory
    READMEs that reference a local codebook are also checked.
    """
    findings = []
    import re as _re
    _CODEBOOK_REF = _re.compile(r'codebook[\w\-]*\.\w+', _re.IGNORECASE)
    all_names = {f.name.lower() for f in all_files}
    readme_files = [
        f for f in all_files
        if f.name.lower() in {'readme.md', 'readme.txt', 'readme.rst'}
    ]
    if not readme_files:
        return findings

    already_reported: set = set()
    for readme_file in readme_files:
        try:
            content = readme_file.read_text(encoding='utf-8', errors='ignore')
        except Exception:
            continue
        for ref in _CODEBOOK_REF.findall(content):
            ref_lower = ref.lower()
            if ref_lower in already_reported:
                continue
            if ref_lower not in all_names:
                already_reported.add(ref_lower)
                near = ', '.join(
                    n for n in all_names
                    if 'codebook' in n or 'dict' in n
                ) or 'none'
                findings.append(finding(
                    'BN', 'LOW CONFIDENCE',
                    f'README references {ref} but file not found',
                    'The README mentions a codebook or data dictionary file '
                    'that does not appear to be present in the repository.',
                    [f'Referenced: {ref}',
                     f'Found in: {readme_file.name}',
                     f'Files present: {near}']
                ))
    return findings


def detect_BP_licence_in_readme_only(repo_dir, all_files):
    """Check if licence is stated in README but no LICENCE file exists."""
    findings = []
    has_licence_file = any(
        f.name.lower() in {'licence', 'license', 'licence.md',
                           'license.md', 'licence.txt', 'license.txt'}
        for f in all_files
    )
    if has_licence_file:
        return findings
    readme_file = None
    for f in all_files:
        if f.name.lower() in {'readme.md', 'readme.txt', 'readme.rst'}:
            readme_file = f
            break
    if not readme_file:
        return findings
    try:
        content = readme_file.read_text(encoding='utf-8', errors='ignore').lower()
    except Exception:
        return findings
    licence_terms = ['cc by', 'cc-by', 'mit license', 'apache 2', 'gpl', 'creative commons']
    if any(term in content for term in licence_terms):
        findings.append(finding(
            'BP', 'LOW CONFIDENCE',
            'Licence stated in README but no LICENCE file found',
            'The README mentions a licence but no dedicated LICENCE file '
            'exists. A separate LICENCE file is standard practice and '
            'required by many repositories and journals.',
            ['Recommendation: create a LICENCE file with the full licence text']
        ))
    return findings


def detect_BR_credentials_exposed(repo_dir, all_files):
    """Check for exposed credentials, API keys, or passwords."""
    findings = []
    import re as _re
    cred_patterns = _re.compile(
        r'(password|passwd|api_key|api_secret|secret_key|token|auth_token'
        r'|private_key|access_key|client_secret|database_url)\s*[=:]\s*(\S+)',
        _re.IGNORECASE
    )
    # Values that are clearly documentation placeholders, not real credentials
    _placeholder_val = _re.compile(
        r'^(your_|my_|example|dummy|test|placeholder|changeme|xxx+'
        r'|<[^>]+>|\[[^\]]+\]|enter_|insert_|add_your|replace_'
        r'|none|null|false|true|n/a|tbd|todo|[a-z0-9_]+_here'
        r'|[A-Z_]+_KEY_HERE|[A-Z_]+_TOKEN_HERE|[A-Z_]+_SECRET_HERE)',
        _re.IGNORECASE
    )
    env_files = [f for f in all_files if f.name.lower() in {
        '.env', '.env.local', '.env.production', '.env.development',
        'secrets.yml', 'secrets.yaml', 'credentials.json', 'credentials.yml'
    }]
    flagged = []
    evidence = []
    # always flag .env files present
    for f in env_files:
        flagged.append(f.name)
        evidence.append(f"Sensitive file present: {f.name}")
    # scan non-code config/secrets files for credential patterns
    # Source code (.py/.r/.jl) is handled by [AG] — avoid duplication
    code_exts = {'.py', '.r', '.rmd', '.jl', '.m'}
    check_exts = {'.yaml', '.yml', '.json', '.toml', '.cfg', '.ini', '.txt', '.md'}
    for f in all_files:
        if f.name.lower() in {ef.name.lower() for ef in env_files}:
            continue
        if f.suffix.lower() in code_exts:
            continue  # [AG] handles source code credentials
        if f.suffix.lower() not in check_exts:
            continue
        try:
            content = f.read_text(encoding='utf-8', errors='ignore')
            matches = cred_patterns.findall(content)
            # matches is list of (key_name, value) tuples; skip placeholder values
            real_matches = [
                (k, v) for k, v in matches
                if not _placeholder_val.match(v.strip('"\''))
            ]
            if real_matches:
                flagged.append(f.name)
                key_name = real_matches[0][0]
                evidence.append(f"{f.name}: credential pattern found ({key_name})")
        except Exception:
            pass
    if flagged:
        findings.append(finding(
            'BR', 'CRITICAL',
            f'Potential credentials or secrets detected in: {", ".join(flagged[:3])}',
            'Files containing passwords, API keys, or secrets must NEVER '
            'be published. Remove these files and rotate any exposed credentials '
            'immediately. Add .env to .gitignore before any further commits.',
            evidence[:5]
        ))
    return findings


def detect_BS_archive_code_present(repo_dir, all_files):
    """Check for vestigial code in archive/old directories."""
    findings = []
    archive_dirs = {"old", "archive", "deprecated", "unused", "backup", "old_versions"}

    def _in_archive_dir(f):
        """True if f sits inside a genuine archive folder.

        'deprecated' (and similar words) are skipped when they appear as a
        Java package path component rather than an intentional archival folder.
        Specifically: if the file is .java/.class AND the suspicious segment is
        immediately followed by another lowercase-word segment (Java package
        naming convention), it is a package name, not an archive folder.
        """
        try:
            parts = f.relative_to(repo_dir).parts
        except ValueError:
            parts = f.parts
        is_java = f.suffix.lower() in {'.java', '.class'}
        for i, part in enumerate(parts[:-1]):          # exclude filename itself
            if part.lower() not in archive_dirs:
                continue
            if is_java and i + 1 < len(parts) - 1:   # there is a next segment
                next_seg = parts[i + 1]
                if next_seg == next_seg.lower() and next_seg.isidentifier():
                    continue                           # Java package component
            return True
        return False

    archive_files = [
        f for f in all_files
        if f.suffix.lower() in CODE_EXTENSIONS
        and _in_archive_dir(f)
    ]
    if archive_files:
        findings.append(finding(
            'BS', 'LOW CONFIDENCE',
            f'Vestigial code files found in archive directories: {", ".join(f.name for f in archive_files[:3])}',
            'Code files in old/, archive/, or deprecated/ directories suggest '
            'version history managed by file duplication rather than git. '
            'Remove these before deposit to avoid confusion about which files '
            'are part of the active pipeline.',
            [f'Archive file: {f.relative_to(repo_dir)}' for f in archive_files[:5]]
        ))
    return findings


def detect_BT_spaces_in_filenames(repo_dir, all_files):
    """Check for problematic characters in path components (filenames or directory names).

    Three categories, in descending severity:
    1. Windows-illegal characters (< > : " | ? *) — prevent extraction on Windows entirely.
    2. Shell special characters ([ ] { }) — cause glob/brace expansion in bash/zsh.
    3. Spaces and non-ASCII characters — shell quoting failures and encoding issues.
    """
    findings = []

    # Characters that are illegal in Windows filenames/directory names.
    # Backslash and forward-slash are path separators and can't appear in
    # a path component, so they're excluded from the set.
    _WIN_ILLEGAL = frozenset('<>:"|?*')
    # Shell glob/brace-expansion characters.
    _GLOB_CHARS = frozenset('[]{}')

    win_illegal: dict = {}   # component -> set of offending chars
    glob_chars: dict = {}    # component -> representative Path
    spaces: dict = {}        # component -> representative Path
    non_ascii: dict = {}     # component -> representative Path

    for f in all_files:
        try:
            rel_parts = f.relative_to(repo_dir).parts
        except ValueError:
            rel_parts = f.parts
        for part in rel_parts:
            bad_win = frozenset(part) & _WIN_ILLEGAL
            bad_glob = frozenset(part) & _GLOB_CHARS
            if bad_win and part not in win_illegal:
                win_illegal[part] = bad_win
            if bad_glob and part not in glob_chars:
                glob_chars[part] = f
            if ' ' in part and part not in spaces:
                spaces[part] = f
            if not part.isascii() and part not in non_ascii:
                non_ascii[part] = f

    # ── SIGNIFICANT: Windows-illegal characters ──────────────────────────────
    if win_illegal:
        _wparts = list(win_illegal.keys())
        if len(_wparts) <= 3:
            _wtitle = (
                'Windows-illegal characters in path names: '
                + ', '.join(repr(p) for p in _wparts)
            )
        else:
            _wtitle = f'Windows-illegal characters in {len(_wparts)} path names'
        findings.append(finding(
            'BT', 'SIGNIFICANT',
            _wtitle,
            'Characters < > : " | ? * are illegal in Windows filenames and directory '
            'names. A validator on Windows cannot unzip this archive without errors — '
            'these names must be renamed before deposit.',
            [f'Problem name: {p!r}  (illegal chars: {" ".join(sorted(win_illegal[p]))})'
             for p in _wparts],
        ))

    # ── LOW CONFIDENCE: spaces, glob chars, non-ASCII ────────────────────────
    other: dict = {}
    issue_labels = []
    if spaces:
        other.update(spaces)
        issue_labels.append('spaces (shell quoting failures when scripting)')
    if glob_chars:
        other.update(glob_chars)
        issue_labels.append('square/curly brackets (shell glob expansion)')
    if non_ascii:
        other.update(non_ascii)
        issue_labels.append('non-ASCII characters (encoding failures on some systems)')

    if other:
        _oparts = list(other.keys())
        if len(_oparts) <= 3:
            _otitle = (
                'Problematic characters in path names: '
                + ', '.join(repr(p) for p in _oparts)
            )
        else:
            _otitle = f'Problematic characters in {len(_oparts)} path names'
        findings.append(finding(
            'BT', 'LOW CONFIDENCE',
            _otitle,
            'Path names contain ' + ' and '.join(issue_labels) + '. '
            'Replace spaces and special characters with underscores before deposit.',
            [f'Problem name: {p!r}' for p in _oparts[:10]],
        ))

    return findings


def detect_FL_long_filenames(repo_dir, all_files):
    """Failure Mode FL: Filenames with stems longer than 64 characters.

    Windows default MAX_PATH is 260 characters total — once nested inside
    any folder structure a long filename can hit this limit. Many tools
    (git, command-line utilities, certain IDEs) also silently truncate at
    64 or 80 characters. Long names with spaces compound the problem.

    Fires SIGNIFICANT when any filename stem exceeds 64 characters.
    """
    findings = []
    _FL_STEM_THRESHOLD = 64

    long_files = []
    seen_names: set = set()
    for f in all_files:
        if len(f.stem) > _FL_STEM_THRESHOLD and f.name not in seen_names:
            seen_names.add(f.name)
            long_files.append(f)

    if not long_files:
        return findings

    details = [
        f'{f.name!r} ({len(f.stem)}-char stem)'
        for f in long_files[:6]
    ]
    if len(long_files) > 6:
        details.append(f'(+{len(long_files) - 6} more)')

    findings.append(finding(
        'FL', 'SIGNIFICANT',
        f'Filename{"s" if len(long_files) != 1 else ""} too long '
        f'({len(long_files)} file{"s" if len(long_files) != 1 else ""} '
        f'exceed {_FL_STEM_THRESHOLD}-char stem limit)',
        'One or more filenames have stems longer than 64 characters. '
        'Windows has a default 260-character total path limit — once '
        'nested inside folders, very long names approach this ceiling. '
        'Many tools also truncate or fail silently at 64–80 characters. '
        'Long names with spaces require quoting in every command-line '
        'context. Recommended: shorten to under 64 characters and '
        'replace spaces with underscores.',
        details
    ))
    return findings


def detect_FD_duplicate_format_pairs(repo_dir, all_files):
    """Check for files sharing the same stem but stored in multiple formats.

    Common cause: Harvard Dataverse auto-generates a .tab copy of every .xlsx
    upload. Validators cannot tell which format is canonical or whether they
    contain identical data.
    """
    findings = []
    # Group files by lower-cased stem, ignoring dot-files and generated output.
    _GENERATED_NAMES = {'ASSESSMENT.md', 'CLEANING_REPORT.md'}
    stem_map: dict = {}
    for f in all_files:
        if f.name in _GENERATED_NAMES:
            continue
        if f.name.endswith('_DRAFT.md') or f.name.endswith('_DRAFT.txt'):
            continue
        stem_key = f.stem.lower()
        stem_map.setdefault(stem_key, []).append(f)

    pairs = []
    for stem_key, files in stem_map.items():
        exts = {f.suffix.lower() for f in files}
        if len(exts) >= 2:
            # Deduplicate by (name, suffix) — ignore depth duplicates.
            seen_names: set = set()
            unique = []
            for f in files:
                key = (f.name.lower(), f.suffix.lower())
                if key not in seen_names:
                    seen_names.add(key)
                    unique.append(f)
            if len({f.suffix.lower() for f in unique}) < 2:
                continue
            # Skip source/rendered-output pairs — these are never duplicates.
            _SCRIPT_EXTS = frozenset({
                '.r', '.rmd', '.rnw', '.qmd', '.py', '.do', '.sps', '.m',
                '.jl', '.sql', '.sh', '.bash', '.bat', '.ps1', '.tex',
            })
            _pair_exts = {f.suffix.lower() for f in unique}
            # PDF + script: printout or compiled LaTeX output
            if '.pdf' in _pair_exts and (_pair_exts - {'.pdf'}) <= _SCRIPT_EXTS:
                continue
            # HTML + analysis script: rendered R Markdown / Quarto / Jupyter output
            _RENDERED_SRC = frozenset({'.r', '.rmd', '.qmd', '.py', '.ipynb'})
            if '.html' in _pair_exts and (_pair_exts - {'.html'}) <= _RENDERED_SRC:
                continue
            # Skip web-font sets — Bootstrap and similar frameworks commit all
            # font formats (.eot, .ttf, .woff, .woff2, .otf, and .svg) because
            # each format targets a different browser.  This is not a
            # reproducibility-relevant duplicate.
            _FONT_EXTS = frozenset({'.eot', '.ttf', '.woff', '.woff2', '.otf'})
            _in_font_dir = any(
                any(part.lower() in {'fonts', 'font'} for part in f.parts)
                for f in unique
            )
            if _pair_exts <= (_FONT_EXTS | {'.svg'}) or (
                    _in_font_dir and _pair_exts <= (_FONT_EXTS | {'.svg', '.png'})):
                continue
            # Skip pairs where file sizes differ by more than 20:1 — these are
            # a primary data file alongside a small sidecar (caption, description,
            # or label), not genuine duplicate-format exports of the same dataset.
            # (e.g. EmpiricalFigures.pdf 31KB + EmpiricalFigures.txt 1.3KB ≈ 24:1)
            try:
                sizes = [f.stat().st_size for f in unique]
                max_sz = max(sizes)
                min_sz = min(s for s in sizes if s > 0)
                if max_sz / min_sz > 20:
                    continue
            except Exception:
                pass
            pairs.append(unique)

    if pairs:
        details = []
        for group in pairs:
            names = ', '.join(sorted(f.name for f in group))
            details.append(f'Duplicate formats: {names}')
        findings.append(finding(
            'FD', 'LOW CONFIDENCE',
            f'Duplicate-format file pairs ({len(pairs)} stem{"s" if len(pairs) != 1 else ""})',
            'Files with the same name but different extensions are present. '
            'Two common causes require different actions: '
            '(1) Platform auto-generation — e.g. Harvard Dataverse creates a .tab '
            'copy of every uploaded spreadsheet. Action: identify the canonical '
            'format and document or remove the auto-generated copy. '
            '(2) Deliberate multi-format export — e.g. the same SPSS dataset '
            'exported as .sav + .por + .csv for cross-software compatibility. '
            'Action: confirm in the README that all copies are in sync and '
            'explain why each format is provided. '
            'Note: duplicate formats do not block reproduction — a validator '
            'can always choose the most accessible format and proceed.',
            details
        ))
    return findings












_3D_MESH_EXTS = frozenset({
    '.off', '.ply', '.stl', '.obj', '.vtk', '.vtu', '.mesh', '.wrl', '.dae',
})

_3D_SOFTWARE_RE = re.compile(
    r'\b(?:meshlab|blender|paraview|open3d|pyvista|trimesh|'
    r'cloudcompare|cloud\s+compare|'
    r'point\s+cloud\s+library|pcl\b|'
    r'3[Dd]\s+(?:viewer|software|tool)|'
    r'mesh\s+(?:viewer|software|tool)|'
    r'view(?:ing|ed?)?\s+(?:the\s+)?(?:mesh(?:es)?|point\s*cloud|3[Dd]\s+file)|'
    r'render(?:ing|ed?)?\s+(?:the\s+)?(?:mesh(?:es)?|point\s*cloud)|'
    r'visuali[sz](?:e|ing|ed?)\s+(?:the\s+)?(?:mesh(?:es)?|point\s*cloud))\b',
    re.IGNORECASE,
)


def detect_3D_mesh_no_viewer(repo_dir, all_files):
    """Fires SIGNIFICANT when a deposit contains 3D mesh or point-cloud files
    but the README does not document any viewer or processing software.

    .off, .ply, .stl, .obj etc. require specific software (MeshLab, Blender,
    ParaView, Open3D, …) to open and render.  Without documentation validators
    cannot confirm their rendering matches the published figures.
    """
    findings = []
    mesh_files = [f for f in all_files if f.suffix.lower() in _3D_MESH_EXTS]
    if not mesh_files:
        return findings

    # Deduplicate by name.
    seen: set = set()
    unique_mesh = []
    for f in mesh_files:
        if f.name.lower() not in seen:
            seen.add(f.name.lower())
            unique_mesh.append(f)

    # Check README(s) for any 3D software mention (shallowest first).
    readme_files = sorted(
        [f for f in all_files
         if f.name.lower() in README_NAMES
         or ('readme' in f.name.lower()
             and f.suffix.lower() in {'.md', '.txt', '.rst', ''})],
        key=lambda f: len(f.relative_to(repo_dir).parts),
    )
    for readme in readme_files:
        content = read_file_safe(readme)
        if _3D_SOFTWARE_RE.search(content):
            return findings  # Suppressed: viewer software is documented

    # Build a compact evidence line grouped by extension.
    ext_groups: dict = {}
    for f in unique_mesh:
        ext_groups.setdefault(f.suffix.lower(), []).append(f.name)
    ext_summary = '; '.join(
        f'{len(v)} × {k} '
        f'({", ".join(v[:2])}{"…" if len(v) > 2 else ""})'
        for k, v in sorted(ext_groups.items())
    )

    n = len(unique_mesh)
    findings.append(finding(
        '3D', 'SIGNIFICANT',
        f'3D mesh/point-cloud files with no viewer documented '
        f'({n} file{"s" if n != 1 else ""})',
        '3D mesh and point-cloud formats (.off, .ply, .stl, .obj, …) require '
        'specific software to open and render — MeshLab, Blender, Open3D, '
        'ParaView, or similar. Without documentation of which software was used '
        'to generate and view these files, validators cannot confirm that their '
        'rendering matches the published figures.',
        [f'3D files present: {ext_summary}',
         'Recommendation: add a "Software requirements" or "Viewing the data" '
         'section to the README specifying which tool (e.g. MeshLab 2023.12, '
         'Blender 4.0) was used and how to load the files.'],
    ))
    return findings


def detect_ND_no_data_files(repo_dir, all_files):
    """Fires CRITICAL when a deposit contains no data or code — only publication
    materials (manuscript, SI document, figures as Word/PDF/image files).

    A deposit with no underlying data or code is fundamentally misdirected:
    validators have nothing to reproduce.
    """
    findings = []
    # Image formats are primary data in imaging, microscopy, photography, and
    # computational biology deposits.  3D mesh / point-cloud formats are primary
    # data in computational biology, palaeontology, and engineering deposits.
    _IMAGE_EXTS = {
        '.jpg', '.jpeg', '.png', '.tif', '.tiff', '.bmp', '.gif',
        '.webp', '.svg', '.eps', '.raw', '.cr2', '.nef', '.dng',
    }
    _MESH_EXTS = {
        '.off', '.ply', '.stl', '.obj', '.vtk', '.vtu', '.mesh',
        '.wrl', '.dae', '.fbx', '.glb', '.gltf',
    }
    _DATA_EXTS = DATA_EXTENSIONS | {'.tab', '.dat', '.nc', '.mat', '.sav', '.dta'} | _IMAGE_EXTS | _MESH_EXTS
    has_data = any(f.suffix.lower() in _DATA_EXTS for f in all_files)
    has_code = any(f.suffix.lower() in CODE_EXTENSIONS for f in all_files)
    if has_data or has_code:
        return findings

    # Has files at all?
    if not all_files:
        return findings

    # Characterise what IS present — only pure document formats count as
    # "publication materials".  Images and 3D mesh files are already handled
    # above as recognised data formats and would have returned early.
    _PUB_EXTS = {'.pdf', '.doc', '.docx', '.ppt', '.pptx'}
    pub_files = [f for f in all_files if f.suffix.lower() in _PUB_EXTS]
    pub_ratio = len(pub_files) / len(all_files) if all_files else 0
    if pub_ratio < 0.5:
        # Mostly unrecognised files — don't fire, may be a specialised deposit.
        return findings

    findings.append(finding(
        'ND', 'CRITICAL',
        'No data or code files found — deposit appears to contain only publication materials',
        'This deposit contains no data files and no code. A reproducibility '
        'deposit must include the underlying data and/or the scripts used to '
        'produce the reported results. Manuscript files, supplementary '
        'documents, and figures alone do not constitute a reproducible deposit.',
        ['No data files detected (e.g. .csv, .xlsx, .tab, .json, .dta, .sav)',
         'No code files detected (e.g. .py, .R, .do, .m, .jl)',
         'Recommendation: deposit the raw or processed data used in the analysis '
         'and, where possible, the analysis scripts']
    ))
    return findings


def detect_FW_figures_in_word(repo_dir, all_files):
    """Fires SIGNIFICANT when figures are stored as Word documents.

    Word files cannot be programmatically rendered, embed figures as
    non-extractable objects, and may display differently across versions and
    operating systems. Figures should be in a standard image format.
    """
    findings = []
    _WORD_EXTS = {'.doc', '.docx'}
    _FIGURE_PATTERNS = re.compile(
        r'\b(?:fig(?:ure)?|figure)[\s_\-]?\d+\b', re.IGNORECASE
    )
    word_figures = [
        f for f in all_files
        if f.suffix.lower() in _WORD_EXTS
        and _FIGURE_PATTERNS.search(f.stem)
    ]
    # Deduplicate by name.
    seen: set = set()
    unique = []
    for f in word_figures:
        if f.name.lower() not in seen:
            seen.add(f.name.lower())
            unique.append(f)
    if not unique:
        return findings
    findings.append(finding(
        'FW', 'SIGNIFICANT',
        f'{"Figure" if len(unique) == 1 else f"{len(unique)} figures"} stored as Word document{"" if len(unique) == 1 else "s"}',
        'Word documents (.doc/.docx) cannot be programmatically rendered and '
        'may display differently across Word versions and operating systems. '
        'Figures in a research deposit should be in a standard image format '
        '(PNG, SVG, PDF, or TIFF) so validators can verify them without '
        'opening proprietary software.',
        [f'Word figure: {f.name}' for f in unique] +
        ['Recommended formats: .png, .svg, .pdf, .tiff']
    ))
    return findings


def detect_UE_unicode_filenames(repo_dir, all_files):
    """Fires SIGNIFICANT when filenames contain non-ASCII characters.

    Em-dash (U+2014, —) is treated as a distinct sub-category because it is
    almost always a Word/typographic autocorrect substitution for a hyphen,
    not a deliberate symbol. Other non-ASCII characters (Greek letters,
    accented characters, etc.) are a separate concern.
    """
    findings = []
    _EM_DASH = '\u2014'      # —
    _EN_DASH = '\u2013'      # – (same cause, same fix)

    seen_names: set = set()
    dash_files = []    # contains em-dash or en-dash only (no other non-ASCII)
    symbol_files = []  # contains other non-ASCII (may also contain dashes)

    for f in all_files:
        if f.name in seen_names:
            continue
        seen_names.add(f.name)
        try:
            f.name.encode('ascii')
        except UnicodeEncodeError:
            non_ascii_chars = {c for c in f.name if ord(c) > 127}
            dash_chars = {_EM_DASH, _EN_DASH}
            if non_ascii_chars <= dash_chars:
                # only dashes — no other unicode
                dash_files.append(f)
            else:
                symbol_files.append(f)

    if not dash_files and not symbol_files:
        return findings

    # Fire one finding per category so the researcher sees the two problems
    # as distinct rather than a single lump count.
    if dash_files:
        nd = len(dash_files)
        findings.append(finding(
            'UE', 'SIGNIFICANT',
            f'Em-dash (—) in {"filename" if nd == 1 else f"{nd} filenames"} '
            f'— likely Word autocorrect substitution for a hyphen',
            'The em-dash (—) or en-dash (–) in these filenames is almost always '
            'inserted automatically by Microsoft Word or similar software when '
            'typing "--". It is not an ASCII hyphen (-) and will cause failures '
            'on case-sensitive filesystems and command-line tools. '
            'Replace — with a hyphen (-) or underscore (_) in all affected filenames.',
            [f'Em-dash in filename: {f.name}' for f in dash_files[:10]] +
            (['...and more'] if nd > 10 else [])
        ))

    if symbol_files:
        ns = len(symbol_files)
        findings.append(finding(
            'UE', 'SIGNIFICANT',
            f'Non-ASCII symbol{"s" if ns != 1 else ""} in '
            f'{"filename" if ns == 1 else f"{ns} filenames"}',
            'Filenames containing non-ASCII characters (Greek letters, accented '
            'characters, special symbols) cause silent failures on older Windows '
            'systems, some HPC environments, and many command-line tools. '
            'Validators on systems with different locale settings may be unable '
            'to open or reference these files. '
            'Replace non-ASCII characters with ASCII equivalents '
            '(e.g. Ω → Ohm, é → e, ñ → n).',
            [f'Non-ASCII filename: {f.name}' for f in symbol_files[:10]] +
            (['...and more'] if ns > 10 else [])
        ))

    return findings


def detect_NX_no_extension(repo_dir, all_files):
    """Fires SIGNIFICANT when files have no extension and are not known
    extensionless build/config files (Makefile, Dockerfile, etc.).

    A file with no extension cannot be identified automatically by type.
    """
    _KNOWN_EXTENSIONLESS = {
        'makefile', 'dockerfile', 'readme', 'license', 'licence',
        'authors', 'changelog', 'contributing', 'manifest', 'snakefile',
        'procfile', 'gemfile', 'rakefile', 'guardfile', 'pipfile',
        'cmakelists', 'justfile', 'vagrantfile', 'brewfile',
    }
    findings = []
    seen_names: set = set()
    no_ext_files = []
    for f in all_files:
        if f.name in seen_names:
            continue
        seen_names.add(f.name)
        if (not f.suffix
                and not f.name.startswith('.')
                and f.name.lower() not in _KNOWN_EXTENSIONLESS):
            no_ext_files.append(f)
    if not no_ext_files:
        return findings
    n = len(no_ext_files)
    findings.append(finding(
        'NX', 'SIGNIFICANT',
        f'{"File" if n == 1 else f"{n} files"} with no extension',
        'Files without a file extension cannot be identified by type automatically. '
        'Most GUI file managers, command-line tools, and validators will be unable '
        'to open them without first guessing the format. Add the correct extension '
        '(e.g. .csv, .png, .pdf) to each file before deposit.',
        [f'No-extension file: {f.name}' for f in no_ext_files[:10]] +
        (['...and more'] if n > 10 else [])
    ))
    return findings


_NN_STEM_RE = re.compile(r'^(.*?)(\d+)$')


def detect_NN_nonsequential_numbering(repo_dir, all_files):
    """Fires LOW CONFIDENCE when a numbered filename series contains a
    non-sequential jump that is likely a typo.

    Example: tapetal001, tapetal002, tapetal003, tapetal004, tapetal2005 —
    the jump from 004 to 2005 almost certainly means tapetal005 with a digit
    inserted by mistake.

    Only fires when the series is otherwise strictly sequential (all adjacent
    gaps are 1) and a single gap exceeds 100.
    """
    findings = []

    # Group files by (parent directory, lowercase extension, non-numeric prefix).
    groups: dict = {}
    for f in all_files:
        m = _NN_STEM_RE.match(f.stem)
        if not m:
            continue
        prefix, num_str = m.group(1), m.group(2)
        key = (f.parent, f.suffix.lower(), prefix)
        groups.setdefault(key, []).append((int(num_str), f.name))

    suspicious = []
    for (_parent, _ext, _prefix), entries in groups.items():
        if len(entries) < 3:
            continue
        entries.sort()
        nums = [e[0] for e in entries]
        fnames = [e[1] for e in entries]

        gaps = [nums[i + 1] - nums[i] for i in range(len(nums) - 1)]

        # Require the series to be otherwise sequential (at most one bad gap).
        n_unit = sum(1 for g in gaps if g == 1)
        if n_unit < len(gaps) - 1:
            continue

        for i, gap in enumerate(gaps):
            if gap > 100:
                suspicious.append(
                    f'{fnames[i]} → {fnames[i + 1]} '
                    f'(gap of {gap:,}; expected ~{nums[i] + 1})'
                )

    if not suspicious:
        return findings
    findings.append(finding(
        'NN', 'LOW CONFIDENCE',
        f'Non-sequential numbering gap in filename series '
        f'({len(suspicious)} gap{"s" if len(suspicious) != 1 else ""})',
        'One or more filename series contain a large non-sequential jump in an '
        'otherwise sequential run. This may be a numbering typo (e.g. tapetal2005 '
        'instead of tapetal005) or may indicate a missing file. Verify that the '
        'sequence is complete and correct the filename if it is a typo.',
        suspicious,
    ))
    return findings


def detect_IC_inconsistent_extension_case(repo_dir, all_files):
    """Fires LOW CONFIDENCE when the same extension appears in both upper and
    lower case (e.g. .csv and .CSV).

    On case-sensitive filesystems (Linux, most HPC) these are different
    extensions. Scripts that glob *.csv on Windows will miss .CSV files on Linux.
    """
    findings = []
    ext_variants: dict = {}   # lower_ext -> set of actual suffixes seen
    ext_examples: dict = {}   # lower_ext -> {actual_suffix -> first_filename}
    for f in all_files:
        if not f.suffix:
            continue
        lo = f.suffix.lower()
        ext_variants.setdefault(lo, set()).add(f.suffix)
        ext_examples.setdefault(lo, {}).setdefault(f.suffix, f.name)

    # Only fire when BOTH upper and lower casings of the same extension coexist
    # in the same deposit.  A single consistent casing (e.g. all .Rmd or all
    # .TIF) is not an inconsistency — it is a convention choice.
    inconsistent = {lo: variants for lo, variants in ext_variants.items()
                    if len(variants) > 1}

    if not inconsistent:
        return findings

    details = []
    for lo, variants in sorted(inconsistent.items()):
        examples = [ext_examples[lo][v] for v in sorted(variants)
                    if v in ext_examples[lo]]
        casing_desc = f'both {" and ".join(sorted(variants))} found'
        details.append(
            f'{lo}: {casing_desc} '
            f'(e.g. {", ".join(examples[:2])})'
        )
    findings.append(finding(
        'IC', 'LOW CONFIDENCE',
        f'Inconsistent extension casing '
        f'({len(inconsistent)} extension{"s" if len(inconsistent) != 1 else ""})',
        'The same extension appears in both uppercase and lowercase. '
        'On case-sensitive filesystems (Linux, most HPC clusters) .csv and .CSV '
        'are different extensions — scripts that glob *.csv on Windows will miss '
        '.CSV files when run on Linux. Standardise all extensions to lowercase.',
        details
    ))
    return findings


def detect_TV_numbering_padding_inconsistency(repo_dir, all_files):
    """LOW CONFIDENCE: numbered scripts in the same series mix zero-padded
    (01_, 02_) and non-padded (3_, 4_) prefixes.

    Zero-padding matters for alphabetical sort tools (ls, Windows Explorer, R
    list.files()) that put '10_' before '2_' when padding is absent.
    Only fires when the series has ≥2 scripts and BOTH padded and unpadded
    forms are present.
    """
    findings = []
    _NUM_PREFIX_RE = re.compile(r'^(\d+)(?:[_\-\s]|\.(?!\d))')
    code_files = [
        f for f in all_files
        if f.suffix.lower() in CODE_EXTENSIONS or f.name == 'Snakefile'
    ]
    padded = []    # have leading zero: 01_, 02_, ...
    unpadded = []  # no leading zero: 3_, 4_, 99_, ...
    for f in code_files:
        m = _NUM_PREFIX_RE.match(f.name)
        if not m:
            continue
        num_str = m.group(1)
        if len(num_str) > 4:
            continue  # date-stamp, skip
        if num_str.startswith('0') and len(num_str) > 1:
            padded.append(f.name)
        else:
            unpadded.append(f.name)
    if padded and unpadded and (len(padded) + len(unpadded)) >= 3:
        findings.append(finding(
            'TV', 'LOW CONFIDENCE',
            'Numbered scripts mix zero-padded and non-padded prefixes',
            'Some scripts use zero-padded numbers (01_, 02_) while others '
            'use unpadded numbers (3_, 4_). Alphabetical sort tools — '
            'including ls, Windows Explorer, and R\'s list.files() — will '
            'sort 10_ before 2_ when padding is inconsistent. '
            'Standardise to the same width (e.g. all two-digit padding: '
            '01_, 02_, 03_, 04_).',
            [f'Zero-padded: {", ".join(padded[:5])}',
             f'Unpadded: {", ".join(unpadded[:5])}']
        ))
    return findings


def detect_IC2_inconsistent_filename_spacing(repo_dir, all_files):
    """Fires LOW CONFIDENCE when file stems that appear to form a matched set
    differ only in whitespace (e.g. 'FigA' vs 'Fig A').

    Groups files by (stem-with-spaces-removed, lower-extension). If a group
    contains files with more than one distinct stem, the spacing is inconsistent.
    """
    findings = []
    seen_names: set = set()
    unique_files = []
    for f in all_files:
        if f.name not in seen_names:
            seen_names.add(f.name)
            unique_files.append(f)

    # Group by normalised key: spaces removed, lowercased stem + lowercased ext.
    norm_map: dict = {}
    for f in unique_files:
        key = (f.stem.replace(' ', '').lower(), f.suffix.lower())
        norm_map.setdefault(key, []).append(f)

    inconsistent_groups = [
        files for files in norm_map.values()
        if len({f.stem for f in files}) > 1
    ]
    if not inconsistent_groups:
        return findings

    details = []
    for group in inconsistent_groups:
        names = ', '.join(sorted(f.name for f in group))
        details.append(f'Spacing inconsistency: {names}')

    n = len(inconsistent_groups)
    findings.append(finding(
        'IC2', 'LOW CONFIDENCE',
        f'Inconsistent spacing in {"filename" if n == 1 else f"{n} filename pairs"}',
        'Files that appear to form a matched set differ only in internal spacing '
        '(e.g. "FigA" vs "Fig A"). This makes scripted pairing by filename '
        'unreliable and suggests an accidental naming inconsistency. '
        'Standardise the spacing convention across all related filenames.',
        details
    ))
    return findings


# Packages that require C/C++ system libraries not installable via pip alone
_SYSTEM_DEP_PACKAGES = {
    'geopandas': {'gdal', 'geos', 'proj'},
    'fiona':     {'gdal', 'geos'},
    'rasterio':  {'gdal', 'geos', 'proj'},
    'pyproj':    {'proj'},
    'shapely':   {'geos'},
    'gdal':      {'gdal'},
    'osgeo':     {'gdal', 'geos', 'proj'},
    'cartopy':   {'geos', 'proj'},
    'opencv-python': {'libopencv'},
    'cv2':       {'libopencv'},
    'lxml':      {'libxml2', 'libxslt'},
    'psycopg2':  {'postgresql'},
    'h5py':      {'hdf5'},
    'netcdf4':   {'netcdf4', 'hdf5'},
    'pyaudio':   {'portaudio'},
}
_SYSTEM_LIB_APT = {
    'gdal':       'libgdal-dev',
    'geos':       'libgeos-dev',
    'proj':       'libproj-dev',
    'libopencv':  'libopencv-dev',
    'libxml2':    'libxml2-dev',
    'libxslt':    'libxslt-dev',
    'postgresql': 'libpq-dev',
    'hdf5':       'libhdf5-dev',
    'netcdf4':    'libnetcdf-dev',
    'portaudio':  'portaudio19-dev',
}
_SYSTEM_LIB_BREW = {
    'gdal': 'gdal', 'geos': 'geos', 'proj': 'proj',
    'libopencv': 'opencv', 'postgresql': 'postgresql',
    'hdf5': 'hdf5', 'netcdf4': 'netcdf', 'portaudio': 'portaudio',
}




# EOL Python versions in Docker base images
_EOL_PYTHON_VERSIONS = {
    '2.7': 'January 2020', '3.4': 'March 2019', '3.5': 'September 2020',
    '3.6': 'December 2021', '3.7': 'June 2023', '3.8': 'October 2024',
}
_CURRENT_PYTHON = '3.12'






# OS-specific commands with their platform restrictions
_OS_SPECIFIC_COMMANDS = {
    # Linux-only (not macOS, not Windows)
    'nproc':          {'fails_on': 'macOS, Windows', 'reason': 'Linux-only CPU count command (use sysctl -n hw.ncpu on macOS)'},
    '/proc/meminfo':  {'fails_on': 'macOS, Windows', 'reason': 'Linux /proc filesystem not available'},
    '/proc/cpuinfo':  {'fails_on': 'macOS, Windows', 'reason': 'Linux /proc filesystem not available'},
    '/proc/':         {'fails_on': 'macOS, Windows', 'reason': 'Linux /proc filesystem not available'},
    '/dev/shm':       {'fails_on': 'macOS, Windows', 'reason': 'Linux shared memory filesystem'},
    'apt-get':        {'fails_on': 'macOS, Windows', 'reason': 'Debian/Ubuntu package manager only'},
    'apt ':           {'fails_on': 'macOS, Windows', 'reason': 'Debian/Ubuntu package manager only'},
    'yum ':           {'fails_on': 'macOS, Windows', 'reason': 'RHEL/CentOS package manager only'},
    'systemctl':      {'fails_on': 'macOS, Windows', 'reason': 'systemd service manager — Linux only'},
    'service ':       {'fails_on': 'macOS, Windows', 'reason': 'SysV init — Linux only'},
    # GNU-only (not macOS BSD tools)
    'grep -P':        {'fails_on': 'macOS (BSD grep), Windows', 'reason': 'PCRE mode (-P) requires GNU grep; macOS ships BSD grep'},
    'grep -P ':       {'fails_on': 'macOS (BSD grep), Windows', 'reason': 'PCRE mode (-P) requires GNU grep; macOS ships BSD grep'},
    'sed -i ':        {'fails_on': 'macOS (BSD sed)', 'reason': 'GNU sed -i syntax differs from BSD sed (needs empty string argument on macOS)'},
    'date -d':        {'fails_on': 'macOS (BSD date)', 'reason': 'GNU date -d not available in BSD date'},
    'readlink -f':    {'fails_on': 'macOS (BSD readlink)', 'reason': 'GNU readlink -f not available in BSD readlink'},
    'timeout ':       {'fails_on': 'macOS', 'reason': 'GNU coreutils timeout — not in macOS by default'},
    'xargs -r':       {'fails_on': 'macOS (BSD xargs)', 'reason': '-r flag not supported in BSD xargs'},
    # Windows-incompatible
    '#!/bin/bash':    {'fails_on': 'Windows (without WSL)', 'reason': 'Bash shebang — not available natively on Windows'},
    '#!/bin/sh':      {'fails_on': 'Windows (without WSL)', 'reason': 'POSIX sh shebang — not available natively on Windows'},
}





def detect_DG_undocumented_gui_steps(repo_dir, all_files):
    """Failure Mode DG: Pipeline has GUI/manual pre-processing steps not documented in README."""
    findings = []
    import re as _re

    # Signals of GUI-based pre-processing in code/comments
    gui_software_pat = _re.compile(
        r'\b(?:imagej|fiji|imagej2|napari|ilastik|imaris|cellprofiler|'
        r'leica|las\s*x|zeiss|zen\s+(?:blue|black)|nikon\s+nis|'
        r'simple\s+neurite\s+tracer|trainable\s+weka|'
        r'photoshop|illustrator|inkscape|gimp|'
        r'prism|graphpad|spss|stata\s+gui|excel\s+manual|'
        r'manual\s+(?:annotation|segmentation|tracing|curation|inspection|review)|'
        r'hand[- ](?:annotated|labelled|labeled|traced|segmented)|'
        r'manually\s+(?:drawn|traced|annotated|segmented|reviewed|curated|selected))',
        _re.IGNORECASE
    )

    # Proprietary/GUI file extensions in directory names or code
    gui_file_pat = _re.compile(
        r'\b\w+\.(?:lif|czi|nd2|oib|oif|vsi|svs|ndpi|scn|'
        r'roi|traces|ano|nrrd|nhdr|mrc|dm3|dm4)\b',
        _re.IGNORECASE
    )

    # Directory structure signals: raw image dir + derived output dir
    image_dirs = {'tiff', 'tif', 'raw', 'images', 'microscopy', 'confocal',
                  'raw_images', 'raw_data', 'acquisition'}
    derived_dirs = {'traces', 'rois', 'masks', 'annotations', 'segmented',
                    'processed', 'annotated', 'labelled', 'labeled'}

    gui_evidence = []
    all_dir_names = {f.parent.name.lower() for f in all_files}

    # Check for image + derived directory co-presence
    has_image_dir = bool(all_dir_names & image_dirs)
    has_derived_dir = bool(all_dir_names & derived_dirs)
    if has_image_dir and has_derived_dir:
        img_d = next(d for d in all_dir_names if d in image_dirs)
        der_d = next(d for d in all_dir_names if d in derived_dirs)
        gui_evidence.append(
            f'Directory structure suggests manual pre-processing: {img_d}/ → {der_d}/'
        )

    # Scan code/comments for GUI software references
    code_files = [f for f in all_files if f.suffix.lower() in {'.py', '.r', '.rmd', '.m', '.jl'}]
    gui_refs = []
    file_refs = []
    for f in code_files:
        try:
            src = f.read_text(encoding='utf-8', errors='ignore')
            for m in gui_software_pat.finditer(src):
                ref = m.group(0).strip()
                if ref.lower() not in [r.lower() for r in gui_refs]:
                    gui_refs.append(ref)
            for m in gui_file_pat.finditer(src):
                ref = m.group(0)
                if ref not in file_refs:
                    file_refs.append(ref)
        except Exception:
            pass

    if gui_refs:
        gui_evidence.append(f'GUI software referenced in code: {", ".join(gui_refs[:4])}')
    if file_refs:
        gui_evidence.append(f'GUI/proprietary file types referenced: {", ".join(file_refs[:4])}')

    if not gui_evidence:
        return findings

    # Check if README documents the manual steps
    readme_documents_gui = False
    for f in all_files:
        if f.name.lower() in {'readme.md', 'readme.txt', 'readme.rst'} and f.parent == repo_dir:
            try:
                src = f.read_text(encoding='utf-8', errors='ignore').lower()
                if any(term in src for term in [
                    'imageJ', 'fiji', 'manual', 'leica', 'las x', 'export', 'plugin',
                    'annotate', 'segment', 'trace', 'roi', 'pre-process', 'preprocess'
                ]):
                    readme_documents_gui = True
            except Exception:
                pass

    if readme_documents_gui:
        return findings

    # Classify the trigger to generate context-appropriate body text.
    _STATS_SW = {'prism', 'graphpad', 'spss', 'stata', 'excel'}
    _IMAGE_SW = {
        'imagej', 'fiji', 'imagej2', 'napari', 'ilastik', 'imaris',
        'cellprofiler', 'leica', 'zeiss', 'zen', 'nikon', 'photoshop',
        'illustrator', 'inkscape', 'gimp',
    }
    _refs_lower = {r.lower() for r in gui_refs}
    _has_stats = any(any(s in ref for s in _STATS_SW) for ref in _refs_lower)
    _has_image = any(any(s in ref for s in _IMAGE_SW) for ref in _refs_lower)
    _has_dir_only = not gui_refs and not file_refs  # triggered purely by dir structure

    if _has_stats and not _has_image:
        _context = (
            'statistical pre-processing steps (data preparation, variable recoding, '
            'transformation, or model fitting in a GUI environment) that must be '
            'performed before the main analysis script can run'
        )
    elif _has_image or _has_dir_only:
        _context = (
            'manual or GUI-based pre-processing steps (e.g. microscopy export, '
            'image annotation, manual tracing) that must be performed before the '
            'automated scripts can run'
        )
    else:
        _context = (
            'manual or GUI-based pre-processing steps that must be performed before '
            'the automated scripts can run'
        )

    details = gui_evidence[:5] + [
        'README does not document these manual/GUI pre-processing steps',
        'Validators cannot reproduce results without knowing these steps',
        'Fix: add a "Pre-processing" section to README documenting:',
        '  (1) Which GUI software is required (name, version)',
        '  (2) Exact steps performed (menus, settings, parameters)',
        '  (3) What files are produced and where to place them',
        '  (4) Any judgment calls made during manual steps',
    ]
    findings.append(finding(
        'DG', 'SIGNIFICANT',
        'Pipeline requires manual/GUI pre-processing steps not documented in README',
        f'The repository contains evidence of {_context}. '
        'These steps are not documented in the README. Validators will be unable to '
        'reproduce the analysis without knowing the exact software, settings, and '
        'procedures used.',
        details
    ))
    return findings


def detect_SP_specialist_software(repo_dir, all_files):
    """Failure Mode SP: Repository requires specialist or proprietary software."""
    findings = []

    # Extensions that imply a specific commercial platform.
    # Grouped by software so one finding fires per platform.
    _EXT_SW = {
        '.mph':      'COMSOL Multiphysics',
        '.odb':      'Abaqus/CAE',
        '.mxd':      'ArcMap (ArcGIS)',
        '.aprx':     'ArcGIS Pro',
        '.sas7bdat': 'SAS',
        '.sas':      'SAS',         # SAS program/syntax file
        '.sps':      'SPSS',        # SPSS syntax file
        '.sav':      'SPSS',        # SPSS data file
        '.por':      'SPSS',        # SPSS Portable format — same tool family
        '.dta':      'Stata',
        '.nb':       'Mathematica',
        '.jmp':      'JMP',
        '.tsc':      'TINA-TI (Texas Instruments circuit simulation)',
        # MATLAB data files — always require MATLAB or GNU Octave to open,
        # regardless of whether .m scripts are also present.
        '.mat':      'MATLAB or GNU Octave',
        # Simulink model files
        '.slx':      'Simulink (MATLAB)',
        '.mdl':      'Simulink (MATLAB)',
    }

    # MATLAB: .m files are the extension, but only flag MATLAB when no
    # Python/R/Julia is present (otherwise .m is likely Objective-C or a
    # mixed-language project where MATLAB is already expected).
    _has_m = any(f.suffix.lower() == '.m' for f in all_files)
    _has_py_r_jl = any(
        f.suffix.lower() in {'.py', '.r', '.rmd', '.jl'} for f in all_files
    )
    _has_ios_markers = any(
        f.suffix.lower() in {'.swift', '.xcodeproj', '.pbxproj'} for f in all_files
    )
    sw_files: dict[str, list[str]] = {}

    if _has_m and not _has_py_r_jl and not _has_ios_markers:
        _m_files = [f.name for f in all_files if f.suffix.lower() == '.m']
        sw_files['MATLAB'] = _m_files

    # Python API imports that identify a proprietary runtime dependency.
    _IMPORT_SW = [
        (re.compile(r'(?:^|[ \t])import\s+mph\b|^from\s+mph\b', re.MULTILINE),
         'COMSOL Multiphysics (mph Python API)'),
        (re.compile(r'(?:^|[ \t])import\s+abaqus\b|^from\s+abaqus\b', re.MULTILINE),
         'Abaqus (abaqus Python API)'),
    ]

    for f in all_files:
        sw = _EXT_SW.get(f.suffix.lower())
        if sw:
            sw_files.setdefault(sw, []).append(f.name)

    for f in all_files:
        if f.suffix.lower() not in {'.py', '.ipynb'}:
            continue
        try:
            src = f.read_text(encoding='utf-8', errors='ignore')
            for pat, sw in _IMPORT_SW:
                if pat.search(src):
                    sw_files.setdefault(sw, []).append(f'import in {f.name}')
        except Exception:
            pass

    for sw, raw_files in sorted(sw_files.items()):
        # deduplicate, cap list length
        seen: list[str] = []
        for fn in raw_files:
            if fn not in seen:
                seen.append(fn)
        display = seen[:6]
        suffix = f' (+ {len(seen) - 6} more)' if len(seen) > 6 else ''

        _octave_note = (
            ' GNU Octave (free) can read many .mat files but may not support '
            'all MATLAB toolbox functions.'
            if 'Octave' in sw else ''
        )
        findings.append(finding(
            'SP', 'LOW CONFIDENCE',
            f'Proprietary software required: {sw}',
            f'Proprietary software required: {sw}. Validators need a valid licence '
            f'for {sw} to reproduce results.{_octave_note} '
            f'Confirm the required version in your README.',
            [f'Files/imports: {", ".join(display)}{suffix}',
             f'Action: add "{sw} vX.Y" to README software requirements']
        ))

    return findings


def detect_EP_data_provenance(repo_dir, all_files):
    """Failure Mode EP: Data-only deposit with no extraction methodology documented."""
    findings = []

    code_files = [f for f in all_files if f.suffix.lower() in CODE_EXTENSIONS]
    data_files = [f for f in all_files if f.suffix.lower() in DATA_EXTENSIONS]

    # Only applies to near-data-only repos (reader scripts are fine, pipelines are not)
    if len(code_files) >= 3 or not data_files:
        return findings

    # No README → [A] is already firing; skip [EP] to avoid redundant noise
    # Depth <= 2 prevents deep sub-project READMEs from suppressing the finding
    readme_files = [f for f in all_files
                    if f.name.lower() in README_NAMES
                    and len(f.relative_to(repo_dir).parts) <= 2]
    if not readme_files:
        return findings
    _methodology_keywords = [
        'extract', 'collect', 'scrap', 'mine', 'generat',
        'methodolog', 'how we', 'how the data', 'data collection',
        'data extraction', 'data generation', 'dataset construction',
        'study', 'survey', 'annotated', 'labeled', 'labelled',
        'taxonomy', 'corpus', 'experiment',
        # retrieval / download provenance (e.g. iNaturalist exports)
        'sourced from', 'retrieved', 'downloaded', 'queried',
        'observations', 'query url', 'retrieval date', 'iNaturalist',
        'portal', 'data source', 'data sourced',
    ]
    has_methodology = False
    for readme in readme_files:
        content = read_file_safe(readme).lower()
        if any(kw in content for kw in _methodology_keywords):
            has_methodology = True
            break

    if not has_methodology:
        findings.append(finding(
            'EP', 'SIGNIFICANT',
            'Data-only repository — no extraction methodology documented',
            'This repository contains primarily data files with no documented '
            'explanation of how the data was produced or collected. Validators '
            'cannot verify collection bias, check for errors in the extraction '
            'process, or regenerate the dataset if needed. Add a README section '
            'describing the collection or extraction process, tools used, and '
            'date of collection.',
            ['No methodology keywords found in README '
             '(e.g. "extracted", "collected", "scraped", "generated")',
             'Recommendation: add a "Data collection" or "Methodology" section '
             'to README describing how this data was produced']
        ))

    return findings


def detect_DF_external_data_no_fetch(repo_dir, all_files):
    """Failure Mode DF: README references external data URL but no fetch script or checksum."""
    findings = []
    import re as _re
    # Data repository URL patterns
    data_repo_pat = _re.compile(
        r'https?://(?:www\.)?'
        r'(?:zenodo\.org/(?:record|doi|deposit)|'
        r'figshare\.com/(?:articles|collections)|'
        r'osf\.io/[a-z0-9]+|'
        r'datadryad\.org/stash/dataset|'
        r'dataverse\.(?:harvard|nl|uc3)\.edu|'
        r'pangaea\.de/10\.1594|'
        r'data\.mendeley\.com/datasets)'
        r'[/\w\-.?=&]*',
        _re.IGNORECASE
    )
    # Fetch script indicators
    fetch_script_names = {
        'download_data.sh', 'download_data.py', 'fetch_data.sh', 'fetch_data.py',
        'get_data.sh', 'get_data.py', 'download.sh', 'download.py',
        'setup_data.sh', 'setup_data.py', 'data_download.sh', 'data_download.py',
    }
    # Checksum patterns in any file
    checksum_pat = _re.compile(
        r'(?:md5|sha256|sha512)[:\s]+[0-9a-f]{32,128}',
        _re.IGNORECASE
    )
    # Find data repo URLs in README
    data_urls = []
    for f in all_files:
        if f.name.lower() in {'readme.md', 'readme.txt', 'readme.rst'}:
            try:
                src = f.read_text(encoding='utf-8', errors='ignore')
                for m in data_repo_pat.finditer(src):
                    data_urls.append(m.group(0))
            except Exception:
                pass
    if not data_urls:
        return findings
    # Check for fetch script
    all_names_lower = {f.name.lower() for f in all_files}
    has_fetch_script = bool(all_names_lower & fetch_script_names)
    # Also check for wget/curl commands in any .sh file
    if not has_fetch_script:
        for f in all_files:
            if f.suffix.lower() in {'.sh', '.bash', '.py'}:
                try:
                    src = f.read_text(encoding='utf-8', errors='ignore')
                    if _re.search(r'wget\s+.*zenodo|curl\s+.*zenodo|requests\.get.*zenodo'
                                  r'|wget\s+.*figshare|pooch\.retrieve|urllib.*download', src, _re.IGNORECASE):
                        has_fetch_script = True
                        break
                except Exception:
                    pass
    # Check for checksums
    has_checksum = False
    for f in all_files:
        try:
            src = f.read_text(encoding='utf-8', errors='ignore')
            if checksum_pat.search(src):
                has_checksum = True
                break
        except Exception:
            pass
    if has_fetch_script and has_checksum:
        return findings
    details = [f'External data URL found: {data_urls[0]}']
    if len(data_urls) > 1:
        details.append(f'  (and {len(data_urls)-1} more)')
    if not has_fetch_script:
        details += [
            'No download/fetch script found in repository',
            'Validators must manually locate and download the data',
            'Fix: add a download_data.sh or download_data.py script, e.g.:',
            '  wget -O data/dataset.tif https://zenodo.org/record/.../files/dataset.tif',
        ]
    if not has_checksum:
        details += [
            'No checksum (MD5/SHA256) provided for the data file',
            'Validators cannot verify they downloaded the identical version',
            'Fix: add to README: SHA256: <hash of downloaded file>',
        ]
    details += [
        'Also document: exact dataset version/DOI used (not just the record URL)',
    ]
    findings.append(finding(
        'DF', 'SIGNIFICANT',
        'External data URL present but no fetch script or checksum',
        'The README references an external data repository (Zenodo, Figshare, OSF, etc.) '
        'but provides no automated download script and no checksum. A URL alone is '
        'insufficient for reproducibility: URLs can break, datasets can be updated, '
        'and validators cannot confirm they downloaded the identical version used in '
        'the original analysis.',
        details
    ))
    return findings


def detect_DE_pytorch_nondeterminism(repo_dir, all_files):
    """Failure Mode DE: PyTorch seeds set but use_deterministic_algorithms absent."""
    findings = []
    import re as _re
    py_files = [f for f in all_files if f.suffix.lower() == '.py']
    if not py_files:
        return findings
    has_torch = False
    has_manual_seed = False
    has_deterministic = False
    has_cudnn_deterministic = False
    for f in py_files:
        try:
            src = f.read_text(encoding='utf-8', errors='ignore')
            if 'torch' in src:
                has_torch = True
            if _re.search(r'torch\.manual_seed|torch\.cuda\.manual_seed', src):
                has_manual_seed = True
            if _re.search(r'torch\.use_deterministic_algorithms\s*\(\s*True', src):
                has_deterministic = True
            if _re.search(r'cudnn\.deterministic\s*=\s*True', src):
                has_cudnn_deterministic = True
        except Exception:
            pass
    if not has_torch or not has_manual_seed:
        return findings
    if has_deterministic:
        return findings
    details = [
        'torch.manual_seed() detected — seeds are set',
        'torch.use_deterministic_algorithms(True) — NOT FOUND',
    ]
    if not has_cudnn_deterministic:
        details.append('torch.backends.cudnn.deterministic = True — NOT FOUND')
    details += [
        'Seeds alone do not guarantee reproducibility in PyTorch:',
        '  certain CUDA ops (atomics, non-deterministic reductions) vary between runs',
        'Fix: add after seed setup:',
        '  torch.use_deterministic_algorithms(True)',
        '  torch.backends.cudnn.deterministic = True',
        '  torch.backends.cudnn.benchmark = False',
        'Note: some ops raise an error under deterministic mode — use',
        '  torch.use_deterministic_algorithms(True, warn_only=True) to identify them',
    ]
    findings.append(finding(
        'DE', 'SIGNIFICANT',
        'PyTorch seeds set but deterministic mode not enabled',
        'torch.manual_seed() is set but torch.use_deterministic_algorithms(True) '
        'is absent. Certain CUDA operations remain non-deterministic even with seeds '
        'set — results will vary between runs on GPU hardware. Seeds are necessary '
        'but not sufficient for PyTorch reproducibility.',
        details
    ))
    return findings


def detect_DD_os_specific_commands(repo_dir, all_files):
    """Failure Mode DD: Shell scripts use OS-specific commands contradicting cross-platform claims."""
    findings = []
    import re as _re
    shell_files = [f for f in all_files if f.suffix.lower() in {'.sh', '.bash', '.zsh'}
                   or (f.suffix == '' and f.name.lower() in {'makefile'})]
    # Also check files with bash shebang
    for f in all_files:
        if f.suffix.lower() not in {'.py', '.r', '.rmd', '.md', '.txt', '.csv', '.tsv',
                                     '.json', '.yml', '.yaml', '.toml', '.lock'} and f not in shell_files:
            try:
                first_line = f.read_text(encoding='utf-8', errors='ignore').split('\n')[0]
                if '#!/bin/bash' in first_line or '#!/bin/sh' in first_line:
                    shell_files.append(f)
            except Exception:
                pass
    if not shell_files:
        return findings
    # Check if README claims cross-platform support
    claims_crossplatform = False
    platform_claim = ''
    for f in all_files:
        if f.name.lower() in {'readme.md', 'readme.txt', 'readme.rst'}:
            try:
                src = f.read_text(encoding='utf-8', errors='ignore')
                if _re.search(r'(windows|macos|mac os|cross.platform|platform.*linux.*mac|platform.*win)',
                              src, _re.IGNORECASE):
                    claims_crossplatform = True
                    m = _re.search(r'\*{0,2}[Pp]latform\*{0,2}:?\s*([^\n]+)', src)
                    if m:
                        platform_claim = m.group(1).strip()
            except Exception:
                pass
    # Scan shell files for OS-specific commands
    hits = []  # (filename, command, info)
    seen_commands = set()
    for f in shell_files:
        try:
            src = f.read_text(encoding='utf-8', errors='ignore')
            for cmd, info in _OS_SPECIFIC_COMMANDS.items():
                cmd_key = cmd.strip()
                if cmd in src and cmd_key not in seen_commands:
                    seen_commands.add(cmd_key)
                    hits.append((f.name, cmd_key, info))
        except Exception:
            pass
    # For non-cross-platform repos, still flag Linux-specific /proc and nproc
    # but only if multiple platforms are claimed
    if not hits:
        return findings
    severity = 'SIGNIFICANT' if claims_crossplatform else 'LOW CONFIDENCE'
    details = []
    if claims_crossplatform and platform_claim:
        details.append(f'README claims: Platform: {platform_claim}')
    details.append('OS-specific commands detected in shell scripts:')
    for fname, cmd, info in hits[:8]:
        details.append(f'  {fname}: `{cmd}` — fails on {info["fails_on"]} ({info["reason"]})')
    if claims_crossplatform:
        details.append('Fix: replace Linux-specific commands with cross-platform equivalents,')
        details.append('  or containerise with Docker to guarantee consistent environment.')
    else:
        details.append('Fix: document OS requirements in README, or use Docker for portability.')
    findings.append(finding(
        'DD', severity,
        'Shell scripts use OS-specific commands' +
        (' — contradicts cross-platform README claim' if claims_crossplatform else ''),
        'Shell scripts contain commands that only work on specific operating systems. '
        + ('The README claims cross-platform support, which is incorrect. '
           if claims_crossplatform else
           'Validators on other operating systems will encounter errors. ') +
        'Affected commands are listed below.',
        details
    ))
    return findings


def detect_DC_monorepo_independent_subprojects(repo_dir, all_files):
    """Failure Mode DC: Repo contains multiple independent sub-projects presented as a pipeline."""
    findings = []
    import re as _re
    # Look for sub-directories that each contain independent entry points
    # Signal: top-level subdirs each containing code + their own dep file
    subdirs = {}
    for f in all_files:
        try:
            rel = f.relative_to(repo_dir)
            parts = rel.parts
            if len(parts) >= 2:
                subdir = parts[0]
                subdirs.setdefault(subdir, []).append(f)
        except Exception:
            pass
    # Peel single common top-level wrapper (e.g. zip extracts as monorepo/paper1/...)
    if len(subdirs) == 1:
        wrapper = list(subdirs.keys())[0]
        wrapper_files = subdirs[wrapper]
        subdirs = {}
        for f in wrapper_files:
            try:
                rel = f.relative_to(repo_dir)
                parts = rel.parts
                if len(parts) >= 3:
                    subdir = parts[1]
                    subdirs.setdefault(subdir, []).append(f)
            except Exception:
                pass
    # A sub-project must have: code files + its own dep file or README
    dep_files = {'requirements.txt', 'environment.yml', 'renv.lock',
                 'pyproject.toml', 'pipfile', 'setup.py', 'project.toml'}
    subprojects = []
    for subdir, files in subdirs.items():
        fnames = {f.name.lower() for f in files}
        suffixes = {f.suffix.lower() for f in files}
        has_code = bool(suffixes & {'.py', '.r', '.rmd', '.jl', '.m', '.do'})
        has_deps = bool(fnames & dep_files)
        has_readme = bool(fnames & {'readme.md', 'readme.txt', 'readme.rst'})
        if has_code and (has_deps or has_readme):
            langs = []
            if '.py' in suffixes: langs.append('Python')
            if suffixes & {'.r', '.rmd'}: langs.append('R')
            if '.jl' in suffixes: langs.append('Julia')
            if '.m' in suffixes: langs.append('MATLAB')
            if '.do' in suffixes: langs.append('Stata')
            entry = next(
                (f.name for f in files
                 if f.suffix.lower() in {'.py', '.r', '.rmd', '.jl', '.m', '.do'}
                 and _re.match(r'^(main|run|analyse|analyze|model|pipeline|app)',
                              f.name.lower())),
                next((f.name for f in files
                      if f.suffix.lower() in {'.py', '.r', '.rmd', '.jl', '.m', '.do'}),
                     None)
            )
            subprojects.append({
                'dir': subdir,
                'langs': langs,
                'entry': entry,
                'has_deps': has_deps,
            })
    if len(subprojects) < 2:
        return findings
    # Check if root README introduces the sub-projects
    root_readme_ok = False
    for f in all_files:
        if f.name.lower() in {'readme.md', 'readme.txt'} and f.parent == repo_dir:
            try:
                src = f.read_text(encoding='utf-8', errors='ignore').lower()
                if all(sp['dir'].lower() in src for sp in subprojects):
                    root_readme_ok = True
            except Exception:
                pass
    details = [
        f'{len(subprojects)} independent sub-projects detected:',
    ]
    for sp in subprojects:
        lang_str = '/'.join(sp['langs']) if sp['langs'] else 'unknown'
        entry_str = f' (entry: {sp["entry"]})' if sp['entry'] else ''
        details.append(f'  {sp["dir"]}/: {lang_str}{entry_str}')
    details += [
        'These are independent pipelines — they do NOT need to run in sequence',
        'Each sub-project has its own dependencies and data',
    ]
    if not root_readme_ok:
        details.append('Root README does not clearly introduce all sub-projects')
    details.append(
        'Fix: root README must explain each sub-project and how to run each independently'
    )
    findings.append(finding(
        'DC', 'SIGNIFICANT',
        f'Monorepo: {len(subprojects)} independent sub-projects in one repository',
        'The repository contains multiple independent sub-projects in separate '
        'subdirectories. These are not sequential pipeline steps — each has its own '
        'language, dependencies, and data. Presenting them as a sequence in QUICKSTART '
        'will confuse validators. The root README must clearly explain each '
        'sub-project and provide separate run instructions for each.',
        details
    ))
    return findings


def detect_DB_shiny_app(repo_dir, all_files):
    """Failure Mode DB: Repository is a Shiny app — needs interactive verification docs."""
    findings = []
    import re as _re
    r_files = _researcher_r_files(all_files, repo_dir)
    if not r_files:
        return findings
    shiny_pat = _re.compile(
        r'shiny::(runApp|fluidPage|navbarPage|tabPanel|renderPlot|renderTable|renderText'
        r'|reactive|observeEvent|eventReactive|shinyApp|shinyUI|shinyServer)'
        r'|library\s*\(\s*shiny\s*\)',
        _re.IGNORECASE
    )
    runapp_pat = _re.compile(r'shiny::runApp|runApp\s*\(', _re.IGNORECASE)
    shiny_files = []
    has_runapp = False
    for f in r_files:
        try:
            src = f.read_text(encoding='utf-8', errors='ignore')
            if shiny_pat.search(src):
                shiny_files.append(f.name)
            if runapp_pat.search(src):
                has_runapp = True
        except Exception:
            pass
    # Also detect by canonical Shiny file names
    shiny_names = {'app.r', 'server.r', 'ui.r', 'app.R', 'server.R', 'ui.R'}
    has_shiny_files = any(f.name in shiny_names for f in all_files)
    if not shiny_files and not has_shiny_files:
        return findings
    # Check for interactive verification docs in README
    has_interaction_docs = False
    for f in all_files:
        if f.name.lower() in {'readme.md', 'readme.txt', 'readme.rst'}:
            try:
                src = f.read_text(encoding='utf-8', errors='ignore').lower()
                if any(term in src for term in
                       ['interact', 'select', 'expected output', 'screenshot',
                        'step-by-step', 'verify', 'ui control']):
                    has_interaction_docs = True
            except Exception:
                pass
    # Determine app directory
    app_dirs = [f.parent for f in all_files
                if f.name.lower() in {'server.r', 'ui.r'} and
                f.parent != repo_dir]
    app_dir_str = (
        "'" + str(app_dirs[0].relative_to(repo_dir)).replace('\\', '/') + "'"
        if app_dirs else '.'
    )
    details = [
        'Shiny files detected: ' + (', '.join(shiny_files[:4]) if shiny_files else ', '.join(shiny_names & {f.name for f in all_files})),
        'shiny::runApp() launches a web server — no output files are generated',
        'Validators must interact with the UI and visually verify outputs',
    ]
    if not has_interaction_docs:
        details.append('README contains no interaction instructions or expected output descriptions')
    details += [
        'Fix: README must include:',
        '  (1) Launch command: Rscript -e "shiny::runApp(' + app_dir_str + ')"',
        '  (2) Step-by-step UI interaction instructions for each figure/table',
        '  (3) Expected values for specific input combinations',
        '  (4) Screenshots or descriptions of expected UI state',
    ]
    findings.append(finding(
        'DB', 'SIGNIFICANT',
        'Repository is a Shiny web application — interactive verification required',
        'This repository contains a Shiny app. Reproduction requires launching '
        'the app, interacting with UI controls, and manually verifying that charts '
        'and tables match published figures. No automated file comparison is possible. '
        'Without detailed interaction instructions in the README, validators cannot '
        'assess reproducibility.',
        details
    ))
    return findings


def detect_DA_nlp_model_not_in_dockerfile(repo_dir, all_files):
    """Failure Mode DA: Code loads spaCy/NLTK models not downloaded in Dockerfile."""
    findings = []
    import re as _re
    dockerfiles = [f for f in all_files
                   if f.name == 'Dockerfile' or f.name.startswith('Dockerfile.')]
    code_files = [f for f in all_files if f.suffix.lower() in CODE_EXTENSIONS]
    if not dockerfiles or not code_files:
        return findings
    # Patterns for NLP model loads in code
    spacy_load = _re.compile(r'spacy\.load\s*\([\"\']([^\"\']+)[\"\']', _re.IGNORECASE)
    nltk_download = _re.compile(r'nltk\.download\s*\([\"\']([^\"\']+)[\"\']', _re.IGNORECASE)
    # Collect models referenced in code
    spacy_models = set()
    nltk_corpora = set()
    for f in code_files:
        try:
            src = f.read_text(encoding='utf-8', errors='ignore')
            for m in spacy_load.finditer(src):
                spacy_models.add(m.group(1))
            for m in nltk_download.finditer(src):
                nltk_corpora.add(m.group(1))
        except Exception:
            pass
    if not spacy_models and not nltk_corpora:
        return findings
    # Check Dockerfile for download commands
    missing_spacy = []
    missing_nltk = []
    for df in dockerfiles:
        try:
            dsrc = df.read_text(encoding='utf-8', errors='ignore')
            for model in spacy_models:
                if model not in dsrc:
                    missing_spacy.append(model)
            for corpus in nltk_corpora:
                if corpus not in dsrc:
                    missing_nltk.append(corpus)
        except Exception:
            pass
    if not missing_spacy and not missing_nltk:
        return findings
    details = []
    for model in missing_spacy:
        details += [
            f'spacy.load("{model}") in code but not downloaded in Dockerfile',
            f'Fix: add to Dockerfile after pip install:',
            f'  RUN python -m spacy download {model}',
        ]
    for corpus in missing_nltk:
        details += [
            f'nltk corpus "{corpus}" used in code but not downloaded in Dockerfile',
            f'Fix: add to Dockerfile:',
            f'  RUN python -c "import nltk; nltk.download(\'{corpus}\')"',
        ]
    details.append('Container will build successfully but crash at runtime')
    findings.append(finding(
        'DA', 'SIGNIFICANT',
        'NLP model/corpus loaded in code but not installed in Dockerfile',
        'Code calls spacy.load() or nltk.download() for models that are not '
        'pip-installable. These must be downloaded separately in the Dockerfile. '
        'The container will build without error but crash immediately at runtime '
        'when the missing model is requested.',
        details
    ))
    return findings


def detect_CZ_eol_docker_base_image(repo_dir, all_files):
    """Failure Mode CZ: Dockerfile uses an end-of-life Python base image."""
    findings = []
    import re as _re
    dockerfiles = [f for f in all_files
                   if f.name == 'Dockerfile' or f.name.startswith('Dockerfile.')]
    if not dockerfiles:
        return findings
    from_pat = _re.compile(
        r'^FROM\s+(\S+)',
        _re.IGNORECASE | _re.MULTILINE
    )
    version_pat = _re.compile(r'python[:/](\d+\.\d+)', _re.IGNORECASE)
    for f in dockerfiles:
        try:
            src = f.read_text(encoding='utf-8', errors='ignore')
        except Exception:
            continue
        eol_images = []
        for m in from_pat.finditer(src):
            image = m.group(1)
            vm = version_pat.search(image)
            if not vm:
                continue
            ver = vm.group(1)
            if ver in _EOL_PYTHON_VERSIONS:
                eol_images.append((image, ver, _EOL_PYTHON_VERSIONS[ver]))
        if eol_images:
            image, ver, eol_date = eol_images[0]
            evidence = [
                f'Dockerfile: FROM {image}',
                f'Python {ver} EOL: {eol_date}',
                f'Fix: update to FROM python:{_CURRENT_PYTHON}-slim',
                'Test that requirements.txt packages are compatible with '
                f'Python {_CURRENT_PYTHON}',
            ]
            if len(eol_images) > 1:
                evidence.append(
                    'Also EOL: ' + ', '.join(
                        f'{img} (Python {v}, EOL {d})' for img, v, d in eol_images[1:]
                    )
                )
            findings.append(finding(
                'CZ', 'SIGNIFICANT',
                f'Dockerfile uses end-of-life Python {ver} base image',
                f'The base image {image} uses Python {ver}, which reached '
                f'end-of-life in {eol_date}. EOL images receive no security '
                f'patches and may become unavailable or broken as registries '
                f'phase out old versions. Validators pulling this image may '
                f'encounter errors or security warnings.',
                evidence
            ))
    return findings


def detect_CY_checksum_not_verified(repo_dir, all_files):
    """Failure Mode CY: README documents a checksum but code never verifies it."""
    findings = []
    import re as _re
    # Find checksum mentions in README
    sha_pat = _re.compile(
        r'(?:sha256|sha512|sha1|md5)[\s:=]+([0-9a-f]{32,128})',
        _re.IGNORECASE
    )
    readme_checksums = []
    for f in all_files:
        if f.name.lower() in {'readme.md', 'readme.txt', 'readme.rst'}:
            try:
                src = f.read_text(encoding='utf-8', errors='ignore')
                for m in sha_pat.finditer(src):
                    readme_checksums.append(m.group(0)[:60])
            except Exception:
                pass
    if not readme_checksums:
        return findings
    # Check if any code file verifies checksums
    verify_pat = _re.compile(
        r'hashlib\.|sha256\(|sha512\(|md5\(|hexdigest\(|checksum',
        _re.IGNORECASE
    )
    code_files = [f for f in all_files if f.suffix.lower() in CODE_EXTENSIONS]
    code_verifies = False
    for f in code_files:
        try:
            if verify_pat.search(f.read_text(encoding='utf-8', errors='ignore')):
                code_verifies = True
                break
        except Exception:
            pass
    if code_verifies:
        return findings
    findings.append(finding(
        'CY', 'SIGNIFICANT',
        'Data checksum documented in README but not verified in code',
        'The README documents a checksum (SHA256/MD5) for a data file, '
        'but no code file verifies it at runtime. A validator who downloads '
        'a corrupted or wrong-version file will get silent wrong results. '
        'Adding a runtime check takes three lines and prevents silent failure.',
        [f'Checksum found in README: {readme_checksums[0]}',
         'No hashlib / checksum verification found in any code file',
         'Fix: add at the start of your analysis script:',
         '  import hashlib',
         '  with open("data/yourfile", "rb") as f:',
         '      assert hashlib.sha256(f.read()).hexdigest() == "<hash>"']
    ))
    return findings


def detect_CX_system_dependencies(repo_dir, all_files):
    """Failure Mode CX: Python packages that require system C/C++ libraries."""
    findings = []
    import re as _re
    # Collect all package names from requirements files
    req_files = [f for f in all_files
                 if _re.match(r'requirements.*\.txt$', f.name.lower())]
    env_files = [f for f in all_files
                 if f.name.lower() in {'environment.yml', 'environment.yaml'}]
    if not req_files and not env_files:
        return findings
    pkgs_found = set()
    for rf in req_files:
        try:
            for line in rf.read_text(encoding='utf-8', errors='ignore').splitlines():
                line = line.strip().lower()
                if not line or line.startswith('#'):
                    continue
                pkg = _re.split(r'[>=<!\[; ]', line)[0].strip()
                if pkg:
                    pkgs_found.add(pkg)
        except Exception:
            pass
    for ef in env_files:
        try:
            for line in ef.read_text(encoding='utf-8', errors='ignore').splitlines():
                s = line.strip().lstrip('- ').lower()
                pkg = _re.split(r'[>=<!\[; =]', s)[0].strip()
                if pkg and not pkg.startswith('#'):
                    pkgs_found.add(pkg)
        except Exception:
            pass
    # Check which found packages need system libs
    triggered = {}
    for pkg in pkgs_found:
        if pkg in _SYSTEM_DEP_PACKAGES:
            for lib in _SYSTEM_DEP_PACKAGES[pkg]:
                triggered.setdefault(lib, set()).add(pkg)
    if not triggered:
        return findings
    # Build fix instructions
    apt_pkgs = [_SYSTEM_LIB_APT[lib] for lib in sorted(triggered) if lib in _SYSTEM_LIB_APT]
    brew_pkgs = [_SYSTEM_LIB_BREW[lib] for lib in sorted(triggered) if lib in _SYSTEM_LIB_BREW]
    py_pkgs = sorted({p for ps in triggered.values() for p in ps})
    details = [
        'Packages requiring system libraries: ' + ', '.join(py_pkgs),
        'System libraries needed: ' + ', '.join(sorted(triggered.keys())),
    ]
    if apt_pkgs:
        details.append('Ubuntu/Debian: sudo apt-get install ' + ' '.join(apt_pkgs))
    if brew_pkgs:
        details.append('macOS: brew install ' + ' '.join(set(brew_pkgs)))
    details += [
        'Recommended: use conda instead of pip for these packages:',
        '  conda install -c conda-forge ' + ' '.join(py_pkgs),
        'Fix: document system library requirements in README before pip install step',
    ]
    findings.append(finding(
        'CX', 'SIGNIFICANT',
        'System-level C/C++ libraries required for ' + ', '.join(py_pkgs),
        'One or more Python packages require system-level C/C++ libraries that '
        'cannot be installed via pip alone. On a clean machine, pip install will '
        'fail with a cryptic compilation error unless these system libraries are '
        'installed first. This must be documented in the README.',
        details
    ))
    return findings


def detect_CW_reticulate_coupling(repo_dir, all_files):
    """Failure Mode CW: R script uses reticulate to call Python at runtime."""
    findings = []
    r_files = _researcher_r_files(all_files, repo_dir)
    if not r_files:
        return findings

    py_invoke_fns = ['py_run_file', 'py_run_string', 'source_python',
                     'py_load_object', 'py_save_object']
    config_pat = re.compile(
        r'(?:reticulate::)?(use_python|use_virtualenv|use_condaenv)\s*\(',
        re.IGNORECASE
    )
    py_dollar_pat = re.compile(r'py\$(\w+)')

    for f in r_files:
        try:
            src = f.read_text(encoding='utf-8', errors='ignore')
        except Exception:
            continue

        has_library = bool(re.search(r'library\s*\(\s*reticulate\s*\)', src))
        has_namespaced = 'reticulate::' in src
        if not has_library and not has_namespaced:
            continue

        evidence = []
        for fn in py_invoke_fns:
            hits = re.findall(
                r'(?:reticulate::)?' + fn + r'\s*\(\s*([^)]+)\)',
                src, re.IGNORECASE
            )
            for h in hits:
                evidence.append(fn + '(' + h.strip()[:40] + ')')
        for m in re.finditer(r'reticulate::import\s*\(\s*(\w+)', src):
            evidence.append('reticulate::import(' + m.group(1) + ')')
        py_vars = list(set(py_dollar_pat.findall(src)))[:3]
        if py_vars:
            evidence.append('Accesses Python objects via py$: ' + ', '.join(py_vars))

        has_config = bool(config_pat.search(src))
        details = ['Reticulate loaded in: ' + f.name]
        if evidence:
            details += ['Evidence: ' + e for e in evidence[:5]]
        details += [
            'Both R and Python environments must be installed before running',
            'reticulate must point to the correct Python interpreter',
        ]
        if not has_config:
            details.append(
                'No interpreter config found -- add '
                'reticulate::use_virtualenv("venv") or set RETICULATE_PYTHON env var'
            )
        details.append('Fix: document Python interpreter requirement in README')

        findings.append(finding(
            'CW', 'SIGNIFICANT',
            'reticulate coupling: R script invokes Python at runtime in ' + f.name,
            'The R script loads reticulate and calls Python code at runtime. '
            'Both R and Python environments must be correctly installed and compatible. '
            'If reticulate points to the wrong Python interpreter, or Python dependencies '
            'are missing, the pipeline fails with cryptic errors.',
            details
        ))

    return findings


def detect_CR_crlf_line_endings(repo_dir, all_files):
    """Failure Mode CR: Shell script has Windows CRLF line endings — will fail on Linux/macOS."""
    findings = []
    shell_files = [f for f in all_files if f.suffix.lower() in {'.sh', '.bash'}
                   or (f.suffix == '' and f.name.lower() in {'makefile'})]
    for f in shell_files:
        try:
            raw = f.read_bytes()
            if b'\r\n' in raw:
                findings.append(finding(
                    'CR', 'SIGNIFICANT',
                    f'Shell script has Windows CRLF line endings — will fail on Linux/macOS: {f.name}',
                    f'{f.name} contains Windows-style CRLF (\\r\\n) line endings. '
                    'On Linux/macOS, bash interprets the \\r as part of the interpreter '
                    'path, causing: /bin/bash^M: bad interpreter: No such file or directory.',
                    [f'File: {f.name} — CRLF endings detected',
                     'Fix: run dos2unix ' + f.name,
                     'Or: sed -i \'s/\\r//\' ' + f.name]
                ))
        except Exception:
            pass
    return findings













def detect_CQ_julia_pkg_add_at_runtime(repo_dir, all_files):
    """Failure Mode CQ: Julia script calls Pkg.add() at runtime with no Project.toml."""
    findings = []
    jl_files = [f for f in all_files if f.suffix.lower() == '.jl']
    if not jl_files:
        return findings
    names = {f.name.lower() for f in all_files}
    if 'project.toml' in names:
        return findings  # Project.toml exists — BY handles missing Manifest
    pkg_add_pat = re.compile(r'Pkg\.add\s*\(', re.IGNORECASE)
    affected = []
    pkgs_found = []
    for f in jl_files:
        try:
            src = f.read_text(encoding='utf-8', errors='ignore')
            if pkg_add_pat.search(src):
                affected.append(f.name)
                for m in re.finditer(r'Pkg\.add\s*\(\s*["\'](\w+)["\'\)]', src):
                    pkgs_found.append(m.group(1))
        except Exception:
            pass
    if affected:
        findings.append(finding(
            'CQ', 'SIGNIFICANT',
            f'Julia script installs packages at runtime via Pkg.add() with no Project.toml',
            'Pkg.add() without a Project.toml/Manifest.toml installs the latest package '
            'versions at the time of execution. Results cannot be reproduced if package '
            'versions change. The correct approach is to commit a Project.toml and '
            'Manifest.toml that lock exact versions.',
            [f'File: {", ".join(affected)} — Pkg.add() called at runtime'] +
            ([f'Packages: {", ".join(pkgs_found[:8])}'] if pkgs_found else []) +
            ['Fix: run `julia --project=. -e "using Pkg; Pkg.add([...]); Pkg.resolve()"` '
             'then commit Project.toml and Manifest.toml']
        ))
    return findings


def detect_CS_committed_model_binary(repo_dir, all_files):
    """Failure Mode CS: Committed model binary loaded via pickle — version-sensitive and security risk."""
    findings = []
    model_artifact_extensions = {'.pkl', '.pickle'}
    model_name_indicators = {'model', 'clf', 'classifier', 'regressor', 'estimator', 'pipeline', 'weights'}

    candidate_files = [
        f for f in all_files
        if f.suffix.lower() in model_artifact_extensions
    ]
    if not candidate_files:
        return findings

    code_content = ''
    py_files = [f for f in all_files if f.suffix.lower() in CODE_EXTENSIONS]
    for f in py_files:
        code_content += read_file_safe(f)

    has_pickle_load = bool(re.search(r'pickle\.load\s*\(', code_content))

    model_files = []
    for f in candidate_files:
        name_lower = f.name.lower()
        in_model_dir = any(part.lower() in {'models', 'model', 'checkpoints'} for part in f.parts)
        has_model_name = any(ind in name_lower for ind in model_name_indicators)
        if has_model_name or in_model_dir:
            model_files.append(f)

    if not model_files and has_pickle_load:
        model_files = candidate_files

    if not model_files or not has_pickle_load:
        return findings

    details = [f'Model file: {", ".join(f.name for f in model_files[:5])}',
               'pickle.load() executes arbitrary code — validators loading this file run untrusted code',
               'Pickle files are version-specific: model trained under one scikit-learn / torch version '
               'may fail silently or produce different results under a different version',
               'Fix: replace with a portable format (safetensors, ONNX) and/or host on HuggingFace Hub']
    findings.append(finding(
        'CS', 'SIGNIFICANT',
        'Committed model binary loaded via pickle — version-sensitive and security risk',
        'A trained model binary is committed to the repository and loaded with pickle.load(). '
        'This creates two reproducibility risks: (1) pickle files encode version information — '
        'a model serialised under one library version may silently produce wrong results under '
        'another; (2) pickle.load() executes arbitrary code, meaning validators are running '
        'untrusted code when loading the file. The fix is to use a portable, safe format '
        '(safetensors for neural networks, ONNX for cross-framework interoperability) or to '
        'host the model on HuggingFace Hub and load it with a version-pinned API call.',
        details
    ))
    return findings


def detect_CU_conda_unpinned_packages(repo_dir, all_files):
    """Failure Mode CU: environment.yml has unpinned or loosely-pinned conda packages."""
    findings = []
    env_file = next((f for f in all_files if f.name.lower() in {'environment.yml', 'environment.yaml'}), None)
    if not env_file:
        return findings
    try:
        src = env_file.read_text(encoding='utf-8', errors='ignore')
    except Exception:
        return findings

    unpinned = []
    loose = []
    in_deps = False
    for line in src.splitlines():
        stripped = line.strip()
        if stripped.startswith('#') or not stripped:
            continue
        if stripped.startswith('dependencies:'):
            in_deps = True
            continue
        if in_deps and stripped.startswith('-') and ':' in stripped and not stripped.startswith('- pip'):
            # sub-key like 'name:', 'channels:' — end of deps
            in_deps = False
        if not in_deps:
            continue
        if stripped.startswith('- pip:') or stripped == '- pip':
            continue
        if stripped.startswith('-'):
            pkg = stripped.lstrip('- ').strip()
            if pkg.startswith('{') or not pkg:
                continue
            # exact pin: pkg=X.Y.Z or pkg=X.Y.Z=pyXXX
            if re.match(r'^[\w\-\.]+=[0-9]', pkg):
                continue  # pinned
            # loose: pkg>=X or pkg>X or no version at all
            if re.match(r'^[\w\-\.]+\s*[><!]', pkg):
                loose.append(pkg)
            elif '=' not in pkg and re.match(r'^[\w\-\.]+$', pkg):
                unpinned.append(pkg)
            elif re.match(r'^[\w\-\.]+$', pkg):
                unpinned.append(pkg)

    all_issues = unpinned + loose
    if all_issues:
        sample = all_issues[:8]
        details = [
            f'Unpinned/loosely-pinned packages ({len(all_issues)}): {", ".join(sample)}'
            + (f' (and {len(all_issues)-8} more)' if len(all_issues) > 8 else ''),
            'Conda uses = for exact pinning: numpy=1.24.3, not numpy>=1.24',
            'Fix: run `conda env export --no-builds > environment.yml` in your original environment',
        ]
        if loose:
            details.insert(1, f'Loose constraints (will install different versions over time): {", ".join(loose[:5])}')
        findings.append(finding(
            'CU', 'SIGNIFICANT',
            f'environment.yml has {len(all_issues)} unpinned or loosely-pinned package(s)',
            'Conda packages without exact version pins will install the latest available '
            'version at the time of environment creation. This means the environment created '
            'by validators may differ from the one used to produce the original results. '
            'Use `conda env export --no-builds` to capture exact versions.',
            details
        ))
    return findings


def detect_CV_hardcoded_cuda_no_fallback(repo_dir, all_files):
    """Failure Mode CV: Code uses torch.device('cuda') with no CPU fallback."""
    findings = []
    py_files = [f for f in all_files if f.suffix.lower() in {'.py', '.ipynb'}]
    if not py_files:
        return findings

    import json as _json
    cuda_pat = re.compile(r'torch\.device\s*\(\s*["\']cuda["\']\s*\)', re.IGNORECASE)
    fallback_pat = re.compile(r'cuda\.is_available\s*\(\)', re.IGNORECASE)

    affected = []
    for f in py_files:
        try:
            if f.suffix.lower() == '.ipynb':
                nb = _json.loads(f.read_text(encoding='utf-8', errors='ignore'))
                src = '\n'.join(
                    ''.join(cell.get('source', []))
                    for cell in nb.get('cells', [])
                    if cell.get('cell_type') == 'code'
                )
            else:
                src = f.read_text(encoding='utf-8', errors='ignore')
            if cuda_pat.search(src) and not fallback_pat.search(src):
                affected.append(f.name)
        except Exception:
            pass

    if affected:
        findings.append(finding(
            'CV', 'SIGNIFICANT',
            f'Hardcoded CUDA device with no CPU fallback in {len(affected)} file(s)',
            'Code calls torch.device("cuda") without checking torch.cuda.is_available(). '
            'This will crash immediately on any machine without a CUDA-capable GPU with: '
            'AssertionError: Torch not compiled with CUDA enabled. '
            'Most validators will not have access to the same GPU hardware. '
            'Fix: replace with device = torch.device("cuda" if torch.cuda.is_available() else "cpu")',
            [f'Affected files: {", ".join(affected)}',
             'Fix: device = torch.device("cuda" if torch.cuda.is_available() else "cpu")',
             'Also document GPU requirement in README System Requirements section']
        ))
    return findings

def detect_CP_python2_syntax(repo_dir, all_files):
    """Failure Mode CP: Python 2 syntax in Python 3 repository."""
    findings = []
    py_files = [f for f in all_files if f.suffix.lower() in {'.py', '.ipynb'}]
    if not py_files:
        return findings
    # Python 2 print statement: print "..." or print var, var
    print_stmt = re.compile(r'^[ \t]*print\s+["\'\w]', re.MULTILINE)
    # Python 2 exec statement
    exec_stmt = re.compile(r'^[ \t]*exec\s+["\'\w]', re.MULTILINE)
    # Python 2 raise: raise Exception, "msg"
    raise_stmt = re.compile(r'raise\s+\w+\s*,\s*["\'\w]')
    # Python 2 unicode/basestring/xrange builtins
    py2_builtins = re.compile(r'\b(unicode|basestring|xrange|raw_input|reduce|execfile|reload)\s*\(')
    # Python 2 integer division note (silent wrong results in Python 2)
    patterns = [
        (print_stmt, 'Python 2 print statement (SyntaxError in Python 3)'),
        (exec_stmt, 'Python 2 exec statement (SyntaxError in Python 3)'),
        (raise_stmt, 'Python 2 raise syntax (SyntaxError in Python 3)'),
        (py2_builtins, 'Python 2 builtin function'),
    ]
    evidence = []
    affected_files = set()
    for f in py_files:
        try:
            src = f.read_text(encoding='utf-8', errors='ignore')
            for pat, desc in patterns:
                if pat.search(src):
                    evidence.append(f'{f.name}: {desc}')
                    affected_files.add(f.name)
        except Exception:
            pass
    if evidence:
        findings.append(finding(
            'CP', 'SIGNIFICANT',
            f'Python 2 syntax detected in {len(affected_files)} file(s) — will fail in Python 3',
            'The repository contains Python 2 syntax that raises SyntaxError in Python 3. '
            'The code cannot run under the stated Python version. '
            'A validator will encounter an immediate error before any logic executes.',
            evidence[:5] + (['Fix: convert to Python 3 syntax (use print(), raise Exception("msg"), etc.)']
                           if evidence else [])
        ))
    return findings

def detect_CO_matlab_undocumented_functions(repo_dir, all_files):
    """Failure Mode CO: MATLAB code uses undocumented internal functions."""
    findings = []
    m_files = [f for f in all_files if f.suffix.lower() == '.m']
    if not m_files:
        return findings
    # Patterns for undocumented/internal MATLAB usage
    internal_patterns = [
        (re.compile(r'matlab\.internal\.', re.IGNORECASE),
         'matlab.internal.* — undocumented internal namespace, may change without notice'),
        (re.compile(r'matlab\.lang\.internal\.', re.IGNORECASE),
         'matlab.lang.internal.* — undocumented internal namespace'),
        (re.compile(r'\bundocumented\b', re.IGNORECASE),
         'comment flags undocumented function usage'),
    ]
    # gradient() on image data is undocumented — imgradient() is the documented alternative
    gradient_img_pat = re.compile(
        r'\[\s*\w+\s*,\s*\w+\s*\]\s*=\s*gradient\s*\(\s*(?:double|single|uint8|uint16|img|image|im|stack|frame)',
        re.IGNORECASE
    )
    evidence_all = []
    for f in m_files:
        try:
            src = f.read_text(encoding='utf-8', errors='ignore')
            for pat, desc in internal_patterns:
                if pat.search(src):
                    evidence_all.append(f'{f.name}: {desc}')
            if gradient_img_pat.search(src):
                evidence_all.append(
                    f'{f.name}: gradient() used on image data — '
                    f'undocumented for this use case; use imgradient() instead'
                )
        except Exception:
            pass
    if evidence_all:
        findings.append(finding(
            'CO', 'SIGNIFICANT',
            f'MATLAB code uses undocumented or internal functions ({len(evidence_all)} instance(s))',
            'Undocumented MATLAB internal functions may change behaviour or be removed '
            'between MATLAB releases without notice. Code depending on these functions '
            'may produce different results or errors on different MATLAB versions.',
            evidence_all[:5] + ['Fix: replace internal/undocumented functions with '
                                 'documented equivalents (see MATLAB documentation)']
        ))
    return findings

def detect_CN_known_version_conflicts(repo_dir, all_files):
    """Failure Mode CN: requirements.txt contains known incompatible package version combinations."""
    findings = []
    req_files = [f for f in all_files if re.match(r'requirements.*\.txt$', f.name.lower())]
    if not req_files:
        return findings
    # Build pinned versions dict from all requirements files
    pinned = {}
    for rf in req_files:
        try:
            for line in rf.read_text(encoding='utf-8', errors='ignore').splitlines():
                line = line.strip()
                m = re.match(r'^([\w.-]+)==([\d.]+)', line)
                if m:
                    pinned[m.group(1).lower()] = m.group(2)
        except Exception:
            pass
    if not pinned:
        return findings
    # Known incompatible combinations: (pkg, version_pred, conflicting_pkg, conflict_desc)
    # Format: (package, max_version_exclusive, requires_pkg, constraint_desc)
    known_conflicts = [
        # tensorflow < 2.13 requires numpy < 1.24
        ('tensorflow', lambda v: tuple(int(x) for x in v.split('.')[:2]) < (2, 13),
         'numpy', lambda v: tuple(int(x) for x in v.split('.')[:2]) >= (1, 24),
         'tensorflow<2.13 requires numpy<1.24'),
        # tensorflow-gpu same constraint
        ('tensorflow-gpu', lambda v: tuple(int(x) for x in v.split('.')[:2]) < (2, 13),
         'numpy', lambda v: tuple(int(x) for x in v.split('.')[:2]) >= (1, 24),
         'tensorflow-gpu<2.13 requires numpy<1.24'),
        # torch < 2.0 and numpy >= 2.0 incompatible
        ('torch', lambda v: tuple(int(x) for x in v.split('.')[:2]) < (2, 0),
         'numpy', lambda v: tuple(int(x) for x in v.split('.')[:1]) >= (2,),
         'torch<2.0 incompatible with numpy>=2.0'),
        # scipy < 1.9 requires numpy < 1.25
        ('scipy', lambda v: tuple(int(x) for x in v.split('.')[:2]) < (1, 9),
         'numpy', lambda v: tuple(int(x) for x in v.split('.')[:2]) >= (1, 25),
         'scipy<1.9 requires numpy<1.25'),
    ]
    conflicts = []
    for pkg, pkg_pred, dep_pkg, dep_pred, desc in known_conflicts:
        try:
            if pkg in pinned and dep_pkg in pinned:
                if pkg_pred(pinned[pkg]) and dep_pred(pinned[dep_pkg]):
                    conflicts.append(
                        f'{pkg}=={pinned[pkg]} conflicts with {dep_pkg}=={pinned[dep_pkg]}: {desc}'
                    )
        except Exception:
            pass
    if conflicts:
        findings.append(finding(
            'CN', 'SIGNIFICANT',
            f'{len(conflicts)} known package version conflict(s) in requirements',
            'The pinned versions contain known incompatible combinations. '
            'pip install will fail with a dependency resolution error.',
            [f'Conflict: {c}' for c in conflicts] +
            ['Fix: update package versions to compatible combinations']
        ))
    return findings

def detect_CM_nextflow_no_container(repo_dir, all_files):
    """Failure Mode CM: Nextflow processes lack container or conda directives."""
    findings = []
    nf_files = [f for f in all_files if f.suffix.lower() == '.nf']
    if not nf_files:
        return findings
    process_pat = re.compile(r'process\s+(\w+)\s*\{([^}]+(?:\{[^}]*\}[^}]*)*)\}', re.DOTALL)
    container_pat = re.compile(r'^\s*(container|conda)\s+', re.MULTILINE)
    bare_processes = []
    for f in nf_files:
        try:
            src = f.read_text(encoding='utf-8', errors='ignore')
            for m in process_pat.finditer(src):
                pname = m.group(1)
                pbody = m.group(2)
                if not container_pat.search(pbody):
                    bare_processes.append(f'{f.name}: process {pname}')
        except Exception:
            pass
    if bare_processes:
        findings.append(finding(
            'CM', 'SIGNIFICANT',
            f'{len(bare_processes)} Nextflow process(es) have no container or conda directive',
            'Without container or conda directives, tool versions depend entirely on '
            'whatever software happens to be installed on the host machine. '
            'A validator on a different system will get different versions and '
            'potentially different results.',
            [f'Process without container/conda: {p}' for p in bare_processes[:5]] +
            ['Fix: add container \'docker://...\'  or conda \'...\'  to each process, '
             'or set process.container globally in nextflow.config']
        ))
    return findings

def detect_CL_bioconductor_unpinned(repo_dir, all_files):
    """Failure Mode CL: BiocManager::install() called without version= argument."""
    findings = []
    r_files = _researcher_r_files(all_files, repo_dir)
    bioc_pat = re.compile(r'BiocManager::install\s*\(', re.IGNORECASE)
    version_pat = re.compile(r'version\s*=', re.IGNORECASE)
    # Extract stated Bioconductor version from README
    stated_version = None
    for f in all_files:
        if f.name.lower() in {'readme.md', 'readme.txt', 'readme.rst'}:
            readme = read_file_safe(f)
            m = re.search(r'Bioconductor\s+(\d+\.\d+)', readme, re.IGNORECASE)
            if m:
                stated_version = m.group(1)
                break
    for f in r_files:
        try:
            src = f.read_text(encoding='utf-8', errors='ignore')
            unversioned = []
            for m in bioc_pat.finditer(src):
                # Find the full call using paren depth
                start = m.start()
                depth = 0
                call_end = start
                for ci, ch in enumerate(src[start:]):
                    if ch == '(': depth += 1
                    elif ch == ')':
                        depth -= 1
                        if depth == 0:
                            call_end = start + ci
                            break
                call_text = src[start:call_end+1]
                if not version_pat.search(call_text):
                    unversioned.append(call_text[:60].replace('\n', ' '))
            if unversioned:
                count = len(unversioned)
                evidence = [
                    f'File: {f.name} — {count} BiocManager::install() call(s) without version=',
                    'Without version=, installs current Bioconductor release (may differ from original)',
                ]
                if stated_version:
                    evidence.append(f'README states Bioconductor {stated_version} — add version="{stated_version}" to enforce this')
                    evidence.append(f'Fix: BiocManager::install(c(...), version="{stated_version}")')
                else:
                    evidence.append('Fix: add version="X.YY" matching the Bioconductor release used')
                findings.append(finding(
                    'CL', 'SIGNIFICANT',
                    f'Bioconductor packages installed without version pin in {f.name}',
                    'BiocManager::install() without version= installs the current Bioconductor '
                    'release, not the one used in the original analysis. Package APIs and '
                    'default parameters change between releases.',
                    evidence
                ))
        except Exception:
            pass
    return findings

def detect_CK_conflicting_readmes(repo_dir, all_files):
    """Failure Mode CK: Multiple README files with conflicting instructions."""
    findings = []
    readme_files = [f for f in all_files
                    if f.name.lower() in {'readme.md', 'readme.txt', 'readme.rst'}]
    if len(readme_files) < 2:
        return findings
    # Extract Python version references from each README
    ver_pat = re.compile(r'python\s+(\d+\.\d+\+?)', re.IGNORECASE)
    conda_pat = re.compile(r'conda\s+env\s+create|conda\s+install', re.IGNORECASE)
    pip_pat = re.compile(r'pip\s+install', re.IGNORECASE)
    run_pat = re.compile(r'(?:python|Rscript)\s+([\w./\-]+\.(?:py|r)(?:\s+--\S+)*)', re.IGNORECASE)
    draft_pat = re.compile(r'draft|outdated|see\s+\S+readme|for\s+latest', re.IGNORECASE)
    readme_info = {}
    for f in readme_files:
        try:
            src = f.read_text(encoding='utf-8', errors='ignore')
            rel = str(f.relative_to(repo_dir)).replace('\\', '/')
            info = {}
            versions = ver_pat.findall(src)
            if versions:
                info['python'] = versions[0]
            if conda_pat.search(src):
                info['install'] = 'conda'
            elif pip_pat.search(src):
                info['install'] = 'pip'
            runs = run_pat.findall(src)
            if runs:
                info['run'] = runs[0].strip()
            if draft_pat.search(src):
                info['draft'] = True
            readme_info[rel] = info
        except Exception:
            pass
    if len(readme_info) < 2:
        return findings
    # Check for conflicts
    conflicts = []
    all_versions = {k: v['python'] for k, v in readme_info.items() if 'python' in v}
    all_installs = {k: v['install'] for k, v in readme_info.items() if 'install' in v}
    all_runs = {k: v['run'] for k, v in readme_info.items() if 'run' in v}
    draft_files = [k for k, v in readme_info.items() if v.get('draft')]

    # If no README contains any code instructions (Python version, install method,
    # or run command), the READMEs are data-only documentation for multiple dataset
    # versions or snapshots — not conflicting execution guides.  Emit a soft
    # informational note instead of a SIGNIFICANT conflict finding.
    _all_data_only = not any(
        v.get('python') or v.get('install') or v.get('run')
        for v in readme_info.values()
    )
    if _all_data_only:
        findings.append(finding(
            'CK', 'LOW CONFIDENCE',
            f'Multiple README files found ({len(readme_files)}) — confirm which is current',
            'Multiple README files were detected but none contains code execution instructions, '
            'suggesting these are data-documentation READMEs for different dataset versions or '
            'snapshots rather than conflicting setup guides.',
            [f'{len(readme_files)} README files: ' + ', '.join(readme_info.keys()),
             'Recommendation: ensure the root README clarifies which dataset version is '
             'current and how the versions relate to each other']
        ))
        return findings

    if len(set(all_versions.values())) > 1:
        conflicts.append('Python version: ' + ' / '.join(f'{k} says {v}' for k, v in all_versions.items()))
    if len(set(all_installs.values())) > 1:
        conflicts.append('Installation method: ' + ' / '.join(f'{k} says {v}' for k, v in all_installs.items()))
    if len(set(all_runs.values())) > 1:
        conflicts.append('Run command: ' + ' / '.join(f'{k} says {v}' for k, v in all_runs.items()))
    if conflicts or len(readme_files) > 1:
        evidence = [f'{len(readme_files)} README files found: ' + ', '.join(readme_info.keys())]
        evidence += conflicts[:4]
        if draft_files:
            evidence.append('Note: ' + ', '.join(draft_files) + ' marked as DRAFT or outdated')
        evidence.append('Fix: if this is a monorepo, ensure root README clearly links to each sub-project README; otherwise consolidate into a single README.md')
        findings.append(finding(
            'CK', 'SIGNIFICANT',
            f'Multiple README files found ({len(readme_files)}) — instructions may conflict',
            'Multiple README files were detected. Conflicting instructions on Python version, '
            'installation method, or run commands will cause validators to follow incorrect steps.',
            evidence[:7]
        ))
    return findings

def detect_CJ_readme_references_missing_files(repo_dir, all_files):
    """Failure Mode CJ: README references config/environment files not in repository."""
    findings = []
    all_names = {f.name.lower() for f in all_files}
    # Scan all READMEs including subdirectory ones
    readme_files = [f for f in all_files
                    if f.name.lower() in {'readme.md', 'readme.txt', 'readme.rst'}]
    if not readme_files:
        return findings
    content = '\n'.join(read_file_safe(f) or '' for f in readme_files)
    # Patterns for referenced files that should exist
    ref_pattern = re.compile(
        r'(?:conda\s+env\s+create\s+-f|'
        r'--config|'
        r'-f\s+|'
        r'conda\s+activate\s+\S+\s+-f\s+)'
        r'([\w./\-]+\.(?:ya?ml|yaml|json|cfg|ini|toml|txt))\b',
        re.IGNORECASE
    )
    # Also catch bare filename references like "config.yaml" near run instructions
    bare_config = re.compile(
        r'(?:--config|config=|-f)\s+([\w./\-]+\.(?:ya?ml|json|cfg|ini|toml))\b',
        re.IGNORECASE
    )
    missing = []
    for pat in [ref_pattern, bare_config]:
        for m in pat.finditer(content):
            ref = m.group(1)
            ref_name = ref.replace('\\', '/').split('/')[-1].lower()
            if ref_name not in all_names and ref not in all_names:
                if ref_name not in [r.split('/')[-1] for r in missing]:
                    missing.append(ref)
    if missing:
        findings.append(finding(
            'CJ', 'SIGNIFICANT',
            f'README references {len(missing)} file(s) not found in repository',
            'The README instructions reference files that do not exist in the repository. '
            'A validator following these instructions will immediately encounter errors.',
            [f'Missing: {m}' for m in missing[:5]] +
            ['Fix: add the missing files or update the README instructions']
        ))
    return findings

def detect_CI_live_data_no_archive(repo_dir, all_files):
    """Failure Mode CI: Code fetches live data at runtime with no local archived copy."""
    findings = []
    code_files = [f for f in all_files if f.suffix.lower() in CODE_EXTENSIONS]
    api_pat = re.compile(
        r'requests\.(get|post)\s*\(\s*["\']https?://'
        r'|urllib.*urlopen\s*\(\s*["\']https?://'
        r'|pd\.read_csv\s*\(\s*["\']https?://',
        re.IGNORECASE
    )
    branch_url_pat = re.compile(
        r'raw\.githubusercontent\.com/[^/]+/[^/]+/(main|master|dev|develop|HEAD)/\S+',
        re.IGNORECASE
    )
    save_pat = re.compile(
        r'to_csv|to_parquet|to_excel|pickle\.dump|np\.save',
        re.IGNORECASE
    )
    url_extract = re.compile(r'https?://[^\s"\'>\)]+')
    for f in code_files:
        try:
            src = f.read_text(encoding='utf-8', errors='ignore')
            if not api_pat.search(src):
                continue
            if save_pat.search(src):
                continue
            evidence = []
            for url in url_extract.findall(src):
                bm = branch_url_pat.search(url)
                if bm:
                    evidence.append('GitHub raw URL @' + bm.group(1) + ' (branch not pinned): ' + url[:80])
                elif url.startswith('http'):
                    evidence.append('Live URL (no local snapshot): ' + url[:80])
            if not evidence:
                evidence.append('Network fetch in ' + f.name + ' -- no local data save found')
            findings.append(finding(
                'CI', 'SIGNIFICANT',
                'Live data fetched at runtime with no archived copy: ' + f.name,
                'Code fetches data from external sources but saves no local snapshot. '
                'Results cannot be reproduced if remote data changes.',
                evidence[:5] + ['Fix: save fetched data to data/ and read locally, '
                                'or document exact snapshot date and version']
            ))
        except Exception:
            pass
    # Authenticated cloud API check — separate from live URL check
    auth_apis = [
        # Google Earth Engine
        (re.compile(r'ee\.(Authenticate|Initialize)\s*\(', re.IGNORECASE),
         'earthengine-api', 'Google Earth Engine',
         'GEE access requires registration at earthengine.google.com — approval is not instant.',
         'ee.Authenticate() / ee.Initialize() detected'),
        # AWS boto3
        (re.compile(r'boto3\.(client|resource|session)\s*\(', re.IGNORECASE),
         'boto3', 'AWS (boto3)',
         'AWS credentials must be configured (aws configure or IAM role).',
         'boto3.client() / boto3.resource() detected'),
        # Google Cloud Storage / BigQuery
        (re.compile(r'(google\.cloud\.(storage|bigquery)|bigquery\.Client|storage\.Client)\s*\(', re.IGNORECASE),
         'google-cloud', 'Google Cloud',
         'GCP credentials must be configured (gcloud auth application-default login).',
         'google.cloud client detected'),
        # Azure
        (re.compile(r'(BlobServiceClient|AzureCliCredential|DefaultAzureCredential)\s*\(', re.IGNORECASE),
         'azure', 'Azure',
         'Azure credentials must be configured (az login).',
         'Azure SDK client detected'),
        # Copernicus / sentinelsat
        (re.compile(r'(SentinelAPI|sentinelsat)\s*\(', re.IGNORECASE),
         'sentinelsat', 'Copernicus/Sentinel Hub',
         'Copernicus account required at scihub.copernicus.eu.',
         'SentinelAPI() detected'),
    ]
    for pat, pkg_hint, api_name, auth_note, evidence_label in auth_apis:
        # Check if pkg is in requirements (optional signal) or pattern found in code
        for f in code_files:
            try:
                src = f.read_text(encoding='utf-8', errors='ignore')
                if not pat.search(src):
                    continue
                findings.append(finding(
                    'CI', 'SIGNIFICANT',
                    f'Authenticated cloud API required: {api_name} in {f.name}',
                    f'Code uses {api_name} which requires account registration, '
                    f'authentication credentials, and potentially approved project access. '
                    f'Validators cannot run this code without setting up credentials. '
                    f'{auth_note}',
                    [f'Evidence: {evidence_label} in {f.name}',
                     f'Fix: document {api_name} account requirement in README with step-by-step '
                     f'authentication instructions',
                     'Consider providing a local data export as a fallback for validators '
                     'who cannot obtain credentials']
                ))
                break  # one finding per API type
            except Exception:
                pass
    return findings


def detect_CH_broken_source_chain(repo_dir, all_files):
    """Failure Mode CH: R source() calls reference files not in the repository."""
    findings = []
    r_files = _researcher_r_files(all_files, repo_dir)
    all_r_names = {f.name.lower() for f in r_files}
    all_r_paths = {str(f.relative_to(repo_dir)).replace('\\', '/').lower() for f in r_files}
    source_pat = re.compile(r'source\s*\(\s*["\']([^"\']+)["\']', re.IGNORECASE)
    missing_sources = []
    for f in r_files:
        try:
            src = f.read_text(encoding='utf-8', errors='ignore')
            for m in source_pat.finditer(src):
                ref = m.group(1).strip()
                ref_name = ref.replace('\\', '/').split('/')[-1].lower()
                ref_norm = ref.replace('\\', '/').lower().lstrip('./')
                if (ref_name not in all_r_names and
                        ref_norm not in all_r_paths and
                        not ref.startswith('http')):
                    missing_sources.append(f'{f.name}: source("{ref}")')
        except Exception:
            pass
    if missing_sources:
        findings.append(finding(
            'CH', 'SIGNIFICANT',
            f'R source() chain references {len(missing_sources)} missing file(s)',
            'source() calls reference R scripts that are not present in the repository. '
            'The pipeline will fail immediately when these files are loaded.',
            [f'Missing: {s}' for s in missing_sources[:5]] +
            ['Fix: add the missing R script(s) to the repository or remove the source() call']
        ))
    return findings

def detect_CG_unpinned_git_requirements(repo_dir, all_files):
    """Failure Mode CG: requirements files contain git+ URLs or unpinned constraints."""
    findings = []
    req_files = [f for f in all_files
                 if re.match(r'requirements.*\.txt$', f.name.lower())]
    unpinned_git = []
    branch_pinned = []
    loose_constraints = []
    for f in req_files:
        try:
            lines = f.read_text(encoding='utf-8', errors='ignore').splitlines()
            for line in lines:
                line = line.strip()
                if not line or line.startswith('#') or line.startswith('-r'):
                    continue
                if line.startswith('git+'):
                    # Check for commit SHA (40 hex chars) or no pin at all
                    at_idx = line.rfind('@')
                    if at_idx == -1 or at_idx == line.index('git+'):
                        unpinned_git.append(f'{f.name}: {line[:80]}')
                    else:
                        ref = line[at_idx+1:].split('#')[0].strip()
                        # Branch names are not reproducible pins
                        if re.match(r'^(main|master|dev|develop|HEAD|latest)$', ref, re.IGNORECASE):
                            branch_pinned.append(f'{f.name}: {line[:80]} (@{ref} is a branch, not a commit)')
                        elif not re.match(r'^[0-9a-f]{7,40}$', ref):
                            branch_pinned.append(f'{f.name}: {line[:80]} (@{ref} may not be a commit SHA)')
                elif re.match(r'^[\w.-]+\s*[><!]=', line) and '==' not in line:
                    loose_constraints.append(f'{f.name}: {line[:80]}')
        except Exception:
            pass
    evidence = []
    if unpinned_git:
        evidence += ['git+ URL with no pin (installs HEAD):'] + [f'  {e}' for e in unpinned_git[:3]]
    if branch_pinned:
        evidence += ['git+ URL pinned to branch (not reproducible):'] + [f'  {e}' for e in branch_pinned[:3]]
    if loose_constraints:
        evidence += ['Loose version constraint (not a reproducible pin):'] + [f'  {e}' for e in loose_constraints[:3]]
    if evidence:
        evidence.append('Fix: pin git+ URLs to a specific commit SHA, e.g. git+https://...@a1b2c3d')
        findings.append(finding(
            'CG', 'SIGNIFICANT',
            'requirements file contains unpinned git+ URLs or loose version constraints',
            'git+ URLs without a commit SHA always install the current HEAD — a '
            'different version than what was used in the original analysis. '
            'Branch references (@main, @master) are not reproducible pins.',
            evidence[:8]
        ))
    return findings

def detect_CF_notebook_outputs_committed(repo_dir, all_files):
    """Failure Mode CF: Jupyter notebook has committed cell outputs — may contain sensitive data or large blobs."""
    findings = []
    import json as _json
    notebooks = [f for f in all_files if f.suffix.lower() == '.ipynb']
    for nb in notebooks:
        try:
            data = _json.loads(nb.read_text(encoding='utf-8', errors='ignore'))
            cells = data.get('cells', [])
            output_cells = []
            large_output = False
            for i, cell in enumerate(cells):
                if cell.get('cell_type') == 'code':
                    outputs = cell.get('outputs', [])
                    if outputs:
                        output_cells.append(i + 1)
                        for out in outputs:
                            # Check for embedded images (large base64 blobs)
                            data_block = out.get('data', {})
                            if 'image/png' in data_block or 'image/jpeg' in data_block:
                                large_output = True
            if output_cells:
                findings.append(finding(
                    'CF', 'LOW CONFIDENCE',
                    f'Notebook has committed cell outputs: {nb.name}',
                    'Cell outputs are embedded in the notebook file. This inflates '
                    'repository size, may contain sensitive data (file paths, user info '
                    'in tracebacks), and makes diffs unreadable. Best practice is to '
                    'strip outputs before committing and regenerate by running the notebook.',
                    [f'Cells with outputs: {len(output_cells)} cells',
                     'Contains embedded images: ' + ('Yes' if large_output else 'No'),
                     'Fix: jupyter nbconvert --ClearOutputPreprocessor.enabled=True '
                     '--to notebook --inplace ' + nb.name]
                ))
        except Exception:
            continue
    return findings

def detect_CE_unpinned_github_packages(repo_dir, all_files):
    """Failure Mode CE: devtools::install_github() calls without commit/tag pin."""
    findings = []
    r_files = _researcher_r_files(all_files, repo_dir)
    github_pattern = re.compile(
        r'(?:devtools|remotes)::install_github\s*\(\s*["\'][^"\']+/([\w.-]+)["\'][^)]*\)',
        re.IGNORECASE
    )
    unpinned = []
    for f in r_files:
        try:
            src = f.read_text(encoding='utf-8', errors='ignore')
            for m in github_pattern.finditer(src):
                pkg = m.group(1)
                # Pinned if @ present (commit sha or tag)
                if '@' not in pkg:
                    unpinned.append(pkg)
        except Exception:
            pass
    if unpinned:
        findings.append(finding(
            'CE', 'SIGNIFICANT',
            f'GitHub R packages installed without commit or version pin: {", ".join(unpinned[:4])}',
            'devtools::install_github() calls found with no @commit or @tag specified. '
            'These will always install the current HEAD — a different version than '
            'what was used in the original analysis. Results may not be reproducible.',
            [f'Unpinned: {p}' for p in unpinned[:5]] +
            ['Fix: pin each call, e.g. install_github("YuLab-SMU/ggtree@a1b2c3d")',
             'Or use renv to lock all package versions including GitHub sources']
        ))
    return findings

def detect_CD_dockerfile_run_before_copy(repo_dir, all_files):
    """Failure Mode CD: Dockerfile has RUN pip install before COPY — build will fail."""
    findings = []
    dockerfiles = [f for f in all_files if f.name.lower() == 'dockerfile']
    for df in dockerfiles:
        try:
            raw = df.read_text(encoding='utf-8', errors='ignore')
        except Exception:
            continue
        if not raw:
            continue
        # Strip comment lines and blank lines for analysis but keep original for evidence
        orig_lines = raw.splitlines()
        # Find index of first COPY or ADD instruction
        first_copy_idx = None
        for i, line in enumerate(orig_lines):
            s = line.strip()
            if s.startswith('#') or not s:
                continue
            if s.upper().startswith('COPY') or s.upper().startswith('ADD '):
                first_copy_idx = i
                break
        if first_copy_idx is None:
            continue
        # Check if any pip/conda install RUN appears before first COPY
        for i, line in enumerate(orig_lines):
            if i >= first_copy_idx:
                break
            s = line.strip()
            if s.startswith('#') or not s:
                continue
            su = s.upper()
            if su.startswith('RUN') and ('PIP INSTALL' in su or 'CONDA INSTALL' in su or 'PIP3 INSTALL' in su):
                findings.append(finding(
                    'CD', 'SIGNIFICANT',
                    'Dockerfile has RUN pip install before COPY — build will fail',
                    f'The RUN pip install command on line {i+1} executes before '
                    f'the COPY instruction on line {first_copy_idx+1}. '
                    'The requirements file does not yet exist in the container '
                    'at build time, causing an immediate build failure.',
                    [f'Line {i+1}: {orig_lines[i].strip()}',
                     f'Line {first_copy_idx+1}: {orig_lines[first_copy_idx].strip()}',
                     'Fix: add "COPY requirements.txt ." before the RUN pip install line']
                ))
                break
    return findings

def detect_CB_snakemake_no_env_isolation(repo_dir, all_files):
    """Failure Mode CB: Snakemake workflow has no per-rule environment isolation."""
    findings = []
    snake_files = [f for f in all_files
                   if f.name == 'Snakefile' or f.suffix.lower() == '.smk']
    if not snake_files:
        return findings
    for f in snake_files:
        content = read_file_safe(f)
        if not content:
            continue
        # Count rules
        rules = re.findall(r'^rule\s+\w+', content, re.MULTILINE)
        if not rules:
            continue
        has_conda = 'conda:' in content
        has_container = 'container:' in content or 'singularity:' in content
        if not has_conda and not has_container:
            findings.append(finding(
                'CB', 'SIGNIFICANT',
                f'Snakemake workflow has no per-rule environment isolation: {f.name}',
                f'No rule in {f.name} has a conda: or container: directive. '
                'Without these, the workflow depends on tools being available '
                'on PATH with no version control. Different tool versions '
                'will produce different results.',
                [f'Rules found: {", ".join(r.split()[1] for r in rules[:5])}',
                 'Fix: add conda: directives with environment YAML files to each rule,',
                 'or use container: with a Docker/Singularity image']
            ))
    return findings


def detect_CC_undocumented_external_tools(repo_dir, all_files):
    """Failure Mode CC: README mentions external tools on PATH with no version specified."""
    findings = []
    # Scan README and shell scripts for tool references
    scan_files = [f for f in all_files
                  if f.name.lower() in {'readme.md', 'readme.txt', 'readme.rst'}
                  and len(f.relative_to(repo_dir).parts) <= 4]
    scan_files += [f for f in all_files if f.suffix.lower() in {'.sh', '.bash', '.nf', '.smk'}]
    scan_files += [f for f in all_files if f.name == 'Snakefile']
    if not scan_files:
        return findings
    content = '\n'.join(read_file_safe(f) or '' for f in scan_files)
    if not content:
        return findings
    # Common bioinformatics/scientific CLI tools
    tool_pattern = re.compile(
        r'\b(bwa|samtools|gatk|bcftools|bowtie2|hisat2|star|kallisto|salmon'
        r'|bedtools|picard|trimmomatic|fastqc|multiqc|varscan|snpeff'
        r'|minimap2|blastn|blastp|makeblastdb|cellranger|seqkit'
        r'|trim_galore|featurecounts|subread|rsem|deseq2|edger|bismark'
        r'|bamtools|deeptools|macs2|homer|stringtie|cufflinks|kraken2|bracken|nextflow|snakemake'
        r'|trim_galore|trimmomatic|fastqc|star|hisat2|bowtie2|bwa|samtools|picard'
        r'|featurecounts|htseq|kallisto|salmon|cellranger|bismark|gatk|bcftools)\b',
        re.IGNORECASE
    )
    tools_found = sorted(set(m.group(1).lower() for m in tool_pattern.finditer(content)))
    if not tools_found:
        return findings
    # Check if versions are mentioned near tool names
    unversioned = []
    for tool in tools_found:
        # Look for version number near tool mention
        tool_ctx = re.search(rf'\b{tool}\b.{{0,80}}', content, re.IGNORECASE)
        if tool_ctx:
            ctx = tool_ctx.group(0)
            if not re.search(r'v?\d+\.\d+', ctx):
                unversioned.append(tool)
    if unversioned:
        findings.append(finding(
            'CC', 'SIGNIFICANT',
            f'External tools required but versions not specified: {", ".join(unversioned[:5])}',
            'The README references external tools that must be on PATH, but no '
            'version numbers are specified. Different versions of these tools '
            '(e.g. GATK v3 vs v4) have completely different command-line interfaces '
            'and may produce different results.',
            [f'Unversioned tools: {", ".join(unversioned)}',
             'Fix: specify exact versions in README System Requirements section']
        ))
    return findings

def detect_CA_readme_script_missing(repo_dir, all_files):
    """Failure Mode CA: Script referenced in README does not exist in repository."""
    findings = []
    # Scan README and shell scripts for tool references
    scan_files = [f for f in all_files
                  if f.name.lower() in {'readme.md', 'readme.txt', 'readme.rst'}]
    scan_files += [f for f in all_files if f.suffix.lower() in {'.sh', '.bash'}]
    if not scan_files:
        return findings
    content = '\n'.join(read_file_safe(f) or '' for f in scan_files)
    if not content:
        return findings
    # Find references to script files in README
    script_pattern = re.compile(
        r'(?:python|Rscript|julia|bash|sh|matlab|jupyter)\s+'
        r'([\w/.-]+\.(?:py|r|jl|sh|m|do|ipynb|rmd|qmd))\b',
        re.IGNORECASE
    )
    # Also catch run-order descriptions: 'preprocess.py -> analyse.py'
    # and numbered notebook lists: '1. rainfall_processing.ipynb'
    runorder_pattern = re.compile(
        r'(?:^|\s|->|,)([\w/.-]+\.(?:py|r|jl|sh|m|do|ipynb|rmd|qmd))(?=\s|->|,|$)',
        re.IGNORECASE | re.MULTILINE
    )
    all_file_paths = {str(f.relative_to(repo_dir)).replace('\\', '/') for f in all_files}
    all_file_names = {f.name.lower() for f in all_files}
    # Build stem→actual-name map for extension-mismatch detection
    stem_to_names: dict = {}
    for f in all_files:
        stem_to_names.setdefault(f.stem.lower(), []).append(f.name)

    missing = []
    ext_mismatches = []   # (referenced_name, actual_name)
    for m in list(script_pattern.finditer(content)) + list(runorder_pattern.finditer(content)):
        ref = m.group(1)
        ref_name = ref.split('/')[-1].lower()
        ref_stem = ref_name.rsplit('.', 1)[0]
        ref_ext  = ref_name.rsplit('.', 1)[1] if '.' in ref_name else ''
        # Exact match or fuzzy stem match
        exact_found = ref_name in all_file_names or ref in all_file_paths
        fuzzy_found = any(ref_stem in fname for fname in all_file_names)
        if exact_found or fuzzy_found:
            # Check for extension mismatch: stem matches but extension differs
            if not exact_found and ref_stem in stem_to_names:
                actual_names = stem_to_names[ref_stem]
                for aname in actual_names:
                    a_ext = aname.rsplit('.', 1)[1] if '.' in aname else ''
                    if a_ext and a_ext != ref_ext:
                        ext_mismatches.append((ref.split('/')[-1], aname))
                        break
        else:
            if ref not in missing:
                missing.append(ref)

    if missing:
        findings.append(finding(
            'CA', 'SIGNIFICANT',
            f'Script(s) referenced in README not found in repository: {", ".join(missing)}',
            'The README describes running scripts that do not exist in the repository. '
            'Validators following the README instructions will immediately encounter '
            'file-not-found errors. Either the script was accidentally omitted from '
            'the deposit or the README refers to an outdated filename.',
            [f'Missing: {s}' for s in missing] +
            ['Fix: add the missing script(s) to the repository or update the README']
        ))
    if ext_mismatches:
        findings.append(finding(
            'CA', 'LOW CONFIDENCE',
            f'README references file(s) with wrong extension: '
            f'{", ".join(f"{ref} → {act}" for ref, act in ext_mismatches[:3])}',
            'The README mentions a filename whose stem matches a real file but '
            'the extension differs (e.g. README says session-info.tex but the '
            'file is session-info.txt). Validators may not find the file by '
            'following the README.',
            [f'README says: {ref}  →  actual file: {act}'
             for ref, act in ext_mismatches] +
            ['Fix: update the filename in the README to match the actual extension']
        ))
    return findings

def detect_BZ_matlab_v73_format(repo_dir, all_files):
    """Failure Mode BZ: MATLAB .mat file saved with -v7.3 flag (HDF5) — version compatibility risk."""
    findings = []
    # Must appear in save() call context, not just comments
    # Match save() with -v7.3 flag, but not in comment lines
    v73_pattern = re.compile(r"^[^%\n]*save\s*\([^)]*-v7\.3", re.IGNORECASE | re.MULTILINE)
    flagged = []
    for f in all_files:
        if f.suffix.lower() in {'.m', '.mlx', '.txt', '.md', '.rst'}:
            try:
                content = read_file_safe(f)
                if v73_pattern.search(content):
                    flagged.append(f.name)
            except Exception:
                pass
    if flagged:
        findings.append(finding(
            'BZ', 'SIGNIFICANT',
            'MATLAB data file uses v7.3 (HDF5) format — version compatibility risk',
            'One or more .mat files appear to have been saved with the -v7.3 flag '
            '(HDF5 format). This requires MATLAB R2011b or later to load. '
            'Validators using older versions will be unable to read the data. '
            'Document this version requirement explicitly in your README.',
            [f'Evidence found in: {", ".join(flagged)}',
             'Fix: add "Requires MATLAB R2011b or later" to README System Requirements']
        ))
    return findings

def detect_BY_julia_missing_manifest(repo_dir, all_files):
    """Failure Mode BY: Julia repo has Project.toml but no Manifest.toml."""
    findings = []
    names_lower = {f.name.lower() for f in all_files}
    if 'project.toml' not in names_lower:
        return findings
    if 'manifest.toml' in names_lower:
        return findings
    findings.append(finding(
        'BY', 'SIGNIFICANT',
        'Julia Manifest.toml missing',
        'Project.toml found but no Manifest.toml present. Without a manifest, '
        'julia --project=. -e "using Pkg; Pkg.instantiate()" resolves packages '
        'to the latest compatible versions, not the exact versions used at '
        'publication. Validators may get different package versions than you used.',
        ['Project.toml present — compat bounds specified',
         'Manifest.toml absent — exact versions unspecified',
         'Fix: run julia --project=. -e "using Pkg; Pkg.resolve(); Pkg.instantiate()" '
         'then commit the generated Manifest.toml']
    ))
    return findings

def detect_BX_pluto_empty_manifest(repo_dir, all_files):
    """Failure Mode BX: Pluto notebook has PLUTO_MANIFEST_TOML_CONTENTS but it is empty."""
    findings = []
    for f in all_files:
        if f.suffix.lower() != '.jl':
            continue
        try:
            content = read_file_safe(f)
            if not content:
                continue
            if 'PLUTO_PROJECT_TOML_CONTENTS' not in content:
                continue
            # Check manifest
            manifest_match = re.search(
                r'PLUTO_MANIFEST_TOML_CONTENTS\s*=\s*"([^"]*)"', content, re.DOTALL)
            if manifest_match and len(manifest_match.group(1).strip()) == 0:
                findings.append(finding(
                    'BX', 'SIGNIFICANT',
                    f'Pluto notebook has empty manifest: {f.name}',
                    'PLUTO_MANIFEST_TOML_CONTENTS is present but empty. '
                    'Without a populated manifest, Pluto resolves packages '
                    'to the latest compatible versions rather than the exact '
                    'versions used at publication. Open the notebook in Pluto, '
                    'allow it to resolve dependencies, then save to populate '
                    'the manifest before depositing.',
                    [f'File: {f.name}',
                     'PLUTO_MANIFEST_TOML_CONTENTS = "" (empty)',
                     'Fix: open in Pluto and save — manifest will be populated automatically']
                ))
        except Exception:
            pass
    return findings

def detect_BU_conda_channel_priority(repo_dir, all_files):
    """Failure Mode BU: Conda environment.yml mixes channels without strict priority."""
    findings = []
    env_files = [f for f in all_files
                 if f.name.lower() in {'environment.yml', 'environment.yaml'}]
    for f in env_files:
        txt = read_file_safe(f)
        if not txt or 'channels:' not in txt:
            continue
        import re as _re2
        m = _re2.search(r'channels:\s*\n((?:\s*-[^\n]+\n)*)', txt)
        if not m:
            continue
        channel_lines = _re2.findall(r'-\s*(\S+)', m.group(1))
        if len(channel_lines) < 2:
            continue
        if 'channel_priority: strict' not in txt:
            findings.append(finding(
                'BU', 'SIGNIFICANT',
                f'Conda channels mixed without strict priority in {f.name}',
                f'Mixing channels ({", ".join(channel_lines)}) without '
                f'channel_priority: strict causes non-deterministic package '
                f'resolution. Conda may silently install packages from '
                f'unexpected channels, producing different environments on '
                f'different machines.',
                [f'Channels listed: {", ".join(channel_lines)}',
                 f'Fix: add "channel_priority: strict" to {f.name} above the channels: block']
            ))
    return findings


def detect_BV_shell_no_set_e(repo_dir, all_files):
    """Failure Mode BV: Shell pipeline script has no error handling (set -e missing)."""
    findings = []
    shell_files = [f for f in all_files if f.suffix.lower() in {'.sh', '.bash'}]
    for f in shell_files:
        txt = read_file_safe(f)
        if not txt or not txt.startswith('#!'):
            continue
        lines = [l for l in txt.splitlines() if l.strip() and not l.strip().startswith('#')]
        if len(lines) < 3:
            continue
        if 'set -e' not in txt and 'set -o errexit' not in txt:
            findings.append(finding(
                'BV', 'SIGNIFICANT',
                f'Shell pipeline has no error handling: {f.name}',
                'Without set -e, the pipeline will continue executing '
                'even if a step fails. Later steps may run on missing '
                'or corrupt inputs without any error being raised, '
                'producing silent garbage output.',
                [f'File: {f.name}',
                 'Fix: add "set -e" on the line immediately after the shebang (#!)',
                 'Optionally also add "set -o pipefail" to catch pipeline errors']
            ))
    return findings


def detect_BW_empty_code_files(repo_dir, all_files):
    """Failure Mode BW: Code file is effectively empty (≤5 bytes — likely a stub or missing upload)."""
    findings = []
    _stata_lib_dirs = {'plus', 'personal', 'stbplus'}
    code_files = [
        f for f in all_files
        if f.suffix.lower() in CODE_EXTENSIONS
        and not ('ado' in f.parts and any(p in _stata_lib_dirs for p in f.parts))
    ]
    empty = []
    for f in code_files:
        try:
            if f.stat().st_size <= 5:
                empty.append(f)
        except Exception:
            pass
    if empty:
        findings.append(finding(
            'BW', 'SIGNIFICANT',
            f'Empty or near-empty code file{"s" if len(empty) > 1 else ""} detected '
            f'({len(empty)} file{"s" if len(empty) > 1 else ""})',
            'One or more code files contain fewer than 6 bytes and are '
            'effectively empty stubs. Empty code files cannot reproduce any '
            'results and suggest a missing upload or an incomplete deposit. '
            'Either replace the file with the actual code, or remove it and '
            'note the omission in your README.',
            [f'{f.name} ({f.stat().st_size} byte{"s" if f.stat().st_size != 1 else ""})'
             for f in empty]
        ))
    return findings


def detect_DUP(repo_dir, all_files):
    """Failure Mode DUP: Data files with identical content — likely accidental duplicates."""
    from collections import defaultdict
    import hashlib

    findings = []

    # Only check data files — code duplicates are usually intentional copies / templates.
    candidates = [f for f in all_files if f.suffix.lower() in DATA_EXTENSIONS]

    # Size-based first pass: only hash-compare files that share an exact byte count.
    size_groups = defaultdict(list)
    for f in candidates:
        try:
            size_groups[f.stat().st_size].append(f)
        except Exception:
            pass

    duplicate_groups = []
    for size, files in size_groups.items():
        if len(files) < 2:
            continue
        if size == 0:
            continue  # Empty files aren't meaningful duplicates
        hash_groups = defaultdict(list)
        for f in files:
            try:
                h = hashlib.md5(f.read_bytes()).hexdigest()
                hash_groups[h].append(f)
            except Exception:
                pass
        for dupes in hash_groups.values():
            if len(dupes) > 1:
                duplicate_groups.append(dupes)

    if not duplicate_groups:
        return findings

    evidence = []
    for group in duplicate_groups:
        names = ', '.join(str(f.relative_to(repo_dir)) for f in group)
        evidence.append(f'Identical files: {names}')
        evidence.append(
            'Confirm these are intentional — if one is a previous version, '
            'remove it or document the difference'
        )

    n_extra = sum(len(g) - 1 for g in duplicate_groups)
    file_word = 'file' if n_extra == 1 else 'files'
    findings.append(finding(
        'DUP', 'SIGNIFICANT',
        f'{n_extra} duplicate data {file_word} detected',
        'Two or more data files have identical content. This may indicate '
        'an accidental duplicate deposit, or files that were intended to '
        'differ but do not. Validators cannot know which file to use.',
        evidence
    ))
    return findings


def detect_DZ_double_zipped(repo_dir, all_files):
    """Failure Mode DZ: Repository file tree is duplicated — likely double-zipped archive."""
    findings = []

    # Pattern: parts[i] == parts[i+1] for any consecutive pair in the repo-relative path.
    # The doubled directory can be at any depth — e.g. a single outer wrapper directory
    # (repo_dir/wrapper/name/name/...) means parts[0]='wrapper', parts[1]=parts[2]='name'.
    # Only checking parts[0]==parts[1] misses that common case.
    doubled_name = None
    doubled_depth = None
    inner_count = 0

    for f in all_files:
        try:
            parts = f.relative_to(repo_dir).parts
        except Exception:
            continue
        for i in range(len(parts) - 2):   # need at least a filename after the doubled dir
            if parts[i] == parts[i + 1]:
                if doubled_name is None:
                    doubled_name = parts[i]
                    doubled_depth = i
                if parts[i] == doubled_name:
                    inner_count += 1
                break

    if doubled_name is None or doubled_depth is None or inner_count < 3:
        return findings

    prefix = '/'.join(['…'] * doubled_depth) + ('/' if doubled_depth else '')
    findings.append(finding(
        'DZ', 'SIGNIFICANT',
        'Repository structure appears duplicated — possible double-zipped archive',
        f'Files appear under both `{prefix}{doubled_name}/` and '
        f'`{prefix}{doubled_name}/{doubled_name}/`, suggesting the deposit '
        'archive was zipped twice. Validators will find the same content at '
        'two different paths, causing confusion about which copy is '
        'authoritative.',
        [f'{inner_count} file(s) duplicated at `{prefix}{doubled_name}/{doubled_name}/`',
         f'Fix: unzip the deposit, enter the `{doubled_name}/` folder, '
         'and re-zip its contents directly']
    ))

    return findings


def detect_NZ(repo_dir, all_files):
    """Failure Mode NZ: Archive files nested inside the deposit — packaging anti-pattern.

    Entry points write a .valichord_nested_archives.json sidecar capturing all
    archive files (including .zip files that will later be extracted and deleted).
    Falls back to scanning all_files for archives when no sidecar exists (tests).
    """
    import json as _json

    sidecar = repo_dir / '.valichord_nested_archives.json'
    if sidecar.exists():
        try:
            records = _json.loads(sidecar.read_text(encoding='utf-8'))
        except Exception:
            records = []
    else:
        # Fallback: scan all_files directly (e.g. test environments)
        records = [
            {'path': str(f.relative_to(repo_dir)), 'size': f.stat().st_size}
            for f in all_files if f.suffix.lower() in ARCHIVE_EXTENSIONS
        ]

    if not records:
        return []

    # Drop single-file compressed files (.csv.gz etc) — they are a data delivery
    # format, not nested archives.  [BH] also skips these.
    from pathlib import Path as _Path
    records = [
        r for r in records
        if not _is_single_file_compressed(_Path(r['path']))
    ]
    if not records:
        return []

    # Build a lookup so we can inspect archives still present on disk
    _files_by_relpath = {str(f.relative_to(repo_dir)): f for f in all_files}

    evidence = []
    for r in records:
        size_kb = r['size'] // 1024
        f = _files_by_relpath.get(r['path'])
        if f is not None:
            note = _archive_contents_note(f)
        else:
            # File was extracted and deleted; use pre-scan contents_note if available
            note = r.get('contents_note', '')
        evidence.append(f'`{r["path"]}` ({size_kb:,} KB){note}')

    n = len(records)
    word = 'archive' if n == 1 else 'archives'
    return [finding(
        'NZ', 'SIGNIFICANT',
        f'{n} nested {word} inside the deposit',
        'Archive files inside a repository require validators to manually '
        'extract additional archives before running the code. '
        'Extract the contents and deposit files directly in a subdirectory.',
        evidence
    )]


# ── [HS] Human subjects data ─────────────────────────────────────────────────

# Direct personal identifiers — each alone is a strong signal.
_HS_STRONG_RE = re.compile(
    r'\b('
    r'patient_?id|participant_?id|subject_?id|respondent_?id|'
    r'date_?of_?birth|dob|birth_?date|'
    r'\bname\b|surname|first_?name|last_?name|full_?name|'
    r'postcode|zip_?code|'
    r'diagnosis|condition|medication|treatment|'
    r'nhs_?number|ssn|national_?insurance|passport_?number|'
    r'email|phone_?number|telephone'
    r')\b',
    re.IGNORECASE,
)

# Weaker signals — demographic, location, or partial identifiers.
_HS_WEAK_RE = re.compile(
    r'\b('
    r'\bage\b|gender|\bsex\b|'
    r'ethnicity|\brace\b|'
    r'\baddress\b|\blocation\b'
    r')\b',
    re.IGNORECASE,
)

# README terms that suppress the finding.
# No trailing \b — stems like 'anonymis' must match 'anonymised', 'anonymization' etc.
_HS_SUPPRESS_RE = re.compile(
    r'(?:'
    r'anonymis|anonymiz|de.?identified|deidentified|'
    r'ethics\s+approval|ethical\s+approval|'
    r'\birb\b|'
    r'rec\s+approval|'
    r'\bgdpr\b|'
    r'\bconsent\b'
    r')',
    re.IGNORECASE,
)

# Filename stems that strongly suggest human subjects data (without reading content).
# Catches administrative data conventions (clmt = claimant, US DoL) and
# behavioural/clinical research identifiers (sleep study, survey, wearable, etc.).
# Applied after CamelCase splitting + separator normalisation so that tokens like
# 'ScreenSleep' → 'Screen Sleep' and 'screen_time' → 'screen time' both match.
_HS_FILENAME_RE = re.compile(
    # Leading \b only — no trailing \b — so stems match all inflected forms.
    # e.g. 'depress' matches depression/depressive/depressed;
    #      'narcissi' matches narcissism/narcissistic.
    # Normalised input is space-separated so the leading \b is sufficient to
    # prevent mid-word matches (e.g. won't hit 'embarrassment').
    r'\b('
    # Administrative / benefits data
    r'clmt|claimant|applicant|beneficiar|respondent|'
    r'enrollee|member.?id|client.?id|'
    # Clinical / research participants
    r'participant|patient|subject|'
    # Behavioural research & wearables
    r'sleep|screen.?time|actigraph|wearable|fitbit|'
    r'diary|survey|questionnaire|interview|'
    # Common study identifiers
    r'cohort|longitudinal|followup|follow.?up|'
    # Clinical / symptom data
    r'symptom|clinical|'
    # Psychometric scales & psychological constructs —
    # stems: match narcissism/narcissistic, depression/depressive, anxiety/anxieties, etc.
    r'narcissi|depress|anxiet|wellbeing|well.being|'
    r'ptsd|personalit|cogniti|psychometr|'
    r'traum|schizophreni|autis|adhd|'
    r'bipolar|mental.health|affectiv|'
    # Educational / ability assessment
    r'achievement|iq.score|iq.data|'
    r'\biq\b|'
    # Survey demographics & psychosocial constructs —
    # individual-level variables that confirm survey / panel data
    r'gender|demographic|psychosocial|'
    r'attitude|belong|ideology|politic|'
    r'scale.item|scale.measure|'
    # Environmental / social psychology
    r'climate.attitude|pro.environmental|pro.social'
    r')',
    re.IGNORECASE,
)

# File stems suggesting synthetic/example data — suppress finding.
_HS_SYNTHETIC_RE = re.compile(
    r'\b(synthetic|simulated|dummy|fake|example|sample|mock)\b',
    re.IGNORECASE,
)


def _hs_read_headers(f) -> list:
    """Return column header strings from a tabular file, or [] on failure."""
    ext = f.suffix.lower()
    try:
        if ext in {'.csv', '.tsv', '.tab', '.txt'}:
            with f.open(encoding='utf-8', errors='ignore') as fh:
                first_line = fh.readline()
            delim = '\t' if (ext in {'.tsv', '.tab'}
                             or first_line.count('\t') > first_line.count(',')) else ','
            return [h.strip().strip('"').strip("'")
                    for h in first_line.split(delim) if h.strip()]
        if ext in {'.xlsx', '.xls'}:
            try:
                import openpyxl
                wb = openpyxl.load_workbook(f, read_only=True, data_only=True)
                ws = wb.active
                if ws is not None:
                    row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True), None)
                    wb.close()
                    if row:
                        return [str(c).strip() for c in row if c is not None]
            except Exception:
                pass
        if ext in {'.sav', '.zsav', '.dta', '.sas7bdat'}:
            try:
                import pyreadstat as _prs
                if ext in {'.sav', '.zsav'}:
                    _, meta = _prs.read_sav(str(f), metadataonly=True)
                elif ext == '.dta':
                    _, meta = _prs.read_dta(str(f), metadataonly=True)
                else:
                    _, meta = _prs.read_sas7bdat(str(f), metadataonly=True)
                return list(meta.column_names)
            except Exception:
                pass
    except Exception:
        pass
    return []


def detect_HS_human_subjects_data(repo_dir, all_files):
    """Failure Mode HS: Data files may contain human subjects data without
    documented anonymisation or ethics approval.

    Scans column headers in readable tabular files (.csv, .tsv, .tab, .xlsx)
    for personal or clinical identifier patterns.

    Severity:
      SIGNIFICANT — 2+ strong signals, or 1 strong + 1 weak signal
      LOW CONFIDENCE — 1 strong signal only, or 2+ weak signals only
    Suppressed when any README documents anonymisation/ethics/IRB, or when
    data file names suggest synthetic or example data.
    """
    findings = []

    # Suppress if README mentions anonymisation or ethics — check first (cheap).
    readme_candidates = sorted(
        [f for f in all_files if f.name.lower() in README_NAMES],
        key=lambda x: len(x.parts),
    )
    for rf in readme_candidates:
        if _HS_SUPPRESS_RE.search(read_file_safe(rf)):
            return findings

    # Filename-level signals — catches binary/proprietary data files (.dta, .sav, etc.)
    # whose headers cannot be read but whose names suggest human subjects content.
    # Image extensions included: survey/study images (e.g. "respondents.jpg") are valid signals.
    _DATA_EXTS_HS = (DATA_EXTENSIONS
                     | {'.csv', '.tsv', '.tab', '.xlsx', '.xls'}
                     | {'.jpg', '.jpeg', '.png', '.tif', '.tiff'})
    filename_hits: list = []
    for f in all_files:
        if (f.suffix.lower() in _DATA_EXTS_HS
                and not _HS_SYNTHETIC_RE.search(f.stem)
                and f.name.lower() not in CODEBOOK_FILENAMES):
            # Normalise to space-separated tokens so \b works correctly:
            # 1. Split CamelCase: 'ScreenSleep' → 'Screen Sleep'
            # 2. Replace separator chars: 'pca_bamdetamt_clmt' → 'pca bamdetamt clmt'
            _stem_norm = re.sub(r'([a-z])([A-Z])', r'\1 \2', f.stem)
            _stem_norm = re.sub(r'[._\-]', ' ', _stem_norm)
            if _HS_FILENAME_RE.search(_stem_norm):
                filename_hits.append(f.name)

    # Tabular files to scan for column headers.
    _SCAN_EXTS = {'.csv', '.tsv', '.tab', '.xlsx', '.xls'}
    candidates = [
        f for f in all_files
        if f.suffix.lower() in _SCAN_EXTS
        and not _HS_SYNTHETIC_RE.search(f.stem)
        and not f.name.lower().startswith('readme')
        and f.name.lower() not in CODEBOOK_FILENAMES
    ]

    strong_hits: list = []
    weak_hits: list = []

    for f in candidates:
        for h in _hs_read_headers(f):
            if _HS_STRONG_RE.search(h) and h not in strong_hits:
                strong_hits.append(h)
            elif _HS_WEAK_RE.search(h) and h not in weak_hits:
                weak_hits.append(h)

    n_strong = len(strong_hits)
    n_weak = len(weak_hits)

    if n_strong == 0 and n_weak < 2 and not filename_hits:
        return findings

    severity = ('SIGNIFICANT'
                if n_strong >= 2 or (n_strong >= 1 and n_weak >= 1)
                else 'LOW CONFIDENCE')

    detected = (strong_hits[:5] + weak_hits[:3])[:6]
    detected_str = ', '.join(f'`{h}`' for h in detected)

    _source_desc = ('Column headers' if detected_str
                    else 'Filename patterns')
    _body = (
        f'{_source_desc} suggest this deposit may contain personal or '
        'clinical/administrative data. '
        'Before depositing openly, confirm that: '
        '(1) data has been fully anonymised or de-identified; '
        '(2) you have ethics/IRB approval to share this data; '
        '(3) the anonymisation process is documented in the README; '
        '(4) sharing complies with your institution\'s data governance policy. '
        'If data cannot be fully anonymised, consider a managed-access repository '
        'such as Vivli (clinical trial data) rather than open deposit.'
    )
    _evidence = []
    if detected_str:
        _evidence.append(f'Sensitive headers detected: {detected_str}')
    if filename_hits:
        _evidence.append(
            f'Filename signals: {", ".join(filename_hits[:5])}'
            + (f' (+ {len(filename_hits)-5} more)' if len(filename_hits) > 5 else '')
        )
    _evidence += [
        'README does not mention anonymisation or ethics approval',
        'Action: add anonymisation statement to README, or restrict '
        'access if data cannot be shared openly',
    ]
    findings.append(finding(
        'HS', severity,
        'Data file(s) may contain human subjects data — anonymisation not documented',
        _body,
        _evidence,
    ))
    return findings


# ─────────────────────────────────────────────────────────────────────────────
# [DH] Undeclared imports
# ─────────────────────────────────────────────────────────────────────────────

# Maps Python import names (lowercase) to their canonical pip package names.
# Only covers well-known cases where the import name differs from the pip name.
_IMPORT_ALIAS_TO_PKG: dict[str, str] = {
    'sklearn':       'scikit-learn',
    'cv2':           'opencv-python',
    'pil':           'pillow',
    'bs4':           'beautifulsoup4',
    'yaml':          'pyyaml',
    'skimage':       'scikit-image',
    'bio':           'biopython',
    'dateutil':      'python-dateutil',
    'attr':          'attrs',
    'umap':          'umap-learn',
    'crypto':        'pycryptodome',
    'dotenv':        'python-dotenv',
    'pkg_resources': 'setuptools',
    'serial':        'pyserial',
    'usb':           'pyusb',
    'gi':            'pygobject',
    'wx':            'wxpython',
}


def _parse_py_dep_packages(all_files) -> set[str]:
    """Return lowercase normalised set of declared Python packages.

    Reads requirements.txt and/or environment.yml/yaml.
    Normalises hyphens to underscores so 'scikit-learn' and 'scikit_learn'
    compare equal.
    """
    declared: set[str] = set()
    names_lower = {f.name.lower(): f for f in all_files}

    req = names_lower.get('requirements.txt')
    if req:
        for line in read_file_safe(req).splitlines():
            line = line.strip()
            if not line or line.startswith('#') or line.startswith('-'):
                continue
            name = re.split(r'[\[<>=!;@\s]', line)[0].strip()
            if name:
                declared.add(name.lower().replace('-', '_'))

    for env_name in ('environment.yml', 'environment.yaml'):
        env = names_lower.get(env_name)
        if not env:
            continue
        for line in read_file_safe(env).splitlines():
            stripped = line.strip()
            if not stripped.startswith('- '):
                continue
            pkg_part = stripped[2:].split('#')[0].strip()
            if pkg_part.lower() in ('pip:', 'pip') or pkg_part.startswith('{'):
                continue
            # conda: channel::package=version
            if '::' in pkg_part:
                pkg_part = pkg_part.split('::')[1]
            name = re.split(r'[=<>!\s]', pkg_part)[0].strip()
            if name and '.' not in name[:2]:
                declared.add(name.lower().replace('-', '_'))

    return declared


def detect_DH_undeclared_imports(repo_dir, all_files):
    """Failure Mode DH: Packages imported in code but absent from dependency files.

    Only fires when a dependency file already exists. If none exists, [B] covers
    the gap and DH would be redundant. Covers Python (.py, .ipynb) and R files.

    Python: cross-references imports against requirements.txt / environment.yml.
    R: cross-references library()/require() calls against renv.lock / DESCRIPTION.

    Severity: SIGNIFICANT — missing packages cause ImportError/package-not-found
    at runtime, blocking reproduction outright.
    """
    findings = []
    names_lower = {f.name.lower(): f for f in all_files}

    # ── Python ────────────────────────────────────────────────────────────────
    _PY_DEP_NAMES = {'requirements.txt', 'environment.yml', 'environment.yaml',
                     'pyproject.toml', 'setup.py', 'setup.cfg', 'pipfile'}
    has_py_dep = any(n in names_lower for n in _PY_DEP_NAMES)

    if has_py_dep:
        declared_py = _parse_py_dep_packages(all_files)
        if declared_py:
            local_mods = {f.stem.lower() for f in all_files if f.suffix.lower() == '.py'}
            pkg_dirs = {f.parent for f in all_files if f.name == '__init__.py'}

            py_files = [
                f for f in all_files
                if f.suffix.lower() in {'.py', '.ipynb'}
                and not _is_minified(f)
                and f.parent not in pkg_dirs
                and not any(part.lower() in VENDOR_DIRS
                            for part in f.relative_to(repo_dir).parts)
            ]

            undeclared: set[str] = set()
            for f in py_files:
                for m in _IMPORT_PAT.finditer(read_file_safe(f)):
                    raw = (m.group(1) or m.group(2) or '').lower()
                    if not raw or raw in _STDLIB_TOPLEVEL or raw in local_mods:
                        continue
                    canonical = _IMPORT_ALIAS_TO_PKG.get(raw, raw).replace('-', '_')
                    if canonical not in declared_py:
                        undeclared.add(canonical)

            if len(undeclared) >= 2:
                pkgs = sorted(undeclared)
                dep_src = next(
                    (names_lower[n].name for n in _PY_DEP_NAMES if n in names_lower),
                    'dependency file'
                )
                findings.append(finding(
                    'DH', 'SIGNIFICANT',
                    f'{len(undeclared)} package(s) imported but not in {dep_src}',
                    'These packages are imported in code files but do not appear in '
                    'the dependency specification. A validator installing only the '
                    'declared dependencies will encounter ImportError at runtime.',
                    [f'Undeclared: {", ".join(pkgs[:10])}']
                    + (['(list truncated)'] if len(pkgs) > 10 else [])
                    + [f'Dependency file: {dep_src}']
                ))

    # ── R ─────────────────────────────────────────────────────────────────────
    renv = names_lower.get('renv.lock')
    desc = next(
        (f for f in all_files
         if f.name.lower() == 'description'
         and len(f.relative_to(repo_dir).parts) <= 2),
        None
    )

    if renv or desc:
        declared_r: set[str] = set()
        if renv:
            try:
                import json as _json_dh
                data = _json_dh.loads(read_file_safe(renv))
                declared_r = {k.lower() for k in data.get('Packages', {})}
            except Exception:
                pass

        if desc and not declared_r:
            in_imports = False
            for line in read_file_safe(desc).splitlines():
                if re.match(r'^Imports\s*:', line, re.IGNORECASE):
                    in_imports = True
                    rest = line.split(':', 1)[1]
                    for pkg in re.findall(r'\b([A-Za-z][A-Za-z0-9.]+)', rest):
                        declared_r.add(pkg.lower())
                elif in_imports:
                    if line[:1] in (' ', '\t'):
                        for pkg in re.findall(r'\b([A-Za-z][A-Za-z0-9.]+)', line):
                            declared_r.add(pkg.lower())
                    else:
                        in_imports = False

        if declared_r:
            _BASE_R = frozenset({
                'base', 'stats', 'utils', 'methods', 'graphics',
                'grdevices', 'datasets', 'tools', 'grid',
                'parallel', 'compiler', 'splines', 'tcltk',
            })
            _R_LIB_PAT = re.compile(
                r'(?:library|require)\s*\(\s*["\']?([A-Za-z][A-Za-z0-9._]+)',
                re.IGNORECASE
            )
            undeclared_r: set[str] = set()
            for f in _researcher_r_files(all_files, repo_dir):
                for m in _R_LIB_PAT.finditer(read_file_safe(f)):
                    pkg = m.group(1)
                    if pkg.lower() not in _BASE_R and pkg.lower() not in declared_r:
                        undeclared_r.add(pkg)

            if undeclared_r:
                dep_name = renv.name if renv else (desc.name if desc else 'DESCRIPTION')
                pkgs = sorted(undeclared_r)
                findings.append(finding(
                    'DH', 'SIGNIFICANT',
                    f'{len(pkgs)} R package(s) used but absent from {dep_name}',
                    'These R packages are called via library() or require() but are '
                    'not listed in the dependency file. Running renv::restore() or '
                    'reading DESCRIPTION will not install them.',
                    [f'Undeclared R packages: {", ".join(pkgs[:10])}']
                    + (['(list truncated)'] if len(pkgs) > 10 else [])
                    + [f'Dependency file: {dep_name}']
                ))

    return findings


# ─────────────────────────────────────────────────────────────────────────────
# [DI] README variable names vs data column headers
# ─────────────────────────────────────────────────────────────────────────────

# README section headers that signal an explicit variable/column documentation block.
_DI_VAR_SECTION_RE = re.compile(
    r'^#{1,4}\s*(?:variables?|columns?|data\s+dict(?:ionary)?'
    r'|field\s+descriptions?|column\s+descriptions?'
    r'|data\s+description|variable\s+descriptions?)',
    re.IGNORECASE | re.MULTILINE
)

# Numbered list with parenthesized code name:
# "1. Forest coverage rate (forest_coverage_rate)"
_DI_NUMBERED_VAR_RE = re.compile(
    r'^\s*\d+\.\s+\w[^\n]*\(([a-zA-Z_][a-zA-Z0-9_]{2,})\)',
    re.MULTILINE
)

# Backtick-quoted identifiers: `column_name`
_DI_BACKTICK_RE = re.compile(r'`([a-zA-Z_][a-zA-Z0-9_]{2,})`')


def _di_looks_like_col_name(s: str) -> bool:
    """True if s resembles a data column name rather than a code expression or command."""
    if any(c in s for c in '()./ \\[]{}:,'):
        return False
    # Require snake_case (underscore present) or ALL_CAPS — bare words like `age`
    # are too ambiguous to treat as column names without additional context.
    return '_' in s or (s.isupper() and len(s) >= 3)


def detect_DI_variable_mismatch(repo_dir, all_files):
    """Failure Mode DI: README-documented variable names not found in data columns.

    Only fires when:
    1. A README contains an explicit variables/columns section header.
    2. That section documents snake_case or parenthesized variable names.
    3. Fewer than 60 % of those names match column headers in the tabular data.
    4. At least 3 names are unmatched (guards against single-alias differences).

    Severity: LOW CONFIDENCE — naming conventions vary legitimately; derived or
    transformed variable names may intentionally differ from raw column names.
    """
    findings = []

    tabular = [
        f for f in all_files
        if f.suffix.lower() in {'.csv', '.tsv', '.tab', '.xlsx', '.xls',
                                 '.dta', '.sav', '.zsav', '.sas7bdat'}
        and not f.name.lower().startswith('readme')
        and f.name.lower() not in CODEBOOK_FILENAMES
        and not _in_asset_dir(f, repo_dir)
    ]
    if not tabular:
        return findings

    # Collect column headers from up to 10 data files (cap avoids slow runs).
    all_headers: set[str] = set()
    for tf in tabular[:10]:
        for h in _hs_read_headers(tf):
            all_headers.add(h.lower().strip())
    if not all_headers:
        return findings

    readme_files = sorted(
        [f for f in all_files if f.name.lower() in README_NAMES],
        key=lambda x: len(x.relative_to(repo_dir).parts)
    )

    readme_vars: set[str] = set()
    for readme in readme_files[:3]:
        try:
            content = readme.read_text(encoding='utf-8', errors='ignore')
        except Exception:
            continue
        if not _DI_VAR_SECTION_RE.search(content):
            continue
        # Split on section headers; examine only the bodies of variable sections.
        parts = _DI_VAR_SECTION_RE.split(content)
        for section_body in parts[1:]:
            next_hdr = re.search(r'^#{1,3}\s+\w', section_body, re.MULTILINE)
            body = section_body[:next_hdr.start()] if next_hdr else section_body
            for m in _DI_NUMBERED_VAR_RE.finditer(body):
                name = m.group(1)
                if _di_looks_like_col_name(name):
                    readme_vars.add(name.lower())
            for m in _DI_BACKTICK_RE.finditer(body):
                name = m.group(1)
                if _di_looks_like_col_name(name):
                    readme_vars.add(name.lower())

    if len(readme_vars) < 3:
        return findings

    def _norm(s: str) -> str:
        return re.sub(r'[\s_\-]+', '_', s.lower().strip())

    norm_headers = {_norm(h) for h in all_headers}
    mismatched = [v for v in sorted(readme_vars) if _norm(v) not in norm_headers]

    if len(mismatched) < 3:
        return findings
    match_rate = (len(readme_vars) - len(mismatched)) / len(readme_vars)
    if match_rate >= 0.6:
        return findings  # majority match — residual differences likely intentional

    findings.append(finding(
        'DI', 'LOW CONFIDENCE',
        f'{len(mismatched)} README variable(s) not found in data column headers',
        'Variable names documented in the README cannot be matched to column '
        'headers in the tabular data files. This may indicate naming '
        'inconsistencies between the documentation and the data, or that the '
        'README documents derived/transformed variables rather than raw column '
        'names. A validator cross-referencing the README against the data will '
        'encounter apparent mismatches.',
        [f'README variables not in data: {", ".join(mismatched[:8])}']
        + (['(list truncated)'] if len(mismatched) > 8 else [])
        + [f'Data files checked: {", ".join(tf.name for tf in tabular[:5])}']
        + (['(more files exist)'] if len(tabular) > 5 else [])
    ))
    return findings


# ── [DJ] Committed build artefacts ───────────────────────────────────────────

_BUILD_ARTEFACT_EXTS = {'.class', '.jar', '.o', '.obj'}
_BUILD_ARTEFACT_DIRS = {
    'target', '__pycache__', 'build', 'bin', 'out', 'dist',
    'cmake-build-debug', 'cmake-build-release', '.gradle',
}
# Directories where bundled JARs are legitimate (lib/, vendor/, libs/)
_JAR_ALLOWED_PARENTS = {'lib', 'libs', 'vendor', 'third_party', 'jars'}


def detect_DJ_build_artefacts(repo_dir, all_files):
    """[DJ] Compiled build artefacts committed to the repository.

    .class, .jar (outside lib/), target/classes/, build/ with compiled output,
    etc. These are machine-specific, version-specific, and should never be
    committed — they prevent clean reproduction on a different machine.
    """
    findings = []
    artefacts = []
    for f in all_files:
        ext = f.suffix.lower()
        parts = [p.lower() for p in f.relative_to(repo_dir).parts]

        # .jar is legitimate inside lib/ / vendor/ / libs/ etc.
        if ext == '.jar' and len(parts) > 1 and parts[0] in _JAR_ALLOWED_PARENTS:
            continue

        # .pyc / .pyo are already caught by [BE] — skip here to avoid duplication
        if ext in {'.pyc', '.pyo'}:
            continue

        if ext in _BUILD_ARTEFACT_EXTS:
            artefacts.append(f)
            continue

        # Flag any file sitting inside a known build output directory
        if any(p in _BUILD_ARTEFACT_DIRS for p in parts[:-1]):
            artefacts.append(f)

    if not artefacts:
        return findings

    class_count = sum(1 for f in artefacts if f.suffix.lower() == '.class')
    jar_count   = sum(1 for f in artefacts if f.suffix.lower() == '.jar')
    other_count = len(artefacts) - class_count - jar_count

    summary_parts = []
    if class_count:
        summary_parts.append(f'{class_count} .class file(s)')
    if jar_count:
        summary_parts.append(f'{jar_count} .jar file(s) outside lib/')
    if other_count:
        summary_parts.append(f'{other_count} other compiled artefact(s)')

    findings.append(finding(
        'DJ', 'SIGNIFICANT',
        f'Compiled build artefacts committed ({", ".join(summary_parts)})',
        'Compiled output files (.class, .jar, .o, etc.) are committed to the '
        'repository. These are machine- and version-specific: a validator on a '
        'different OS, JVM version, or compiler will produce different binaries, '
        'making the build non-reproducible. Add the relevant output directories '
        '(target/, build/, bin/) to .gitignore and remove these files from git '
        'history with `git rm --cached`.',
        [f'Examples: {", ".join(f.name for f in artefacts[:6])}']
        + (['(list truncated)'] if len(artefacts) > 6 else [])
        + ['Fix: echo "target/\\nbuild/\\nbin/\\n*.class" >> .gitignore',
           'Then: git rm -r --cached target/ build/ bin/ 2>/dev/null']
    ))
    return findings


# ── [DK] Filename typo detector ───────────────────────────────────────────────

# Common research/academic words frequently seen in filenames.
# Conservative — only words long enough (≥6 chars) to have unambiguous
# near-neighbours, and common enough to appear in deposits.
_DK_WORDLIST = [
    'analysis', 'results', 'output', 'figure', 'script', 'method',
    'sample', 'filter', 'matrix', 'import', 'export', 'report',
    'summary', 'config', 'parameter', 'calibration', 'validation',
    'training', 'testing', 'preprocessing', 'processing', 'cleaned',
    'merged', 'subset', 'random', 'baseline', 'prediction',
    'classification', 'regression', 'correlation', 'simulation',
    'experiment', 'dataset', 'metadata', 'citation', 'reference',
    'comparison', 'decision', 'selection', 'collection', 'extraction',
    'estimation', 'evaluation', 'distribution', 'threshold', 'response',
    'variable', 'encoding', 'normalisation', 'normalization',
    'aggregation', 'transformation', 'visualisation', 'visualization',
    'statistics', 'statistical',
]


def detect_DK_filename_typos(repo_dir, all_files):
    """[DK] Likely typos in filenames.

    Splits each filename stem into alphabetic tokens of 6+ characters and
    checks each against a curated list of common research terms using
    difflib close-match. Fires LOW CONFIDENCE only.
    """
    import difflib as _difflib

    findings = []
    _tok_pat = re.compile(r'[A-Za-z]{6,}')
    flagged = []
    seen_tokens: set = set()

    for f in all_files:
        tokens = _tok_pat.findall(f.stem)
        for tok in tokens:
            tok_lower = tok.lower()
            if tok_lower in _DK_WORDLIST:
                continue
            if tok_lower in seen_tokens:
                continue
            matches = _difflib.get_close_matches(
                tok_lower, _DK_WORDLIST, n=1, cutoff=0.82
            )
            if matches:
                seen_tokens.add(tok_lower)
                flagged.append((f.name, tok, matches[0]))

    if not flagged:
        return findings

    findings.append(finding(
        'DK', 'LOW CONFIDENCE',
        f'{len(flagged)} filename token(s) may contain typos',
        'One or more filenames contain tokens that closely resemble common '
        'research terms but are not exact matches. This may indicate a typo '
        'that could make scripts harder to read or cause string-matching '
        'failures in code that constructs filenames programmatically. '
        'Review each case — abbreviations and domain-specific terms may be '
        'intentional.',
        [f'"{tok}" in {fname} — did you mean "{sugg}"?'
         for fname, tok, sugg in flagged[:8]]
        + (['(list truncated)'] if len(flagged) > 8 else [])
    ))
    return findings


# ── [DL] MATLAB autosave artefacts (.asv files) ───────────────────────────────

def detect_DL_matlab_autosave(repo_dir, all_files):
    """[DL] MATLAB editor autosave files (.asv) committed to the repository.

    .asv files are created automatically by the MATLAB editor as crash-recovery
    backups. They are transient editor state and should never be in a deposit.
    Their presence suggests the deposit was assembled by copying a live working
    directory rather than preparing a clean release.
    """
    findings = []
    asv_files = [f for f in all_files if f.suffix.lower() == '.asv']
    if not asv_files:
        return findings

    findings.append(finding(
        'DL', 'LOW CONFIDENCE',
        f'{len(asv_files)} MATLAB autosave file(s) committed (.asv)',
        'MATLAB editor autosave files (.asv) are committed to the repository. '
        'These are crash-recovery backups created automatically by the MATLAB '
        'editor — they are not part of the codebase and should not be included '
        'in a deposit. Their presence suggests the deposit was made by copying '
        'a live working directory rather than preparing a clean release. '
        'Remove them and add *.asv to .gitignore.',
        [f'Files: {", ".join(f.name for f in asv_files[:6])}']
        + (['(list truncated)'] if len(asv_files) > 6 else [])
        + ['Fix: git rm --cached *.asv && echo "*.asv" >> .gitignore']
    ))
    return findings

